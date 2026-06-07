# -*- coding: utf-8 -*-

"""DOCX writer for internal static source-page anchored fidelity mode."""

import re

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .analyzer import (
    build_source_page_layout_map,
    clean_visible_text,
    families_for_page,
    family_observation_for_page,
    group_line_items,
    labels_by_source_page,
    static_item_from_observation,
)


def apply_static_anchored_plan(document, plan: dict) -> dict:
    """Apply static source-page header/footer text to a python-docx document."""
    source_analysis = plan.get('source_analysis') or {}
    source_map = plan.get('source_page_layout_map') or build_source_page_layout_map(source_analysis)
    static_items = plan.get('static_items') or []
    page_count = int(source_analysis.get('page_count') or source_map.get('source_page_count') or 0)
    if page_count <= 0:
        return {
            'applied': False,
            'warning_codes': ['empty_source_page_map'],
            'summary': {'sections_written': 0},
        }

    sections = ensure_source_page_sections(document, page_count)
    items_by_page = labels_by_source_page(static_items)
    pages = source_map.get('pages', []) or [
        {'source_page_index': index}
        for index in range(page_count)
    ]

    header_groups_written = 0
    footer_groups_written = 0
    items_written = 0
    for page_entry in pages[:page_count]:
        page_index = page_entry.get('source_page_index', 0)
        section = sections[page_index]
        unlink_section_parts(section)
        header_items = static_items_for_page(
            source_analysis,
            items_by_page,
            page_index,
            'header')
        footer_items = static_items_for_page(
            source_analysis,
            items_by_page,
            page_index,
            'footer')
        header_result = write_static_part(section.header, section, header_items)
        footer_result = write_static_part(section.footer, section, footer_items)
        header_groups_written += header_result['line_groups_written']
        footer_groups_written += footer_result['line_groups_written']
        items_written += header_result['items_written'] + footer_result['items_written']

    return {
        'applied': True,
        'warning_codes': [],
        'summary': {
            'sections_written': page_count,
            'header_line_groups_written': header_groups_written,
            'footer_line_groups_written': footer_groups_written,
            'static_items_written': items_written,
            'page_number_field_generation': 'static_visual_text',
            'word_PAGE_field_count': 0,
        },
    }


def ensure_source_page_sections(document, page_count: int) -> list:
    sections = list(document.sections)
    while len(sections) < page_count:
        document.add_section(WD_SECTION.NEW_PAGE)
        sections = list(document.sections)
    return sections


def unlink_section_parts(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def clear_part(part):
    element = part._element
    for child in list(element):
        element.remove(child)


def static_items_for_page(source_analysis: dict, items_by_page: dict, page_index: int, target_part: str) -> list:
    family_key = 'selected_header_families' if target_part == 'header' else 'selected_footer_families'
    items = []
    for family in families_for_page(source_analysis.get(family_key, []), page_index):
        observation = family_observation_for_page(family, page_index)
        items.append(static_item_from_observation(
            family.get('text', ''),
            observation,
            target_part,
            'family'))
    for label in items_by_page.get(page_index, []):
        if label.get('target_part') == target_part:
            items.append(static_item_from_observation(
                label.get('text', ''),
                label,
                target_part,
                'static_label'))
    return items


def write_static_part(part, section, items: list) -> dict:
    clear_part(part)
    result = {'line_groups_written': 0, 'items_written': 0}
    for group in group_line_items(items):
        group_result = write_static_line_group(part, section, group)
        result['line_groups_written'] += group_result['line_groups_written']
        result['items_written'] += group_result['items_written']
    return result


def write_static_line_group(part, section, group: dict) -> dict:
    items = group.get('items', []) or []
    if not items:
        return {'line_groups_written': 0, 'items_written': 0}
    paragraph = part.add_paragraph()
    normalize_paragraph_spacing(paragraph)
    zones = {item.get('zone', 'left') for item in items}
    if len(zones) == 1:
        zone = next(iter(zones))
        paragraph.alignment = alignment_value(zone)
        write_zone_runs(paragraph, items)
        return {'line_groups_written': 1, 'items_written': len(items)}

    by_zone = {
        'left': [item for item in items if item.get('zone') == 'left'],
        'center': [item for item in items if item.get('zone') == 'center'],
        'right': [item for item in items if item.get('zone') == 'right'],
    }
    if zones == {'left', 'right'}:
        add_right_tab_stop_only(paragraph, section)
        write_zone_runs(paragraph, by_zone['left'])
        paragraph.add_run('\t')
        write_zone_runs(paragraph, by_zone['right'])
    elif zones == {'left', 'center'}:
        add_center_tab_stop_only(paragraph, section)
        write_zone_runs(paragraph, by_zone['left'])
        paragraph.add_run('\t')
        write_zone_runs(paragraph, by_zone['center'])
    elif zones == {'center', 'right'}:
        add_center_and_right_tab_stops(paragraph, section)
        paragraph.add_run('\t')
        write_zone_runs(paragraph, by_zone['center'])
        paragraph.add_run('\t')
        write_zone_runs(paragraph, by_zone['right'])
    else:
        add_center_and_right_tab_stops(paragraph, section)
        if by_zone['left']:
            write_zone_runs(paragraph, by_zone['left'])
        if by_zone['center']:
            paragraph.add_run('\t')
            write_zone_runs(paragraph, by_zone['center'])
        if by_zone['right']:
            paragraph.add_run('\t')
            write_zone_runs(paragraph, by_zone['right'])
    return {'line_groups_written': 1, 'items_written': len(items)}


def write_zone_runs(paragraph, items: list):
    for index, item in enumerate(items or []):
        if index:
            paragraph.add_run(' ')
        run = paragraph.add_run(clean_visible_text(item.get('text', '')))
        apply_simple_style(run, item.get('observation') or item)


def add_center_and_right_tab_stops(paragraph, section):
    width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        int(width / 2),
        WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        width,
        WD_TAB_ALIGNMENT.RIGHT)


def add_center_tab_stop_only(paragraph, section):
    width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        int(width / 2),
        WD_TAB_ALIGNMENT.CENTER)


def add_right_tab_stop_only(paragraph, section):
    width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        width,
        WD_TAB_ALIGNMENT.RIGHT)


def normalize_paragraph_spacing(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = None


def alignment_value(name: str):
    values = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    return values.get(str(name or '').lower(), WD_ALIGN_PARAGRAPH.LEFT)


def apply_simple_style(run, observation: dict):
    font_name = safe_font_name((observation or {}).get('font_name', ''))
    if font_name:
        run.font.name = font_name
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font_size = optional_float((observation or {}).get('font_size'))
    if font_size and font_size > 0:
        run.font.size = Pt(font_size)
    bold = optional_bool((observation or {}).get('bold'))
    if bold is not None:
        run.bold = bold
    italic = optional_bool((observation or {}).get('italic'))
    if italic is not None:
        run.italic = italic
    color = hex_color((observation or {}).get('color'))
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def safe_font_name(value: str) -> str:
    text = str(value or '').strip()
    if not text or len(text) > 80:
        return ''
    if re.search(r'[\x00-\x1f<>]', text):
        return ''
    return text


def optional_float(value):
    if value in ('', None):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def optional_bool(value):
    if isinstance(value, bool):
        return value
    if value in ('', None):
        return None
    text = str(value).strip().lower()
    if text in {'true', '1', 'yes'}:
        return True
    if text in {'false', '0', 'no'}:
        return False
    return None


def hex_color(value) -> str:
    if value in ('', None):
        return ''
    if isinstance(value, int) and 0 <= value <= 0xFFFFFF:
        return f'{value:06X}'
    text = str(value).strip().upper()
    if text.startswith('#'):
        text = text[1:]
    if text.startswith('0X'):
        text = text[2:]
    if len(text) == 6 and re.fullmatch(r'[0-9A-F]{6}', text):
        return text
    return ''
