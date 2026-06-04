'''docx operation methods based on ``python-docx``.'''

import re

from docx.shared import Pt, RGBColor
try:
    # python-docx <= 0.8.11 or python-docx > 1.0.0
    from docx.oxml import OxmlElement, parse_xml, register_element_cls
except ImportError:
    # python-docx >= 1.0.0
    from docx.oxml.parser import OxmlElement, parse_xml, register_element_cls
from docx.oxml.ns import qn, nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_TAB_ALIGNMENT
from docx.image.exceptions import UnrecognizedImageError
from docx.table import _Cell
from docx.opc.constants import RELATIONSHIP_TYPE
from .share import rgb_value
from lxml import etree


# ---------------------------------------------------------
# section and paragraph
# ---------------------------------------------------------
def set_equal_columns(section, num=2, space=0):
    """Set section column count and space. All the columns have same width.

    Args:
        section : ``python-docx`` Section instance.
        num (int): Column count. Defaults to 2.
        space (int, optional): Space between adjacent columns. Unit: Pt. Defaults to 0.
    """
    col = section._sectPr.xpath('./w:cols')[0]
    col.set(qn('w:num'), str(num))
    col.set(qn('w:space'), str(int(20*space))) # basic unit 1/20 Pt


def set_columns(section, width_list:list, space=0):
    """Set section column count and space.

    Args:
        section : ``python-docx`` Section instance.
        width_list (list|tuple): Width of each column.
        space (int, optional): Space between adjacent columns. Unit: Pt. Defaults to 0.
    
    Scheme::

        <w:cols w:num="2" w:space="0" w:equalWidth="0">
            <w:col w:w="2600" w:space="0"/>
            <w:col w:w="7632"/>
        </w:cols>
    """
    cols = section._sectPr.xpath('./w:cols')[0]

    # do nothing if only one column
    if len(width_list)==1:
        # recovery to default column setting in case previous section has multiple columns
        if len(cols)!=1: cols.clear()
        return

    # set multiple column properties
    # clear columns in advance because the latest column setting seems to be inherited
    cols.clear()
    cols.set(qn('w:num'), str(len(width_list)))
    cols.set(qn('w:equalWidth'), '0')

    # insert column with width
    for w in width_list:
        e = OxmlElement('w:col')
        e.set(qn('w:w'), str(int(20*w)))
        e.set(qn('w:space'), str(int(20*space))) # basic unit 1/20 Pt
        cols.append(e)

def delete_paragraph(paragraph):
    '''Delete a paragraph.

    Reference:    
        https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    '''
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


_PAGE_NUMBER_BEHAVIOR_PLACEHOLDER_ONLY = 'placeholder_only'
_PAGE_NUMBER_BEHAVIOR_STATIC_TEXT = 'static_text'
_PAGE_NUMBER_BEHAVIOR_WORD_FIELD = 'word_field'
_PAGE_NUMBER_BEHAVIOR_UNSUPPORTED = 'unsupported'
_PAGE_NUMBER_BEHAVIORS = {
    _PAGE_NUMBER_BEHAVIOR_PLACEHOLDER_ONLY,
    _PAGE_NUMBER_BEHAVIOR_STATIC_TEXT,
    _PAGE_NUMBER_BEHAVIOR_WORD_FIELD,
    _PAGE_NUMBER_BEHAVIOR_UNSUPPORTED,
}
_HEADER_FOOTER_PLAN_ROLE_HEADER = 'header'
_HEADER_FOOTER_PLAN_ROLE_FOOTER = 'footer'
_HEADER_FOOTER_PLAN_ROLE_PAGE_NUMBER = 'page_number'
_HEADER_FOOTER_PLAN_TARGET_HEADER = 'header'
_HEADER_FOOTER_PLAN_TARGET_FOOTER = 'footer'
_HEADER_FOOTER_PLAN_REGION_TOP = 'top'
_HEADER_FOOTER_PLAN_REGION_BODY = 'body'
_HEADER_FOOTER_PLAN_REGION_BOTTOM = 'bottom'


def apply_header_footer_text_plan(
        document,
        plan: dict,
        enabled: bool = False,
        page_number_behavior: str = _PAGE_NUMBER_BEHAVIOR_PLACEHOLDER_ONLY) -> dict:
    '''Apply a simple internal header/footer text plan to a python-docx document.

    This helper is not called by the normal conversion path. It supports only
    simple header/footer content. When the internal plan carries style or
    alignment hints, the helper applies those hints conservatively. Page
    numbers are explicit: placeholder and static-text modes are diagnostic,
    while word_field writes an internal OpenXML PAGE field only when this
    helper is explicitly enabled.
    '''
    sections = list(getattr(document, 'sections', []) or [])
    page_number_behavior = _normalize_page_number_behavior(page_number_behavior)
    if not enabled:
        return {
            'enabled': False,
            'applied': False,
            'policy': 'internal_docx_header_footer_text_plan_only',
            'summary': {
                'section_count': len(sections),
                'header_paragraphs_written': 0,
                'footer_paragraphs_written': 0,
                'page_number_field_generation': 'deferred_placeholder_only',
                'page_number_behavior': page_number_behavior,
                'page_number_fields_written': 0,
                'styled_runs_written': 0,
                'style_properties_applied': 0,
                'alignment_paragraphs_written': 0,
                'paragraph_spacing_normalized_count': 0,
                'page_number_prefix_suffix_runs_written': 0,
                'page_number_start_number': None,
                'page_number_start_applied': False,
                'line_groups_written': 0,
                'grouped_line_item_count': 0,
                'tabbed_paragraphs_written': 0,
                'tab_runs_written': 0,
            },
            'safety_warnings': [],
        }

    plan_sections = list((plan or {}).get('sections', []) or [])
    warnings = []
    plan_safety_warnings = list((plan or {}).get('safety_warnings', []) or [])
    recommendation = (plan or {}).get('recommendation') or {}
    header_footer_policy = (plan or {}).get('header_footer_policy') or {}
    policy_type = header_footer_policy.get('policy_type')
    warnings.extend(_header_footer_plan_role_warnings(plan))
    if page_number_behavior == _PAGE_NUMBER_BEHAVIOR_UNSUPPORTED:
        warnings.append({'type': 'unsupported_page_number_behavior'})
    if policy_type and policy_type != 'default':
        warnings.append({
            'type': 'header_footer_policy_not_supported_for_simple_writer',
            'policy_type': policy_type,
        })
    if header_footer_policy.get('fail_closed'):
        warnings.append({
            'type': 'header_footer_policy_fail_closed',
            'policy_type': policy_type or 'unknown',
        })
    if plan_safety_warnings:
        warnings.append({
            'type': 'header_footer_plan_has_safety_warnings',
            'count': len(plan_safety_warnings),
        })
    if (
            'safe_for_internal_docx_header_footer_experiment' in recommendation and
            not recommendation.get('safe_for_internal_docx_header_footer_experiment')):
        warnings.append({'type': 'header_footer_plan_not_safe_to_apply'})
    if not sections:
        warnings.append({'type': 'docx_document_has_no_sections'})
    if not plan_sections:
        warnings.append({'type': 'header_footer_plan_has_no_sections'})

    header_count = 0
    footer_count = 0
    page_number_fields_written = 0
    page_number_placeholders_written = 0
    styled_runs_written = 0
    style_properties_applied = 0
    alignment_paragraphs_written = 0
    paragraph_spacing_normalized_count = 0
    page_number_prefix_suffix_runs_written = 0
    page_number_start_number = None
    page_number_start_applied = False
    line_groups_written = 0
    grouped_line_item_count = 0
    tabbed_paragraphs_written = 0
    tab_runs_written = 0
    if sections and plan_sections and not warnings:
        section = sections[0]
        section_plan = plan_sections[0]
        header_items = _header_footer_plan_items(
            section_plan,
            'header_items',
            'header_texts')
        footer_items = _header_footer_plan_items(
            section_plan,
            'footer_items',
            'footer_texts')
        page_number_items = _header_footer_plan_items(
            section_plan,
            'page_number_items',
            'page_number_placeholders')
        header_groups = _header_footer_plan_line_groups(
            section_plan,
            'header_line_groups')
        footer_groups = _header_footer_plan_line_groups(
            section_plan,
            'footer_line_groups')
        page_number_start_number = _page_number_start_number(
            page_number_items,
            section_plan)
        if header_groups:
            header_result = _replace_header_footer_part_line_groups(
                section.header,
                header_groups,
                section,
                page_number_behavior,
                page_number_start_number)
        else:
            header_result = _replace_header_footer_part_items(section.header, header_items)
        if footer_groups:
            footer_result = _replace_header_footer_part_line_groups(
                section.footer,
                footer_groups,
                section,
                page_number_behavior,
                page_number_start_number)
        else:
            if page_number_behavior in {
                    _PAGE_NUMBER_BEHAVIOR_PLACEHOLDER_ONLY,
                    _PAGE_NUMBER_BEHAVIOR_STATIC_TEXT}:
                footer_items.extend(page_number_items)
                page_number_placeholders_written = len(page_number_items)
            footer_result = _replace_header_footer_part_items(section.footer, footer_items)
        header_count = header_result['paragraphs_written']
        footer_count = footer_result['paragraphs_written']
        page_number_fields_written += (
            header_result['fields_written'] +
            footer_result['fields_written'])
        page_number_placeholders_written += (
            header_result['page_number_placeholders_written'] +
            footer_result['page_number_placeholders_written'])
        styled_runs_written += (
            header_result['styled_runs_written'] +
            footer_result['styled_runs_written'])
        style_properties_applied += (
            header_result['style_properties_applied'] +
            footer_result['style_properties_applied'])
        alignment_paragraphs_written += (
            header_result['alignment_paragraphs_written'] +
            footer_result['alignment_paragraphs_written'])
        paragraph_spacing_normalized_count += (
            header_result['paragraph_spacing_normalized_count'] +
            footer_result['paragraph_spacing_normalized_count'])
        page_number_prefix_suffix_runs_written += (
            header_result['prefix_suffix_runs_written'] +
            footer_result['prefix_suffix_runs_written'])
        line_groups_written += (
            header_result['line_groups_written'] +
            footer_result['line_groups_written'])
        grouped_line_item_count += (
            header_result['grouped_line_item_count'] +
            footer_result['grouped_line_item_count'])
        tabbed_paragraphs_written += (
            header_result['tabbed_paragraphs_written'] +
            footer_result['tabbed_paragraphs_written'])
        tab_runs_written += (
            header_result['tab_runs_written'] +
            footer_result['tab_runs_written'])
        if page_number_behavior == _PAGE_NUMBER_BEHAVIOR_WORD_FIELD:
            page_number_start_applied = _set_section_page_number_start(
                section,
                page_number_start_number)
            if not footer_groups:
                page_number_result = _append_page_number_fields(
                    section.footer,
                    page_number_items,
                    page_number_start_number)
                page_number_fields_written += page_number_result['fields_written']
                footer_count += page_number_result['fields_written']
                styled_runs_written += page_number_result['styled_runs_written']
                style_properties_applied += page_number_result['style_properties_applied']
                alignment_paragraphs_written += (
                    page_number_result['alignment_paragraphs_written'])
                paragraph_spacing_normalized_count += (
                    page_number_result['paragraph_spacing_normalized_count'])
                page_number_prefix_suffix_runs_written += (
                    page_number_result['prefix_suffix_runs_written'])

    return {
        'enabled': True,
        'applied': not warnings,
        'policy': 'internal_docx_header_footer_text_plan_only',
        'summary': {
            'section_count': len(sections),
            'plan_section_count': len(plan_sections),
            'header_paragraphs_written': header_count,
            'footer_paragraphs_written': footer_count,
            'page_number_behavior': page_number_behavior,
            'page_number_field_generation': _page_number_field_generation(
                page_number_behavior),
            'page_number_fields_written': page_number_fields_written,
            'page_number_placeholders_written': page_number_placeholders_written,
            'styled_runs_written': styled_runs_written,
            'style_properties_applied': style_properties_applied,
            'alignment_paragraphs_written': alignment_paragraphs_written,
            'paragraph_spacing_normalized_count': paragraph_spacing_normalized_count,
            'page_number_prefix_suffix_runs_written': page_number_prefix_suffix_runs_written,
            'page_number_start_number': page_number_start_number,
            'page_number_start_applied': page_number_start_applied,
            'line_groups_written': line_groups_written,
            'grouped_line_item_count': grouped_line_item_count,
            'tabbed_paragraphs_written': tabbed_paragraphs_written,
            'tab_runs_written': tab_runs_written,
        },
        'safety_warnings': warnings,
        'plan_safety_warnings': plan_safety_warnings,
    }


def _normalize_page_number_behavior(value: str) -> str:
    behavior = str(value or _PAGE_NUMBER_BEHAVIOR_PLACEHOLDER_ONLY).strip().lower()
    return behavior if behavior in _PAGE_NUMBER_BEHAVIORS else _PAGE_NUMBER_BEHAVIOR_UNSUPPORTED


def _page_number_field_generation(behavior: str) -> str:
    if behavior == _PAGE_NUMBER_BEHAVIOR_WORD_FIELD:
        return 'word_field'
    if behavior == _PAGE_NUMBER_BEHAVIOR_STATIC_TEXT:
        return 'static_text_diagnostic_only'
    if behavior == _PAGE_NUMBER_BEHAVIOR_UNSUPPORTED:
        return 'unsupported'
    return 'deferred_placeholder_only'


def _header_footer_plan_texts(section_plan: dict, key: str) -> list:
    values = []
    for value in section_plan.get(key, []) or []:
        text = str(value).strip()
        if text:
            values.append(text)
    return values


def _header_footer_plan_items(
        section_plan: dict,
        item_key: str,
        text_key: str) -> list:
    values = []
    raw_items = section_plan.get(item_key, []) or []
    if raw_items:
        for value in raw_items:
            item = _normalize_header_footer_plan_item(value)
            if item.get('text'):
                values.append(item)
        return values

    return [
        {'text': text}
        for text in _header_footer_plan_texts(section_plan, text_key)
    ]


def _header_footer_plan_line_groups(section_plan: dict, key: str) -> list:
    groups = []
    for group in section_plan.get(key, []) or []:
        if not isinstance(group, dict):
            continue
        items = [
            _normalize_header_footer_plan_item(item)
            for item in group.get('items', []) or []
        ]
        items = [item for item in items if item.get('text')]
        if not items:
            continue
        copied = dict(group)
        copied['items'] = items
        groups.append(copied)
    return groups


def _normalize_header_footer_plan_item(value) -> dict:
    if isinstance(value, dict):
        item = dict(value)
        style = item.get('style') or {}
        if isinstance(style, dict):
            item.setdefault('font_name', style.get('font_name', ''))
            item.setdefault('font_size', style.get('font_size'))
            item.setdefault('bold', style.get('bold'))
            item.setdefault('italic', style.get('italic'))
            item.setdefault('color', style.get('color'))
        item['text'] = str(item.get('text', '')).strip()
        item['alignment'] = str(item.get('alignment', '')).strip().lower()
        if isinstance(item.get('page_number_template'), dict):
            item['page_number_template'] = dict(item.get('page_number_template') or {})
        else:
            item['page_number_template'] = {}
        return item
    return {'text': str(value or '').strip()}


def _header_footer_plan_role_warnings(plan: dict) -> list:
    warnings = []
    for entry in (plan or {}).get('entries', []) or []:
        role = str(entry.get('role', '')).strip().lower()
        target_part = str(entry.get('target_part', '')).strip().lower()
        regions = {
            str(region).strip().lower()
            for region in entry.get('regions', []) or []
            if str(region).strip()
        }
        if role == _HEADER_FOOTER_PLAN_ROLE_HEADER:
            if target_part and target_part != _HEADER_FOOTER_PLAN_TARGET_HEADER:
                warnings.append(_header_footer_plan_entry_warning(
                    entry,
                    'header_entry_target_part_mismatch'))
            if _header_footer_plan_regions_mismatch(
                    regions,
                    required=_HEADER_FOOTER_PLAN_REGION_TOP,
                    forbidden={
                        _HEADER_FOOTER_PLAN_REGION_BODY,
                        _HEADER_FOOTER_PLAN_REGION_BOTTOM,
                    }):
                warnings.append(_header_footer_plan_entry_warning(
                    entry,
                    'header_entry_region_mismatch'))
        elif role == _HEADER_FOOTER_PLAN_ROLE_FOOTER:
            if target_part and target_part != _HEADER_FOOTER_PLAN_TARGET_FOOTER:
                warnings.append(_header_footer_plan_entry_warning(
                    entry,
                    'footer_entry_target_part_mismatch'))
            if _header_footer_plan_regions_mismatch(
                    regions,
                    required=_HEADER_FOOTER_PLAN_REGION_BOTTOM,
                    forbidden={
                        _HEADER_FOOTER_PLAN_REGION_BODY,
                        _HEADER_FOOTER_PLAN_REGION_TOP,
                    }):
                warnings.append(_header_footer_plan_entry_warning(
                    entry,
                    'footer_entry_region_mismatch'))
        elif role == _HEADER_FOOTER_PLAN_ROLE_PAGE_NUMBER:
            if target_part and target_part != _HEADER_FOOTER_PLAN_TARGET_FOOTER:
                warnings.append(_header_footer_plan_entry_warning(
                    entry,
                    'page_number_entry_target_part_mismatch'))
            if _header_footer_plan_regions_mismatch(
                    regions,
                    required=_HEADER_FOOTER_PLAN_REGION_BOTTOM,
                    forbidden={
                        _HEADER_FOOTER_PLAN_REGION_BODY,
                        _HEADER_FOOTER_PLAN_REGION_TOP,
                    }):
                warnings.append(_header_footer_plan_entry_warning(
                    entry,
                    'page_number_entry_region_mismatch'))
    return warnings


def _header_footer_plan_regions_mismatch(
        regions: set,
        required: str,
        forbidden: set) -> bool:
    return required not in regions or bool(regions.intersection(forbidden))


def _header_footer_plan_entry_warning(entry: dict, warning_type: str) -> dict:
    return {
        'type': warning_type,
        'candidate_id': entry.get('candidate_id', ''),
        'role': entry.get('role', ''),
        'target_part': entry.get('target_part', ''),
        'regions': sorted({
            str(region).strip().lower()
            for region in entry.get('regions', []) or []
            if str(region).strip()
        }),
    }


def _append_page_number_fields(
        part,
        items: list,
        page_number_start_number=None) -> dict:
    result = _header_footer_write_result()
    for value in items or []:
        item = _normalize_header_footer_plan_item(value)
        if not item.get('text'):
            continue
        paragraph = part.add_paragraph()
        result['paragraph_spacing_normalized_count'] += (
            _prepare_header_footer_paragraph(paragraph))
        result['alignment_paragraphs_written'] += (
            _apply_header_footer_paragraph_alignment(paragraph, item))
        template = _page_number_template(item)
        prefix = template.get('prefix', '')
        suffix = template.get('suffix', '')
        if prefix:
            prefix_run = paragraph.add_run(prefix)
            style_count = _apply_header_footer_run_style(prefix_run, item)
            if style_count:
                result['styled_runs_written'] += 1
                result['style_properties_applied'] += style_count
            result['prefix_suffix_runs_written'] += 1
        run = add_page_number_field(
            paragraph,
            display_text=str(page_number_start_number or 1))
        style_count = _apply_header_footer_run_style(run, item)
        if style_count:
            result['styled_runs_written'] += 1
            result['style_properties_applied'] += style_count
        if suffix:
            suffix_run = paragraph.add_run(suffix)
            style_count = _apply_header_footer_run_style(suffix_run, item)
            if style_count:
                result['styled_runs_written'] += 1
                result['style_properties_applied'] += style_count
            result['prefix_suffix_runs_written'] += 1
        result['fields_written'] += 1
    return result


def add_page_number_field(paragraph, display_text='1'):
    '''Append an internal Word PAGE field to a paragraph.

    This helper is intentionally low-level and internal. It is not used by the
    default converter path; tests inspect the resulting OpenXML before any
    public migration path can depend on it.
    '''
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = str(display_text or '1')
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    return run


def _replace_header_footer_part_text(part, texts: list) -> int:
    return _replace_header_footer_part_items(
        part,
        [{'text': text} for text in texts or []])['paragraphs_written']


def _replace_header_footer_part_items(part, items: list) -> dict:
    result = _header_footer_write_result()
    paragraphs = list(part.paragraphs)
    for index, value in enumerate(items or []):
        item = _normalize_header_footer_plan_item(value)
        text = item.get('text', '')
        if not text:
            continue
        paragraph = paragraphs[index] if index < len(paragraphs) else part.add_paragraph()
        _clear_paragraph_text(paragraph)
        result['paragraph_spacing_normalized_count'] += (
            _prepare_header_footer_paragraph(paragraph))
        result['alignment_paragraphs_written'] += (
            _apply_header_footer_paragraph_alignment(paragraph, item))
        run = paragraph.add_run(text)
        style_count = _apply_header_footer_run_style(run, item)
        if style_count:
            result['styled_runs_written'] += 1
            result['style_properties_applied'] += style_count
        result['paragraphs_written'] += 1
    return result


def _replace_header_footer_part_line_groups(
        part,
        line_groups: list,
        section,
        page_number_behavior: str,
        page_number_start_number=None) -> dict:
    result = _header_footer_write_result()
    paragraphs = list(part.paragraphs)
    for index, group in enumerate(line_groups or []):
        items = [
            _normalize_header_footer_plan_item(item)
            for item in group.get('items', []) or []
        ]
        items = [item for item in items if item.get('text')]
        if not items:
            continue
        paragraph = paragraphs[index] if index < len(paragraphs) else part.add_paragraph()
        _clear_paragraph_text(paragraph)
        result['paragraph_spacing_normalized_count'] += (
            _prepare_header_footer_paragraph(paragraph))
        use_tab_layout = _line_group_needs_tab_layout(items)
        if use_tab_layout:
            tab_count = _apply_header_footer_line_group_tabs(
                paragraph,
                section,
                items)
            if tab_count:
                result['tabbed_paragraphs_written'] += 1
        else:
            tab_count = 0
            result['alignment_paragraphs_written'] += (
                _apply_header_footer_paragraph_alignment(
                    paragraph,
                    _line_group_paragraph_alignment_item(items)))
        result['line_groups_written'] += 1
        result['grouped_line_item_count'] += len(items)
        previous_alignment = None
        wrote_any = False
        for item in _ordered_line_group_items(items):
            if use_tab_layout:
                tab_runs = _tabs_before_line_group_item(
                    paragraph,
                    item,
                    previous_alignment,
                    wrote_any)
            else:
                tab_runs = _space_before_line_group_item(paragraph, wrote_any)
            result['tab_runs_written'] += tab_runs
            item_result = _write_header_footer_item_runs(
                paragraph,
                item,
                page_number_behavior,
                page_number_start_number)
            _merge_header_footer_write_result(result, item_result)
            previous_alignment = _line_group_item_alignment(item)
            wrote_any = True
        result['paragraphs_written'] += 1
    return result


def _header_footer_write_result() -> dict:
    return {
        'paragraphs_written': 0,
        'fields_written': 0,
        'styled_runs_written': 0,
        'style_properties_applied': 0,
        'alignment_paragraphs_written': 0,
        'paragraph_spacing_normalized_count': 0,
        'prefix_suffix_runs_written': 0,
        'page_number_placeholders_written': 0,
        'line_groups_written': 0,
        'grouped_line_item_count': 0,
        'tabbed_paragraphs_written': 0,
        'tab_runs_written': 0,
    }


def _clear_paragraph_text(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn('w:pPr'):
            paragraph._p.remove(child)


def _prepare_header_footer_paragraph(paragraph) -> int:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = None
    return 1


def _apply_header_footer_line_group_tabs(paragraph, section, items: list) -> int:
    alignments = {
        _line_group_item_alignment(item)
        for item in items or []
    }
    tab_count = 0
    writable_width = _section_writable_width(section)
    if writable_width <= 0:
        return 0
    tab_stops = paragraph.paragraph_format.tab_stops
    if 'center' in alignments:
        tab_stops.add_tab_stop(int(writable_width / 2), WD_TAB_ALIGNMENT.CENTER)
        tab_count += 1
    if 'right' in alignments:
        tab_stops.add_tab_stop(int(writable_width), WD_TAB_ALIGNMENT.RIGHT)
        tab_count += 1
    if tab_count:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return tab_count


def _section_writable_width(section) -> int:
    try:
        return int(section.page_width - section.left_margin - section.right_margin)
    except Exception:
        return int(Pt(468))


def _line_group_needs_tab_layout(items: list) -> bool:
    if len(items or []) <= 1:
        return False
    zones = {
        _line_group_item_alignment_zone(item)
        for item in items or []
    }
    return len(zones) > 1


def _line_group_paragraph_alignment_item(items: list) -> dict:
    alignments = [
        _line_group_item_alignment(item)
        for item in items or []
    ]
    known = [
        alignment
        for alignment in alignments
        if alignment in {'left', 'center', 'right'}
    ]
    if not known:
        return {}
    if len(set(known)) == 1:
        return {'alignment': known[0]}
    return {'alignment': 'left'}


def _line_group_item_alignment_zone(item: dict) -> str:
    alignment = _line_group_item_alignment(item)
    return alignment if alignment in {'center', 'right'} else 'left'


def _ordered_line_group_items(items: list) -> list:
    return sorted(
        items or [],
        key=lambda item: (
            _alignment_sort_key(_line_group_item_alignment(item)),
            _header_footer_optional_float(item.get('x_center')) or 0.0,
            str(item.get('text', ''))))


def _alignment_sort_key(alignment: str) -> int:
    return {
        'left': 0,
        'unknown': 0,
        'center': 1,
        'right': 2,
    }.get(alignment, 0)


def _line_group_item_alignment(item: dict) -> str:
    alignment = str((item or {}).get('alignment', '')).strip().lower()
    return alignment if alignment in {'left', 'center', 'right'} else 'unknown'


def _tabs_before_line_group_item(
        paragraph,
        item: dict,
        previous_alignment: str,
        wrote_any: bool) -> int:
    alignment = _line_group_item_alignment(item)
    if alignment == 'center':
        paragraph.add_run().add_tab()
        return 1
    if alignment == 'right':
        paragraph.add_run().add_tab()
        return 1
    if wrote_any and previous_alignment in {'left', 'unknown'}:
        paragraph.add_run(' ')
    return 0


def _space_before_line_group_item(paragraph, wrote_any: bool) -> int:
    if wrote_any:
        paragraph.add_run(' ')
    return 0


def _write_header_footer_item_runs(
        paragraph,
        item: dict,
        page_number_behavior: str,
        page_number_start_number=None) -> dict:
    result = _header_footer_write_result()
    role = str((item or {}).get('role', '')).strip().lower()
    text = str((item or {}).get('text', '')).strip()
    if role == _HEADER_FOOTER_PLAN_ROLE_PAGE_NUMBER:
        if page_number_behavior == _PAGE_NUMBER_BEHAVIOR_WORD_FIELD:
            template = _page_number_template(item)
            prefix = template.get('prefix', '')
            suffix = template.get('suffix', '')
            if prefix:
                _write_styled_text_run(paragraph, prefix, item, result)
                result['prefix_suffix_runs_written'] += 1
            run = add_page_number_field(
                paragraph,
                display_text=str(page_number_start_number or 1))
            style_count = _apply_header_footer_run_style(run, item)
            if style_count:
                result['styled_runs_written'] += 1
                result['style_properties_applied'] += style_count
            result['fields_written'] += 1
            if suffix:
                _write_styled_text_run(paragraph, suffix, item, result)
                result['prefix_suffix_runs_written'] += 1
            return result
        _write_styled_text_run(paragraph, text, item, result)
        result['page_number_placeholders_written'] += 1
        return result

    _write_styled_text_run(paragraph, text, item, result)
    return result


def _write_styled_text_run(paragraph, text: str, item: dict, result: dict):
    run = paragraph.add_run(text)
    style_count = _apply_header_footer_run_style(run, item)
    if style_count:
        result['styled_runs_written'] += 1
        result['style_properties_applied'] += style_count


def _merge_header_footer_write_result(target: dict, source: dict):
    for key in source or {}:
        if key in target and isinstance(target[key], int):
            target[key] += source[key]


def _apply_header_footer_paragraph_alignment(paragraph, item: dict) -> int:
    alignment = str((item or {}).get('alignment', '')).strip().lower()
    values = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    if alignment not in values:
        return 0
    paragraph.alignment = values[alignment]
    return 1


def _apply_header_footer_run_style(run, item: dict) -> int:
    count = 0
    font_name = _safe_header_footer_font_name((item or {}).get('font_name', ''))
    if font_name:
        run.font.name = font_name
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        count += 1

    font_size = _header_footer_optional_float((item or {}).get('font_size'))
    if font_size and font_size > 0:
        run.font.size = Pt(font_size)
        count += 1

    bold = _header_footer_optional_bool((item or {}).get('bold'))
    if bold is not None:
        run.bold = bold
        count += 1

    italic = _header_footer_optional_bool((item or {}).get('italic'))
    if italic is not None:
        run.italic = italic
        count += 1

    color = _header_footer_hex_color((item or {}).get('color'))
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
        count += 1
    return count


def _safe_header_footer_font_name(value: str) -> str:
    text = str(value or '').strip()
    if not text or len(text) > 80:
        return ''
    if re.search(r'[\x00-\x1f<>]', text):
        return ''
    return text


def _header_footer_optional_float(value):
    if value in ('', None):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _header_footer_optional_bool(value):
    if isinstance(value, bool):
        return value
    if value in ('', None):
        return None
    text = str(value).strip().lower()
    if text in {'true', '1', 'yes'}:
        return True
    if text in {'false', '0', 'no'}:
        return False
    return None


def _header_footer_hex_color(value) -> str:
    if value in ('', None):
        return ''
    if isinstance(value, int) and 0 <= value <= 0xFFFFFF:
        return f'{value:06X}'
    text = str(value).strip().upper()
    if text.startswith('#'):
        text = text[1:]
    if text.startswith('0X'):
        text = text[2:]
    if len(text) == 6 and re.fullmatch(r'[0-9A-F]{6}', text):
        return text
    return ''


def _page_number_template(item: dict) -> dict:
    template = (item or {}).get('page_number_template') or {}
    if not isinstance(template, dict):
        return {}
    if template.get('supported') and template.get('consecutive'):
        return template
    return {}


def _page_number_start_number(items: list, section_plan: dict = None):
    values = []
    for value in items or []:
        item = _normalize_header_footer_plan_item(value)
        template = _page_number_template(item)
        start = _header_footer_optional_int(template.get('start_number'))
        if start and start > 0:
            values.append(start)
    section_start = _header_footer_optional_int(
        (section_plan or {}).get('page_number_start_number'))
    if section_start and section_start > 0:
        values.append(section_start)
    if len(set(values)) == 1:
        return values[0]
    return None


def _header_footer_optional_int(value):
    if value in ('', None):
        return None
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return None


def _set_section_page_number_start(section, start_number) -> bool:
    start = _header_footer_optional_int(start_number)
    if not start or start <= 0:
        return False
    sect_pr = getattr(section, '_sectPr', None)
    if sect_pr is None:
        return False
    pg_num_type = sect_pr.find(qn('w:pgNumType'))
    if pg_num_type is None:
        pg_num_type = OxmlElement('w:pgNumType')
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn('w:start'), str(start))
    return True


def reset_paragraph_format(p, line_spacing:float=1.05):
    '''Reset paragraph format, especially line spacing.

    Two kinds of line spacing, corresponding to the setting in MS Office Word:

    * line_spacing=1.05: single or multiple
    * line_spacing=Pt(1): exactly
    
    Args:
        p (Paragraph): ``python-docx`` paragraph instance.
        line_spacing (float, optional): Line spacing. Defaults to 1.05.
    
    Returns:
        paragraph_format: Paragraph format.
    '''
    pf = p.paragraph_format
    pf.line_spacing = line_spacing # single by default
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.widow_control = True

    # do not adjust spacing between Chinese and Latin/number
    xml = r'<w:autoSpaceDE {} w:val="0"/>'.format(nsdecls('w'))
    p._p.get_or_add_pPr().insert(0, parse_xml(xml))

    xml = r'<w:autoSpaceDN {} w:val="0"/>'.format(nsdecls('w'))
    p._p.get_or_add_pPr().insert(0, parse_xml(xml))

    return pf


def set_hidden_property(p):
    '''Hide paragraph. This method just sets the paragraph property, while the added text must
    be hided explicitly.

        r = p.add_run()
        r.text = "Hidden"
        r.font.hidden = True

    Args:
        p (Paragraph): python-docx created paragraph.
    '''
    pPr = OxmlElement('w:pPr') # paragraph property
    rPr = OxmlElement('w:rPr') # run property
    v = OxmlElement('w:vanish') # hidden
    rPr.append(v)
    pPr.append(rPr)
    p._p.append(rPr)


# ---------------------------------------------------------
# text properties
# ---------------------------------------------------------
def set_char_scaling(p_run, scale:float=1.0):
    '''Set character spacing: scaling. 
    
    Manual operation in MS Word: Font | Advanced | Character Spacing | Scaling.
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        scale (float, optional): scaling factor. Defaults to 1.0.
    '''
    p_run._r.get_or_add_rPr().insert(0, 
        parse_xml(r'<w:w {} w:val="{}"/>'.format(nsdecls('w'), 100*scale)))


def set_char_spacing(p_run, space:float=0.0):
    '''Set character spacing. 
    
    Manual operation in MS Word: Font | Advanced | Character Spacing | Spacing.
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        space (float, optional): Spacing value in Pt. Expand if positive else condense. Defaults to 0.0.
    '''
    p_run._r.get_or_add_rPr().insert(0, 
        parse_xml(r'<w:spacing {} w:val="{}"/>'.format(nsdecls('w'), 20*space)))


def set_char_shading(p_run, srgb:int):
    '''Set character shading color, in case the color is out of highlight color scope.
    
    Reference: 
        http://officeopenxml.com/WPtextShading.php
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        srgb (int): Color value.
    '''
    # try to set highlight first using python-docx built-in method
    # Here give 6/16 of the valid highlight colors
    color_map = {        
        rgb_value((1,0,0)): WD_COLOR_INDEX.RED,
        rgb_value((0,1,0)): WD_COLOR_INDEX.BRIGHT_GREEN,
        rgb_value((0,0,1)): WD_COLOR_INDEX.BLUE,
        rgb_value((1,1,0)): WD_COLOR_INDEX.YELLOW,
        rgb_value((1,0,1)): WD_COLOR_INDEX.PINK,
        rgb_value((0,1,1)): WD_COLOR_INDEX.TURQUOISE
    }
    if srgb in color_map:
        p_run.font.highlight_color = color_map[srgb]

    # set char shading
    else:
        c = hex(srgb)[2:].zfill(6)
        xml = r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), c)
        p_run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def set_char_underline(p_run, srgb:int):
    '''Set underline and color.
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        srgb (int): Color value.
    '''
    c = hex(srgb)[2:].zfill(6)
    xml = r'<w:u {} w:val="single" w:color="{}"/>'.format(nsdecls('w'), c)
    p_run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def add_hyperlink(paragraph, url, text):
    """Create a hyperlink within a paragraph object.

    Reference:

        https://github.com/python-openxml/python-docx/issues/74#issuecomment-215678765

    Args:
        paragraph (Paragraph): ``python-docx`` paragraph adding the hyperlink to.
        url (str): The required url.
        text (str): The text displayed for the url.

    Returns: 
        Run: A Run object containing the hyperlink.
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    hyperlink.set(qn('w:history'), '1')

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Create a w:rStyle element, note this currently does not add the hyperlink style as its not in
    # the default template, I have left it here in case someone uses one that has the style in it
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    return r

EMU_PER_PT = 12700
def pt_to_emu(value_pt):
    return int(value_pt * EMU_PER_PT)


def add_floating_picture_pt(
    paragraph,
    image_path,
    x_pt,
    y_pt,
    width_pt,
    height_pt,
):
    """
    Insert floating image at absolute position using points (pt).

    :param paragraph: target paragraph
    :param image_path: image path
    :param x_pt: horizontal position in pt
    :param y_pt: vertical position in pt
    :param width_pt: width in pt
    :param height_pt: height in pt
    """

    # Convert to EMU
    x_emu = pt_to_emu(x_pt)
    y_emu = pt_to_emu(y_pt)
    width_emu = pt_to_emu(width_pt)
    height_emu = pt_to_emu(height_pt)

    run = paragraph.add_run()
    run.add_picture(image_path)

    drawing = run._r.xpath(".//w:drawing")[0]
    inline = drawing.xpath(".//wp:inline")[0]

    # Update size in wp:extent
    extent = inline.xpath("./wp:extent")[0]
    extent.set("cx", str(width_emu))
    extent.set("cy", str(height_emu))

    # Update size in a:ext (VERY IMPORTANT)
    a_ext = inline.xpath(".//a:ext")[0]
    a_ext.set("cx", str(width_emu))
    a_ext.set("cy", str(height_emu))

    # Extract required elements
    docPr = inline.xpath("./wp:docPr")[0]
    cNvGraphicFramePr = inline.xpath("./wp:cNvGraphicFramePr")[0]
    graphic = inline.xpath("./a:graphic")[0]

    extent_xml = etree.tostring(extent, encoding="unicode")
    docPr_xml = etree.tostring(docPr, encoding="unicode")
    cNvGraphicFramePr_xml = etree.tostring(cNvGraphicFramePr, encoding="unicode")
    graphic_xml = etree.tostring(graphic, encoding="unicode")

    # Build anchor XML
    anchor_xml = f"""
    <wp:anchor {nsdecls('wp','a','pic','r')}
        simplePos="0"
        relativeHeight="0"
        behindDoc="0"
        locked="0"
        layoutInCell="1"
        allowOverlap="1">

        <wp:simplePos x="0" y="0"/>

        <wp:positionH relativeFrom="page">
            <wp:posOffset>{x_emu}</wp:posOffset>
        </wp:positionH>

        <wp:positionV relativeFrom="page">
            <wp:posOffset>{y_emu}</wp:posOffset>
        </wp:positionV>

        {extent_xml}

        <wp:wrapNone/>

        {docPr_xml}
        {cNvGraphicFramePr_xml}
        {graphic_xml}

    </wp:anchor>
    """

    anchor = parse_xml(anchor_xml)

    drawing.remove(inline)
    drawing.append(anchor)


# ---------------------------------------------------------
# image properties
# ---------------------------------------------------------
def add_image(p, image_path_or_stream, x_pos, y_pos, width, height):
    '''Add a floating image to a paragraph at a specific position.

    The image is inserted as a floating picture (not inline) using ``add_floating_picture_pt``,
    anchored to the given paragraph and positioned by ``x_pos`` and ``y_pos`` in Pt.

    Args:
        p (Paragraph): ``python-docx`` paragraph instance.
        image_path_or_stream (str | bytes): Image file path or in‑memory image stream.
        x_pos (float): Horizontal position of the image in Pt, relative to the page.
        y_pos (float): Vertical position of the image in Pt, relative to the page.
        width (float): Image width in Pt.
        height (float): Image height in Pt.
    '''
    docx_span = p.add_run()
    try:
        docx_span.add_picture(image_path_or_stream, width=Pt(width), height=Pt(height))
        '''
        add_floating_picture_pt(
            paragraph=p,
            image_path=image_path_or_stream,
            x_pt=x_pos,      # 2 inches
            y_pt=y_pos,      # 2 inches
            width_pt=width,  # 3 inches
            height_pt=height, # 2 inches
        )
        '''
    except UnrecognizedImageError:
        print('Unrecognized Image.')
        return
    
    # exactly line spacing will destroy image display, so set single line spacing instead
    p.paragraph_format.line_spacing = 1.00


class _CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )

register_element_cls('wp:anchor', _CT_Anchor)


def add_float_image(p, image_path_or_stream, width, pos_x=None, pos_y=None):
    '''Add float image behind text.
    
    Args:
        p (Paragraph): ``python-docx`` Paragraph object this picture belongs to.
        image_path_or_stream (str, bytes): Image path or stream.
        width (float): Displaying width of picture, in unit Pt.
        pos_x (float): X-position (English Metric Units) to the top-left point of page valid region
        pos_y (float): Y-position (English Metric Units) to the top-left point of page valid region
    '''
    run = p.add_run()
    # parameters for picture, e.g. id, name
    rId, image = run.part.get_or_add_image(image_path_or_stream)
    cx, cy = image.scaled_dimensions(Pt(width), None)
    shape_id, filename = run.part.next_id, image.filename
    anchor = _CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, Pt(pos_x), Pt(pos_y))
    run._r.add_drawing(anchor)


# ---------------------------------------------------------
# table properties
# ---------------------------------------------------------
def indent_table(table, indent:float):
    '''Indent a table.
    
    Args:
        table (Table): ``python-docx`` Table object.
        indent (float): Indent value, the basic unit is 1/20 pt.
    '''
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(20*indent)) # basic unit 1/20 pt for openxml 
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)


def set_cell_margins(cell:_Cell, **kwargs):
    '''Set cell margins. Provided values are in twentieths of a point (1/1440 of an inch).
    
    Reference: 

        * https://blog.csdn.net/weixin_44312186/article/details/104944773
        * http://officeopenxml.com/WPtableCellMargins.php
    
    Args:
        cell (_Cell): ``python-docx`` Cell instance you want to modify.
        kwargs (dict): Dict with keys: top, bottom, start, end.
        
    Usage::
    
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)    
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
 
    for m in ['top', 'start', 'bottom', 'end']:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
 
    tcPr.append(tcMar)


def set_cell_shading(cell:_Cell, srgb:int):
    '''Set cell background-color.

    Reference:
        https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color
    
    Args:
        cell (_Cell): ``python-docx`` Cell instance you want to modify
        srgb (int): RGB color value.
    '''
    c = hex(srgb)[2:].zfill(6)
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), c)))


def set_cell_border(cell:_Cell, **kwargs):
    '''Set cell`s border.
    
    Reference:
        * https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
        * https://blog.csdn.net/weixin_44312186/article/details/104944110

    Args:
        cell (_Cell): ``python-docx`` Cell instance you want to modify.
        kwargs (dict): Dict with keys: top, bottom, start, end.

    Usage::
    
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_vertical_cell_direction(cell:_Cell, direction:str='btLr'):
    '''Set vertical text direction for cell.

    Reference:
        https://stackoverflow.com/questions/47738013/how-to-rotate-text-in-table-cells
    
    Args:
        direction (str): Either "tbRl" (top to bottom) or "btLr" (bottom to top).
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)
