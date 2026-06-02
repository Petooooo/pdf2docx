import json
import tempfile
import zipfile
import unittest
from collections import Counter
from importlib import util
from pathlib import Path


MODULE_PATH = Path(__file__).resolve().parents[1] / 'pdf2docx' / 'page' / 'LayoutAnalyzer.py'
SPEC = util.spec_from_file_location('LayoutAnalyzer', MODULE_PATH)
LayoutAnalyzer = util.module_from_spec(SPEC)
# Load pure helpers directly; synthetic PDF tests import PyMuPDF only when available.
SPEC.loader.exec_module(LayoutAnalyzer)

PAGE_NUMBER_PLACEHOLDER = LayoutAnalyzer.PAGE_NUMBER_PLACEHOLDER
REGION_BODY = LayoutAnalyzer.REGION_BODY
REGION_BOTTOM = LayoutAnalyzer.REGION_BOTTOM
REGION_TOP = LayoutAnalyzer.REGION_TOP
IMAGE_PLACEHOLDER = LayoutAnalyzer.IMAGE_PLACEHOLDER
ACTION_KEEP = LayoutAnalyzer.ACTION_KEEP
ACTION_REVIEW = LayoutAnalyzer.ACTION_REVIEW
ACTION_WOULD_EXCLUDE = LayoutAnalyzer.ACTION_WOULD_EXCLUDE
ROLE_FOOTER = LayoutAnalyzer.ROLE_FOOTER
ROLE_HEADER = LayoutAnalyzer.ROLE_HEADER
ROLE_KEEP_BODY = LayoutAnalyzer.ROLE_KEEP_BODY
ROLE_LAYOUT_PLACEHOLDER = LayoutAnalyzer.ROLE_LAYOUT_PLACEHOLDER
ROLE_PAGE_NUMBER = LayoutAnalyzer.ROLE_PAGE_NUMBER
ROLE_REVIEW_ONLY = LayoutAnalyzer.ROLE_REVIEW_ONLY
classify_y_band = LayoutAnalyzer.classify_y_band
find_repeated_text_candidates = LayoutAnalyzer.find_repeated_text_candidates
build_header_footer_exclusion_dry_run = LayoutAnalyzer.build_header_footer_exclusion_dry_run
build_reviewed_filtering_internal_config = LayoutAnalyzer.build_reviewed_filtering_internal_config
build_reviewed_filtering_internal_config_report = LayoutAnalyzer.build_reviewed_filtering_internal_config_report
build_docx_header_footer_generation_plan = LayoutAnalyzer.build_docx_header_footer_generation_plan
build_body_filtering_diff_report = LayoutAnalyzer.build_body_filtering_diff_report
build_body_table_geometry_delta_safety_report = LayoutAnalyzer.build_body_table_geometry_delta_safety_report
build_body_table_delta_root_cause_report = LayoutAnalyzer.build_body_table_delta_root_cause_report
build_table_geometry_visual_approval_gate_report = LayoutAnalyzer.build_table_geometry_visual_approval_gate_report
build_table_geometry_visual_review_pack = LayoutAnalyzer.build_table_geometry_visual_review_pack
build_filtered_docx_generation_comparison_report = LayoutAnalyzer.build_filtered_docx_generation_comparison_report
build_filtered_docx_residual_structure_report = LayoutAnalyzer.build_filtered_docx_residual_structure_report
build_local_corpus_validation_summary_report = LayoutAnalyzer.build_local_corpus_validation_summary_report
build_local_corpus_manual_review_pack = LayoutAnalyzer.build_local_corpus_manual_review_pack
build_local_corpus_manual_review_summary_report = LayoutAnalyzer.build_local_corpus_manual_review_summary_report
build_local_corpus_approval_validation_report = LayoutAnalyzer.build_local_corpus_approval_validation_report
build_local_corpus_approval_validation_summary_report = LayoutAnalyzer.build_local_corpus_approval_validation_summary_report
build_reviewed_filtering_feature_readiness_report = LayoutAnalyzer.build_reviewed_filtering_feature_readiness_report
build_document_parse_copied_raw_page_filtering_apply_report = LayoutAnalyzer.build_document_parse_copied_raw_page_filtering_apply_report
build_document_parse_filtering_hook_report = LayoutAnalyzer.build_document_parse_filtering_hook_report
build_document_parse_filtered_parse_experiment_report = LayoutAnalyzer.build_document_parse_filtered_parse_experiment_report
build_document_parse_guarded_raw_page_apply_restore_report = LayoutAnalyzer.build_document_parse_guarded_raw_page_apply_restore_report
build_document_parse_raw_object_mapping_report = LayoutAnalyzer.build_document_parse_raw_object_mapping_report
build_document_parse_filtering_simulation_report = LayoutAnalyzer.build_document_parse_filtering_simulation_report
build_filter_insertion_point_analysis_report = LayoutAnalyzer.build_filter_insertion_point_analysis_report
build_indentation_rule_comparison_report = LayoutAnalyzer.build_indentation_rule_comparison_report
build_paragraph_integrity_report = LayoutAnalyzer.build_paragraph_integrity_report
build_paragraph_mismatch_analysis_report = LayoutAnalyzer.build_paragraph_mismatch_analysis_report
build_paragraph_production_comparison_report = LayoutAnalyzer.build_paragraph_production_comparison_report
build_paragraph_reconstruction_validation_report = LayoutAnalyzer.build_paragraph_reconstruction_validation_report
build_reviewed_header_footer_filter_report = LayoutAnalyzer.build_reviewed_header_footer_filter_report
build_layout_analysis_report = LayoutAnalyzer.build_layout_analysis_report
build_table_delta_investigation_report = LayoutAnalyzer.build_table_delta_investigation_report
reviewed_filtering_config_to_document_parse_settings = LayoutAnalyzer.reviewed_filtering_config_to_document_parse_settings
find_paragraph_continuation_candidates = LayoutAnalyzer.find_paragraph_continuation_candidates
make_text_fingerprint = LayoutAnalyzer.make_text_fingerprint
normalize_page_number = LayoutAnalyzer.normalize_page_number
normalize_text = LayoutAnalyzer.normalize_text
parse_exclusion_review_markdown = LayoutAnalyzer.parse_exclusion_review_markdown
parse_table_geometry_visual_review_markdown = LayoutAnalyzer.parse_table_geometry_visual_review_markdown
text_block_records = LayoutAnalyzer.text_block_records

try:
    from pdf2docx.page.Pages import Pages
except Exception:
    Pages = None

try:
    import fitz
    from pdf2docx import Converter
    from pdf2docx.common import docx as docx_utils
    from docx import Document as DocxDocument
except Exception:
    fitz = None
    Converter = None
    docx_utils = None
    DocxDocument = None


class TestLayoutAnalyzer(unittest.TestCase):

    def test_normalize_text_collapses_whitespace(self):
        self.assertEqual(normalize_text('  Annual\n   Report\t2026  '), 'Annual Report 2026')
        self.assertEqual(normalize_text(None), '')

    def test_normalize_page_number_replaces_simple_page_number(self):
        for text in ['1', 'Page 2', 'p. 3', '4 / 10', 'Page 5 of 12', '- 6 -']:
            with self.subTest(text=text):
                self.assertEqual(normalize_page_number(text), PAGE_NUMBER_PLACEHOLDER)

    def test_normalize_page_number_keeps_non_page_number_text(self):
        self.assertEqual(normalize_page_number('Chapter 3'), 'Chapter 3')
        self.assertEqual(normalize_page_number('Section 1 Introduction'), 'Section 1 Introduction')

    def test_make_text_fingerprint_uses_normalized_page_number_and_band(self):
        fingerprint = make_text_fingerprint(' Page 7 ', y_band='Bottom')

        self.assertEqual(fingerprint, {
            'text': PAGE_NUMBER_PLACEHOLDER.lower(),
            'y_band': 'bottom',
            'style': '',
            'key': f'{PAGE_NUMBER_PLACEHOLDER.lower()}||bottom',
        })

    def test_classify_y_band_uses_page_height_ratios(self):
        self.assertEqual(classify_y_band([0, 10, 100, 30], page_height=1000), REGION_TOP)
        self.assertEqual(classify_y_band([0, 300, 100, 330], page_height=1000), REGION_BODY)
        self.assertEqual(classify_y_band([0, 950, 100, 980], page_height=1000), REGION_BOTTOM)

    def test_classify_y_band_rejects_invalid_ratios(self):
        with self.assertRaises(ValueError):
            classify_y_band([0, 10, 100, 30], page_height=1000, top_ratio=0.6, bottom_ratio=0.4)

    def test_text_block_records_are_json_serializable(self):
        records = text_block_records([
            {
                'page_index': 0,
                'height': 1000,
                'blocks': [
                    {'text': 'Annual Report', 'bbox': (50, 20, 300, 40), 'font': 'Arial', 'size': 10.0},
                    {'text': 'Body text', 'bbox': (50, 300, 500, 330)},
                ],
            },
        ])

        self.assertEqual(records[0]['region'], REGION_TOP)
        self.assertEqual(records[0]['bbox'], [50.0, 20.0, 300.0, 40.0])
        json.dumps(records)

    def test_find_repeated_text_candidates_clusters_top_and_bottom_text_only_by_default(self):
        pages = [
            {
                'page_index': 0,
                'height': 1000,
                'blocks': [
                    {'text': 'Annual Report', 'bbox': [50, 20, 300, 40]},
                    {'text': 'Repeated body line', 'bbox': [50, 300, 500, 330]},
                    {'text': 'Page 1', 'bbox': [260, 960, 310, 980]},
                ],
            },
            {
                'page_index': 1,
                'height': 1000,
                'blocks': [
                    {'text': 'Annual Report', 'bbox': [50, 22, 300, 42]},
                    {'text': 'Repeated body line', 'bbox': [50, 310, 500, 340]},
                    {'text': 'Page 2', 'bbox': [260, 960, 310, 980]},
                ],
            },
            {
                'page_index': 2,
                'height': 1000,
                'blocks': [
                    {'text': 'Annual Report', 'bbox': [50, 19, 300, 39]},
                    {'text': 'Repeated body line', 'bbox': [50, 320, 500, 350]},
                    {'text': 'Page 3', 'bbox': [260, 960, 310, 980]},
                ],
            },
        ]

        candidates = find_repeated_text_candidates(pages)
        by_text = {candidate['text']: candidate for candidate in candidates}

        self.assertIn('annual report', by_text)
        self.assertIn(PAGE_NUMBER_PLACEHOLDER.lower(), by_text)
        self.assertNotIn('repeated body line', by_text)
        self.assertEqual(by_text['annual report']['pages'], [0, 1, 2])
        self.assertEqual(by_text['annual report']['confidence'], 1.0)
        self.assertEqual(by_text['annual report']['signals']['support_pages'], 3)
        self.assertEqual(by_text[PAGE_NUMBER_PLACEHOLDER.lower()]['regions'], [REGION_BOTTOM])
        json.dumps(candidates)

    def test_build_layout_analysis_report_contains_page_summary_and_candidates(self):
        pages = [
            {
                'page_index': 0,
                'width': 600,
                'height': 1000,
                'blocks': [
                    {'text': 'Annual Report', 'bbox': [50, 20, 300, 40]},
                    {'text': 'First body paragraph', 'bbox': [50, 300, 500, 330]},
                    {'text': 'Page 1', 'bbox': [260, 960, 310, 980]},
                ],
            },
            {
                'page_index': 1,
                'width': 600,
                'height': 1000,
                'blocks': [
                    {'text': 'Annual Report', 'bbox': [50, 20, 300, 40]},
                    {'text': 'Second body paragraph', 'bbox': [50, 300, 500, 330]},
                    {'text': 'Page 2', 'bbox': [260, 960, 310, 980]},
                ],
            },
        ]

        report = build_layout_analysis_report(pages)
        by_text = {candidate['text']: candidate for candidate in report['repeated_text_candidates']}

        self.assertEqual(report['page_count'], 2)
        self.assertEqual(report['pages'][0]['region_counts'][REGION_TOP], 1)
        self.assertEqual(report['pages'][0]['region_counts'][REGION_BODY], 1)
        self.assertEqual(report['pages'][0]['region_counts'][REGION_BOTTOM], 1)
        self.assertIn('First body paragraph', report['pages'][0]['text'])
        self.assertIn('annual report', by_text)
        self.assertIn(PAGE_NUMBER_PLACEHOLDER.lower(), by_text)
        self.assertEqual(report['signals']['repeated_text_candidate_count'], 2)
        self.assertIn('paragraph_continuation_candidates', report)
        json.dumps(report)

    def test_paragraph_continuation_candidate_when_text_runs_across_pages(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('This paragraph starts on one page and continues', 50, 805, 520, 835),
            ]),
            _page(1, [
                _block('with the same sentence on the next page.', 50, 170, 520, 200),
            ]),
        ])
        candidate = report['paragraph_continuation_candidates'][0]

        self.assertEqual(candidate['label'], 'candidate')
        self.assertGreaterEqual(candidate['score'], 0.65)
        self.assertIn('previous_text_open_ended', candidate['positive_signals'])
        self.assertIn('previous_near_body_bottom', candidate['positive_signals'])
        self.assertIn('next_near_body_top', candidate['positive_signals'])

    def test_paragraph_continuation_unlikely_after_sentence_end(self):
        candidates = find_paragraph_continuation_candidates(build_layout_analysis_report([
            _page(0, [
                _block('This paragraph clearly ends here.', 50, 805, 520, 835),
            ]),
            _page(1, [
                _block('A new paragraph starts on the next page.', 50, 170, 520, 200),
            ]),
        ])['pages'])
        candidate = candidates[0]

        self.assertEqual(candidate['label'], 'unlikely')
        self.assertIn('previous_strong_sentence_end', candidate['negative_signals'])

    def test_paragraph_continuation_hyphenated_split_is_strong_candidate(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('The policy applies to inter-', 50, 805, 520, 835),
            ]),
            _page(1, [
                _block('national transactions recorded later', 50, 170, 520, 200),
            ]),
        ])
        candidate = report['paragraph_continuation_candidates'][0]

        self.assertEqual(candidate['label'], 'candidate')
        self.assertIn('previous_hyphenated_word', candidate['positive_signals'])
        self.assertIn('Hyphenated page break', candidate['reason'])

    def test_paragraph_continuation_unlikely_when_next_block_looks_like_heading(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('The preceding discussion continues without punctuation', 50, 805, 520, 835),
            ]),
            _page(1, [
                _block('CHAPTER TWO', 50, 170, 250, 200),
            ]),
        ])
        candidate = report['paragraph_continuation_candidates'][0]

        self.assertEqual(candidate['label'], 'unlikely')
        self.assertIn('next_looks_like_heading', candidate['negative_signals'])

    def test_paragraph_continuation_unlikely_for_footer_to_header_only(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Page 1', 270, 955, 330, 980),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 45),
            ]),
        ])
        candidate = report['paragraph_continuation_candidates'][0]

        self.assertEqual(candidate['label'], 'unlikely')
        self.assertEqual(candidate['previous_text_preview'], '')
        self.assertEqual(candidate['next_text_preview'], '')
        self.assertIn('no_previous_body_block', candidate['negative_signals'])
        self.assertIn('no_next_body_block', candidate['negative_signals'])

    def test_short_text_does_not_become_strong_continuation_evidence(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Note', 50, 805, 520, 835),
            ]),
            _page(1, [
                _block('x', 50, 170, 520, 200),
            ]),
        ])
        candidate = report['paragraph_continuation_candidates'][0]

        self.assertEqual(candidate['label'], 'unlikely')
        self.assertIn('previous_short_text', candidate['negative_signals'])
        self.assertIn('next_short_text', candidate['negative_signals'])
        self.assertIn('too short', candidate['reason'])

    def test_page_number_placeholder_is_not_normal_body_continuation_text(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Page 1', 50, 805, 520, 835),
            ]),
            _page(1, [
                _block('continued body text with matching layout', 50, 170, 520, 200),
            ]),
        ])
        candidate = report['paragraph_continuation_candidates'][0]

        self.assertEqual(candidate['label'], 'unlikely')
        self.assertIn('previous_placeholder_text', candidate['negative_signals'])
        self.assertIn('placeholder-like', candidate['reason'])

    def test_image_placeholder_is_reported_but_not_strong_semantic_repeated_text(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body paragraph one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body paragraph two.', 50, 300, 520, 330),
            ]),
            _page(2, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body paragraph three.', 50, 300, 520, 330),
            ]),
        ])
        by_text = {candidate['text']: candidate for candidate in report['repeated_text_candidates']}
        candidate = by_text[IMAGE_PLACEHOLDER.lower()]

        self.assertEqual(candidate['confidence'], 1.0)
        self.assertLess(candidate['semantic_confidence'], candidate['confidence'])
        self.assertEqual(candidate['confidence_label'], 'placeholder')
        self.assertEqual(candidate['signals']['text_quality']['placeholder_kind'], 'image')

    def test_two_page_repeated_cluster_is_more_cautious_than_all_page_header(self):
        pages = [
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Boundary Note', 50, 955, 300, 980),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Boundary Note', 50, 955, 300, 980),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
            _page(3, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
        ]
        report = build_layout_analysis_report(pages)
        by_text = {candidate['text']: candidate for candidate in report['repeated_text_candidates']}

        self.assertGreater(
            by_text['annual report']['semantic_confidence'],
            by_text['boundary note']['semantic_confidence'])
        self.assertEqual(by_text['annual report']['confidence_label'], 'strong')
        self.assertEqual(by_text['boundary note']['confidence_label'], 'cautious')
        self.assertEqual(by_text['boundary note']['signals']['support_level'], 'low')

    def test_continuation_avoids_likely_repeated_header_footer_blocks(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Running Footer', 50, 960, 300, 980),
                _block('The paragraph continues across pages', 50, 805, 520, 835),
                _block('Running Footer', 50, 830, 300, 850),
            ]),
            _page(1, [
                _block('Running Footer', 50, 20, 300, 40),
                _block('Running Footer', 50, 170, 300, 190),
                _block('on the next page with matching text', 50, 205, 520, 235),
                _block('Running Footer', 50, 960, 300, 980),
            ]),
            _page(2, [
                _block('Running Footer', 50, 20, 300, 40),
                _block('Another body paragraph.', 50, 300, 520, 330),
                _block('Running Footer', 50, 960, 300, 980),
            ]),
        ])
        candidate = report['paragraph_continuation_candidates'][0]

        self.assertEqual(
            candidate['previous_text_preview'],
            'The paragraph continues across pages')
        self.assertEqual(
            candidate['next_text_preview'],
            'on the next page with matching text')

    def test_dry_run_marks_strong_all_page_top_text_as_header_candidate(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body two.', 50, 300, 520, 330),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body three.', 50, 300, 520, 330),
            ]),
        ])
        dry_run = _dry_run_by_fingerprint(report)
        candidate = dry_run['annual report||top']

        self.assertEqual(candidate['action'], ACTION_WOULD_EXCLUDE)
        self.assertEqual(candidate['proposed_role'], ROLE_HEADER)
        self.assertIn('top_region', candidate['positive_signals'])

    def test_dry_run_marks_strong_all_page_bottom_text_as_footer_candidate(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Body one.', 50, 300, 520, 330),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
            _page(1, [
                _block('Body two.', 50, 300, 520, 330),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
            _page(2, [
                _block('Body three.', 50, 300, 520, 330),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
        ])
        dry_run = _dry_run_by_fingerprint(report)
        candidate = dry_run['confidential footer||bottom']

        self.assertEqual(candidate['action'], ACTION_WOULD_EXCLUDE)
        self.assertEqual(candidate['proposed_role'], ROLE_FOOTER)
        self.assertIn('bottom_region', candidate['positive_signals'])

    def test_dry_run_marks_page_number_placeholder_as_page_number_candidate(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Body one.', 50, 300, 520, 330),
                _block('Page 1', 270, 960, 330, 980),
            ]),
            _page(1, [
                _block('Body two.', 50, 300, 520, 330),
                _block('Page 2', 270, 960, 330, 980),
            ]),
            _page(2, [
                _block('Body three.', 50, 300, 520, 330),
                _block('Page 3', 270, 960, 330, 980),
            ]),
        ])
        dry_run = _dry_run_by_fingerprint(report)
        candidate = dry_run[f'{PAGE_NUMBER_PLACEHOLDER.lower()}||bottom']

        self.assertEqual(candidate['action'], ACTION_WOULD_EXCLUDE)
        self.assertEqual(candidate['proposed_role'], ROLE_PAGE_NUMBER)
        self.assertNotEqual(candidate['proposed_role'], ROLE_KEEP_BODY)
        self.assertIn('page_number_placeholder', candidate['positive_signals'])

    def test_dry_run_marks_image_placeholder_as_review_layout_signal(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body two.', 50, 300, 520, 330),
            ]),
            _page(2, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body three.', 50, 300, 520, 330),
            ]),
        ])
        dry_run = _dry_run_by_fingerprint(report)
        candidate = dry_run[f'{IMAGE_PLACEHOLDER.lower()}||top']

        self.assertEqual(candidate['action'], ACTION_REVIEW)
        self.assertEqual(candidate['proposed_role'], ROLE_LAYOUT_PLACEHOLDER)
        self.assertIn('placeholder_not_semantic_text', candidate['negative_signals'])

    def test_dry_run_keeps_two_page_cautious_cluster_out_of_automatic_exclusion(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Boundary Note', 50, 960, 300, 980),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Boundary Note', 50, 960, 300, 980),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
            _page(3, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
        ])
        dry_run = _dry_run_by_fingerprint(report)
        candidate = dry_run['boundary note||bottom']

        self.assertEqual(candidate['action'], ACTION_REVIEW)
        self.assertNotEqual(candidate['action'], ACTION_WOULD_EXCLUDE)
        self.assertIn('low_support', candidate['negative_signals'])

    def test_dry_run_keeps_body_region_repetition_as_body_content(self):
        pages = [
            _page(0, [
                _block('Repeated Body Line', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Repeated Body Line', 50, 310, 520, 340),
            ]),
            _page(2, [
                _block('Repeated Body Line', 50, 320, 520, 350),
            ]),
        ]
        repeated = find_repeated_text_candidates(pages, regions=(REGION_BODY,))
        dry_run = build_header_footer_exclusion_dry_run(repeated, page_count=3)
        candidate = dry_run['candidates'][0]

        self.assertEqual(candidate['action'], ACTION_KEEP)
        self.assertEqual(candidate['proposed_role'], ROLE_KEEP_BODY)
        self.assertIn('body_region_repetition', candidate['negative_signals'])

    def test_dry_run_does_not_mutate_input_blocks(self):
        pages = [
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body two.', 50, 300, 520, 330),
            ]),
        ]
        before = json.loads(json.dumps(pages))

        build_layout_analysis_report(pages)

        self.assertEqual(pages, before)

    def test_parse_exclusion_review_markdown_counts_manual_decisions(self):
        decisions = parse_exclusion_review_markdown('\n'.join([
            '### repeated-1 | header | would_exclude',
            '- fingerprint: `annual report||top`',
            '- human_decision: approve_exclude: [x]    reject_exclude: [ ]    unsure: [ ]',
            '### repeated-2 | review_only | review',
            '- fingerprint: `body note||bottom`',
            '- human_decision: approve_exclude: [ ]    reject_exclude: [X]    unsure: [ ]',
            '### repeated-3 | layout_placeholder | review',
            '- fingerprint: `<image>||top`',
            '- human_decision: approve_exclude: [ ]    reject_exclude: [ ]    unsure: [x]',
        ]))

        self.assertEqual(decisions['summary']['decision_counts'], {
            'approve_exclude': 1,
            'reject_exclude': 1,
            'unsure': 1,
        })
        self.assertEqual(decisions['decisions'][0]['fingerprint'], 'annual report||top')

    def test_reviewed_filter_removes_approved_candidate_only_when_opted_in(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph two.', 50, 300, 520, 330),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph three.', 50, 300, 520, 330),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))

        disabled = build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions)
        dry_run = build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)
        applied = build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True,
            apply=True)

        self.assertFalse(disabled['enabled'])
        self.assertEqual(disabled['summary']['removed_block_count'], 0)
        self.assertEqual(disabled['summary']['would_remove_block_count'], 0)
        self.assertEqual(dry_run['summary']['would_remove_block_count'], 3)
        self.assertEqual(dry_run['summary']['removed_block_count'], 0)
        self.assertEqual(applied['summary']['removed_block_count'], 3)
        self.assertEqual(applied['summary']['kept_block_count'], 3)
        self.assertNotIn('Annual Report', applied['filtered_pages'][0]['text'])

    def test_reviewed_filter_does_not_remove_rejected_or_unsure_candidates(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
        ])
        decisions = parse_exclusion_review_markdown('\n'.join([
            _review_markdown(
                'repeated-1',
                'annual report||top',
                'header',
                'would_exclude',
                'reject_exclude'),
            _review_markdown(
                'repeated-2',
                'confidential footer||bottom',
                'footer',
                'would_exclude',
                'unsure'),
        ]))

        applied = build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True,
            apply=True)

        self.assertEqual(applied['approved_candidate_count'], 0)
        self.assertEqual(applied['summary']['removed_block_count'], 0)
        self.assertEqual(applied['summary']['kept_block_count'], 6)

    def test_reviewed_filter_does_not_consume_raw_would_exclude_without_approval(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
        ])

        applied = build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            review_decisions=[],
            enabled=True,
            apply=True)

        self.assertEqual(applied['approved_candidate_count'], 0)
        self.assertEqual(applied['summary']['removed_block_count'], 0)
        self.assertEqual(applied['summary']['kept_block_count'], 3)

    def test_reviewed_filter_does_not_remove_layout_placeholder_even_if_approved(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body paragraph one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body paragraph two.', 50, 300, 520, 330),
            ]),
            _page(2, [
                _block(IMAGE_PLACEHOLDER, 50, 20, 120, 40),
                _block('Body paragraph three.', 50, 300, 520, 330),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            f'{IMAGE_PLACEHOLDER.lower()}||top',
            'layout_placeholder',
            'review',
            'approve_exclude'))

        applied = build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True,
            apply=True)

        self.assertEqual(applied['approved_candidate_count'], 0)
        self.assertEqual(applied['summary']['removed_block_count'], 0)
        self.assertIn(
            'dry_run_action_not_would_exclude',
            {item['reason'] for item in applied['blocked_candidates']})

    def test_reviewed_filter_default_does_not_mutate_page_summaries(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('More body text.', 50, 300, 520, 330),
            ]),
        ])
        before = json.loads(json.dumps(report['pages']))
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))

        build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions)

        self.assertEqual(report['pages'], before)

    def test_reviewed_filter_report_includes_removed_and_kept_counts(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Page 1', 270, 960, 330, 980),
                _block('Body paragraph one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Page 2', 270, 960, 330, 980),
                _block('Body paragraph two.', 50, 300, 520, 330),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            f'{PAGE_NUMBER_PLACEHOLDER.lower()}||bottom',
            'page_number',
            'would_exclude',
            'approve_exclude'))

        applied = build_reviewed_header_footer_filter_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True,
            apply=True)

        self.assertEqual(applied['summary']['original_block_count'], 4)
        self.assertEqual(applied['summary']['would_remove_block_count'], 2)
        self.assertEqual(applied['summary']['removed_block_count'], 2)
        self.assertEqual(applied['summary']['kept_block_count'], 2)

    def test_body_filtering_diff_report_includes_removed_and_kept_counts(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph two.', 50, 300, 520, 330),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph three.', 50, 300, 520, 330),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))

        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)

        self.assertEqual(diff['summary']['original_block_count'], 6)
        self.assertEqual(diff['summary']['would_remove_block_count'], 3)
        self.assertEqual(diff['summary']['kept_block_count'], 3)
        self.assertEqual(
            sum(page['removed_count'] for page in diff['removed_blocks_by_page']),
            3)

    def test_body_filtering_diff_groups_removed_blocks_by_approved_candidate(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
            ]),
        ])
        decisions = parse_exclusion_review_markdown('\n'.join([
            _review_markdown(
                'repeated-1',
                'annual report||top',
                'header',
                'would_exclude',
                'approve_exclude'),
            _review_markdown(
                'repeated-2',
                'confidential footer||bottom',
                'footer',
                'would_exclude',
                'approve_exclude'),
        ]))

        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)
        grouped = {
            item['fingerprint']: item
            for item in diff['removed_blocks_by_candidate']
        }

        self.assertEqual(grouped['annual report||top']['removed_count'], 3)
        self.assertEqual(grouped['confidential footer||bottom']['removed_count'], 3)

    def test_body_filtering_diff_keeps_rejected_unsure_and_placeholders(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
                _block(IMAGE_PLACEHOLDER, 50, 60, 120, 80),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
                _block(IMAGE_PLACEHOLDER, 50, 60, 120, 80),
            ]),
            _page(2, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Confidential Footer', 50, 960, 300, 980),
                _block(IMAGE_PLACEHOLDER, 50, 60, 120, 80),
            ]),
        ])
        decisions = parse_exclusion_review_markdown('\n'.join([
            _review_markdown(
                'repeated-1',
                'annual report||top',
                'header',
                'would_exclude',
                'reject_exclude'),
            _review_markdown(
                'repeated-2',
                'confidential footer||bottom',
                'footer',
                'would_exclude',
                'unsure'),
            _review_markdown(
                'repeated-3',
                f'{IMAGE_PLACEHOLDER.lower()}||top',
                'layout_placeholder',
                'review',
                'approve_exclude'),
        ]))

        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)

        self.assertEqual(diff['summary']['would_remove_block_count'], 0)
        self.assertEqual(diff['summary']['kept_block_count'], 9)
        self.assertEqual(diff['safety']['warnings'], [])

    def test_body_filtering_diff_report_does_not_mutate_input(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('More body text.', 50, 300, 520, 330),
            ]),
        ])
        before = json.loads(json.dumps(report['pages']))
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))

        build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)

        self.assertEqual(report['pages'], before)

    def test_body_filtering_diff_warns_when_unapproved_candidate_would_be_removed(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'reject_exclude'))
        unsafe_filter_report = {
            'approved_fingerprints': [],
            'approved_candidate_count': 0,
            'blocked_candidate_count': 1,
            'blocked_candidates': [{
                'candidate_id': 'repeated-1',
                'fingerprint': 'annual report||top',
                'proposed_role': 'header',
                'action': 'would_exclude',
                'manual_decision': 'reject_exclude',
            }],
            'pages': [{
                'page_index': 0,
                'removed_blocks': [{
                    'block_index': 0,
                    'fingerprint': 'annual report||top',
                    'candidate_id': 'repeated-1',
                    'proposed_role': 'header',
                    'manual_decision': 'reject_exclude',
                    'explicit_approval': False,
                    'removal_reason': 'unsafe test removal',
                }],
            }],
        }

        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True,
            filtering_report=unsafe_filter_report)

        self.assertIn('Unapproved candidates would be removed.', diff['safety']['warnings'])
        self.assertIn('Rejected candidates would be removed.', diff['safety']['warnings'])
        self.assertEqual(diff['safety']['unapproved_removed_candidate_count'], 1)
        self.assertEqual(diff['safety']['rejected_removed_candidate_count'], 1)

    def test_body_filtering_diff_disabled_mode_removes_zero_blocks(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph one.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph two.', 50, 300, 520, 330),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))

        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions)

        self.assertFalse(diff['enabled'])
        self.assertEqual(diff['summary']['would_remove_block_count'], 0)
        self.assertEqual(diff['summary']['kept_block_count'], 4)

    def test_paragraph_integrity_report_has_no_body_loss_for_top_bottom_removal(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph first line', 50, 300, 520, 330),
                _block('Body paragraph second line', 50, 335, 520, 365),
                _block('Page 1', 270, 960, 330, 980),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph continues', 50, 300, 520, 330),
                _block('Body paragraph finishes', 50, 335, 520, 365),
                _block('Page 2', 270, 960, 330, 980),
            ]),
        ])
        decisions = parse_exclusion_review_markdown('\n'.join([
            _review_markdown(
                'repeated-1',
                f'{PAGE_NUMBER_PLACEHOLDER.lower()}||bottom',
                'page_number',
                'would_exclude',
                'approve_exclude'),
            _review_markdown(
                'repeated-2',
                'annual report||top',
                'header',
                'would_exclude',
                'approve_exclude'),
        ]))
        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)

        integrity = build_paragraph_integrity_report(
            report['pages'],
            diff,
            enabled=True)

        self.assertEqual(integrity['summary']['removed_block_count'], 4)
        self.assertEqual(integrity['summary']['body_region_removed_count'], 0)
        self.assertEqual(integrity['summary']['top_bottom_removed_count'], 4)
        self.assertEqual(integrity['suspicious_body_loss_warnings'], [])
        self.assertEqual(integrity['suspicious_paragraph_gap_warnings'], [])
        self.assertTrue(integrity['summary']['line_level_body_blocks_available'])

    def test_paragraph_integrity_warns_if_body_region_block_is_removed(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Body paragraph first line', 50, 300, 520, 330),
                _block('Body paragraph removed line', 50, 335, 520, 365),
                _block('Body paragraph final line', 50, 370, 520, 400),
            ]),
        ])
        diff = _unsafe_diff_report_for_removed_blocks(
            report['pages'],
            [(0, 1)])

        integrity = build_paragraph_integrity_report(
            report['pages'],
            diff,
            enabled=True)

        self.assertEqual(integrity['summary']['body_region_removed_count'], 1)
        self.assertIn(
            'body_region_removed',
            {warning['type'] for warning in integrity['suspicious_body_loss_warnings']})
        self.assertIn(
            'body_flow_gap',
            {warning['type'] for warning in integrity['suspicious_paragraph_gap_warnings']})

    def test_paragraph_integrity_warns_on_high_page_body_loss(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Body paragraph first line', 50, 300, 520, 330),
                _block('Body paragraph second line', 50, 335, 520, 365),
                _block('Body paragraph third line', 50, 370, 520, 400),
            ]),
        ])
        diff = _unsafe_diff_report_for_removed_blocks(
            report['pages'],
            [(0, 0), (0, 1)])

        integrity = build_paragraph_integrity_report(
            report['pages'],
            diff,
            enabled=True,
            high_body_loss_ratio=0.5)

        self.assertIn(
            'high_body_loss_ratio',
            {warning['type'] for warning in integrity['suspicious_body_loss_warnings']})

    def test_paragraph_integrity_keeps_line_level_body_blocks_available(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body line one', 50, 300, 520, 330),
                _block('Body line two', 50, 335, 520, 365),
                _block('Body line three', 50, 370, 520, 400),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body line four', 50, 300, 520, 330),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))
        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)

        integrity = build_paragraph_integrity_report(
            report['pages'],
            diff,
            enabled=True)
        kept_body_blocks = [
            block
            for page in integrity['filtered_pages']
            for block in page['text_blocks']
            if block['region'] == REGION_BODY
        ]

        self.assertEqual(len(kept_body_blocks), 4)
        self.assertTrue(integrity['summary']['line_level_body_blocks_available'])

    def test_paragraph_integrity_report_does_not_mutate_input(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('More body text.', 50, 300, 520, 330),
            ]),
        ])
        before = json.loads(json.dumps(report['pages']))
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))
        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)

        build_paragraph_integrity_report(
            report['pages'],
            diff,
            enabled=True)

        self.assertEqual(report['pages'], before)

    def test_paragraph_integrity_disabled_mode_keeps_original_summaries(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body paragraph.', 50, 300, 520, 330),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('More body text.', 50, 300, 520, 330),
            ]),
        ])
        decisions = parse_exclusion_review_markdown(_review_markdown(
            'repeated-1',
            'annual report||top',
            'header',
            'would_exclude',
            'approve_exclude'))
        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)

        integrity = build_paragraph_integrity_report(
            report['pages'],
            diff)

        self.assertFalse(integrity['enabled'])
        self.assertEqual(integrity['summary']['removed_block_count'], 0)
        self.assertEqual(integrity['summary']['filtered_block_count'], 4)
        self.assertEqual(integrity['filtered_pages'], report['pages'])

    def test_paragraph_reconstruction_groups_consistent_body_lines(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('This visual line starts a paragraph', 50, 300, 520, 320),
                _block('and this visual line continues it', 50, 326, 520, 346),
                _block('before the same paragraph finishes.', 50, 352, 520, 372),
                _block('Page 1', 270, 960, 330, 980),
            ]),
            _page(1, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Another body paragraph.', 50, 300, 520, 320),
                _block('Page 2', 270, 960, 330, 980),
            ]),
        ])
        decisions = parse_exclusion_review_markdown('\n'.join([
            _review_markdown(
                'repeated-1',
                'annual report||top',
                'header',
                'would_exclude',
                'approve_exclude'),
            _review_markdown(
                'repeated-2',
                f'{PAGE_NUMBER_PLACEHOLDER.lower()}||bottom',
                'page_number',
                'would_exclude',
                'approve_exclude'),
        ]))
        diff = build_body_filtering_diff_report(
            report['pages'],
            report['header_footer_exclusion_dry_run'],
            decisions,
            enabled=True)
        integrity = build_paragraph_integrity_report(
            report['pages'],
            diff,
            enabled=True)

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            paragraph_integrity_report=integrity,
            enabled=True)

        self.assertEqual(
            validation['summary']['body_block_count_before_filtering'], 4)
        self.assertEqual(
            validation['summary']['body_block_count_after_filtering'], 4)
        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 1)
        self.assertEqual(
            validation['pages'][0]['estimated_paragraph_groups'][0]['block_count'], 3)
        self.assertEqual(validation['summary']['warning_count'], 0)

    def test_paragraph_reconstruction_groups_same_row_fragments_as_one_line(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Click on', 50, 300, 110, 320),
                _block('Header', 116, 300, 170, 320, style={'font': 'Arial', 'size': 11.0}),
                _block('and then choose Edit.', 50, 326, 520, 346),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        group = validation['pages'][0]['estimated_paragraph_groups'][0]
        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 1)
        self.assertEqual(group['line_count'], 2)
        self.assertEqual(group['block_count'], 3)

    def test_paragraph_reconstruction_detects_hard_break_signals(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('First paragraph line one', 50, 300, 520, 320),
                _block('First paragraph line two.', 50, 326, 250, 346),
                _block('Indented new paragraph starts', 80, 352, 520, 372),
                _block(
                    'Different style paragraph starts',
                    80, 378, 520, 398,
                    style={'font': 'Arial', 'size': 12.0}),
                _block(
                    'Large gap paragraph starts',
                    80, 520, 520, 540,
                    style={'font': 'Arial', 'size': 12.0}),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        groups = validation['pages'][0]['estimated_paragraph_groups']
        break_reasons = [
            reason
            for group in groups
            for reason in group['break_before_reasons']
        ]
        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 4)
        self.assertIn('indentation_change', break_reasons)
        self.assertIn('style_change', break_reasons)
        self.assertIn('large_vertical_gap', break_reasons)
        self.assertIn('split_boundaries', validation['pages'][0])
        boundary_reasons = [
            reason
            for boundary in validation['pages'][0]['split_boundaries']
            for reason in boundary['reasons']
        ]
        self.assertIn('indentation_change', boundary_reasons)
        self.assertIn('sentence_end_with_trailing_space', boundary_reasons)

    def test_paragraph_reconstruction_ignores_weak_indentation_change(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('This visual line continues without punctuation', 50, 300, 520, 320),
                _block('and this indented line is still continuation text', 80, 326, 520, 346),
                _block('with one more aligned continuation line.', 80, 352, 520, 372),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        page = validation['pages'][0]
        self.assertEqual(page['estimated_paragraph_group_count'], 1)
        self.assertEqual(
            page['ignored_split_boundaries'][0]['ignored_reasons'],
            ['weak_indentation_change'])
        self.assertEqual(
            validation['summary']['ignored_split_reason_counts'],
            {'weak_indentation_change': 1})

    def test_paragraph_reconstruction_indentation_with_free_space_can_split(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('A short complete sentence.', 50, 300, 250, 320),
                _block('Indented new paragraph starts after free space', 90, 326, 520, 346),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        split_reasons = validation['pages'][0]['split_boundaries'][0]['reasons']
        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 2)
        self.assertIn('indentation_change', split_reasons)
        self.assertIn('sentence_end_with_trailing_space', split_reasons)

    def test_paragraph_reconstruction_heading_like_indentation_still_splits(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Overview', 50, 300, 170, 320),
                _block('Indented body starts below the heading.', 85, 326, 520, 346),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        split_reasons = validation['pages'][0]['split_boundaries'][0]['reasons']
        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 2)
        self.assertIn('previous_heading_like', split_reasons)
        self.assertIn('indentation_change', split_reasons)

    def test_paragraph_reconstruction_list_indentation_still_splits(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Introductory prose before a list', 50, 300, 520, 320),
                _block('- Indented list item', 80, 326, 420, 346),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        split_reasons = validation['pages'][0]['split_boundaries'][0]['reasons']
        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 2)
        self.assertIn('list_marker', split_reasons)
        self.assertIn('indentation_change', split_reasons)

    def test_paragraph_reconstruction_weak_indent_relaxation_reduces_group_count(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('A continuing instruction line without a period', 50, 300, 520, 320),
                _block('wrapped continuation line with indentation', 80, 326, 520, 346),
                _block('another wrapped continuation line', 80, 352, 520, 372),
                _block('final wrapped continuation line.', 80, 378, 520, 398),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 1)
        self.assertEqual(
            validation['diagnostics']['ignored_split_reason_counts'],
            {'weak_indentation_change': 1})

    def test_paragraph_reconstruction_sentence_end_does_not_force_every_line_split(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('This complete sentence fills the line.', 50, 300, 550, 320),
                _block('The next visual line can still belong with it', 50, 326, 550, 346),
                _block('Short sentence.', 50, 352, 250, 372),
                _block('A new paragraph starts after visible free space', 50, 378, 550, 398),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 2)
        split_reasons = [
            reason
            for boundary in validation['pages'][0]['split_boundaries']
            for reason in boundary['reasons']
        ]
        self.assertIn('sentence_end_with_trailing_space', split_reasons)

    def test_paragraph_reconstruction_splits_heading_like_short_lines(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Overview', 50, 300, 180, 320),
                _block('This body paragraph starts below the heading', 50, 326, 550, 346),
                _block('and continues on a second visual line.', 50, 352, 550, 372),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 2)
        split_reasons = [
            reason
            for boundary in validation['pages'][0]['split_boundaries']
            for reason in boundary['reasons']
        ]
        self.assertIn('previous_heading_like', split_reasons)

    def test_paragraph_reconstruction_keeps_list_items_separate_from_prose(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Introductory prose before a list.', 50, 300, 550, 320),
                _block('- First list item', 70, 326, 400, 346),
                _block('Closing prose after the list item', 50, 352, 550, 372),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 3)
        split_reasons = [
            reason
            for boundary in validation['pages'][0]['split_boundaries']
            for reason in boundary['reasons']
        ]
        self.assertIn('list_marker', split_reasons)
        self.assertIn('previous_list_item', split_reasons)

    def test_paragraph_reconstruction_hyphenated_line_ending_is_continuation_signal(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('This line ends with a hyphen-', 50, 300, 350, 320),
                _block('ated word despite indentation change.', 80, 326, 550, 346),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(validation['pages'][0]['estimated_paragraph_group_count'], 1)
        self.assertEqual(validation['pages'][0]['split_boundaries'], [])

    def test_paragraph_reconstruction_warns_on_excessive_one_line_fragments(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Small bit one', 50, 300, 520, 320),
                _block('Small bit two', 50, 380, 520, 400),
                _block('Small bit three', 50, 460, 520, 480),
                _block('Small bit four', 50, 540, 520, 560),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(
            validation['summary']['suspicious_single_line_paragraph_count'], 4)
        self.assertEqual(
            validation['summary']['suspicious_short_fragment_count'], 4)
        self.assertIn(
            'excessive_one_line_fragmentation',
            {warning['type'] for warning in validation['warnings']})
        self.assertGreater(
            validation['diagnostics']['one_line_group_ratio'],
            0.0)
        self.assertEqual(
            validation['diagnostics']['pages_with_worst_fragmentation'][0]['page_number'],
            1)
        self.assertIn('1', validation['diagnostics']['groups_by_line_count'])

    def test_paragraph_reconstruction_reports_cross_page_continuation_without_merge(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('This paragraph starts on one page and continues', 50, 805, 520, 835),
            ]),
            _page(1, [
                _block('with the same sentence on the next page.', 50, 170, 520, 200),
            ]),
        ])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(
            validation['summary']['possible_cross_page_continuation_count'], 1)
        self.assertEqual(
            validation['summary']['estimated_paragraph_group_count'], 2)
        self.assertIn(
            'possible_cross_page_continuation',
            {warning['type'] for warning in validation['warnings']})

    def test_paragraph_reconstruction_report_does_not_mutate_input(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Body line one', 50, 300, 520, 320),
                _block('Body line two.', 50, 326, 520, 346),
            ]),
        ])
        before = json.loads(json.dumps(report['pages']))

        build_paragraph_reconstruction_validation_report(
            report['pages'],
            enabled=True)

        self.assertEqual(report['pages'], before)

    def test_paragraph_reconstruction_disabled_mode_keeps_original_summaries(self):
        report = build_layout_analysis_report([
            _page(0, [
                _block('Annual Report', 50, 20, 300, 40),
                _block('Body line one', 50, 300, 520, 320),
                _block('Body line two.', 50, 326, 520, 346),
            ]),
        ])
        diff = _unsafe_diff_report_for_removed_blocks(
            report['pages'],
            [(0, 1)])

        validation = build_paragraph_reconstruction_validation_report(
            report['pages'],
            body_filtering_diff_report=diff)

        self.assertFalse(validation['enabled'])
        self.assertEqual(
            validation['summary']['body_block_count_before_filtering'], 2)
        self.assertEqual(
            validation['summary']['body_block_count_after_filtering'], 2)
        self.assertEqual(
            validation['summary']['estimated_paragraph_group_count'], 0)
        self.assertEqual(validation['filtered_pages'], report['pages'])

    def test_paragraph_production_comparison_includes_estimator_metrics(self):
        estimator = build_paragraph_reconstruction_validation_report(
            build_layout_analysis_report([
                _page(0, [
                    _block('Body line one', 50, 300, 520, 320),
                    _block('Body line two.', 50, 326, 520, 346),
                ]),
            ])['pages'],
            enabled=True)

        comparison = build_paragraph_production_comparison_report(
            estimator,
            production_pages=[],
            enabled=True)

        self.assertTrue(comparison['estimator']['available'])
        self.assertEqual(comparison['estimator']['paragraph_group_count'], 1)
        self.assertFalse(comparison['production_observed']['available'])
        self.assertEqual(
            comparison['warnings'][0]['type'],
            'production_metrics_unavailable')

    def test_paragraph_production_comparison_includes_production_metrics(self):
        estimator = build_paragraph_reconstruction_validation_report(
            build_layout_analysis_report([
                _page(0, [
                    _block('Body line one', 50, 300, 520, 320),
                    _block('Body line two.', 50, 326, 520, 346),
                    _block('Second paragraph.', 50, 420, 300, 440),
                ]),
            ])['pages'],
            enabled=True)
        production_pages = [
            _production_page(0, [
                _production_text_block(
                    ['Body line one', 'Body line two.'],
                    [50, 300, 520, 346]),
                _production_text_block(
                    ['Second paragraph.'],
                    [50, 420, 300, 440]),
            ]),
        ]

        comparison = build_paragraph_production_comparison_report(
            estimator,
            production_pages=production_pages,
            enabled=True)

        self.assertTrue(comparison['production_observed']['available'])
        self.assertEqual(comparison['production_observed']['paragraph_group_count'], 2)
        self.assertEqual(comparison['production_observed']['total_body_line_count'], 3)
        self.assertIn('average_lines_per_group', comparison['production_observed'])

    def test_paragraph_production_comparison_computes_mismatch_ratio(self):
        estimator = build_paragraph_reconstruction_validation_report(
            build_layout_analysis_report([
                _page(0, [
                    _block('One.', 50, 300, 200, 320),
                    _block('Two.', 50, 380, 200, 400),
                    _block('Three.', 50, 460, 200, 480),
                ]),
            ])['pages'],
            enabled=True)
        production_pages = [
            _production_page(0, [
                _production_text_block(
                    ['One.', 'Two.', 'Three.'],
                    [50, 300, 200, 480]),
            ]),
        ]

        comparison = build_paragraph_production_comparison_report(
            estimator,
            production_pages=production_pages,
            enabled=True)

        self.assertEqual(comparison['mismatch']['estimator_group_count'], 3)
        self.assertEqual(comparison['mismatch']['production_group_count'], 1)
        self.assertEqual(comparison['mismatch']['group_count_delta_ratio'], 2.0)
        self.assertEqual(
            comparison['warnings'][0]['type'],
            'high_group_count_mismatch')

    def test_paragraph_production_comparison_does_not_mutate_inputs(self):
        estimator = build_paragraph_reconstruction_validation_report(
            build_layout_analysis_report([
                _page(0, [
                    _block('Body line one', 50, 300, 520, 320),
                    _block('Body line two.', 50, 326, 520, 346),
                ]),
            ])['pages'],
            enabled=True)
        production_pages = [
            _production_page(0, [
                _production_text_block(
                    ['Body line one', 'Body line two.'],
                    [50, 300, 520, 346]),
            ]),
        ]
        before_estimator = json.loads(json.dumps(estimator))
        before_production = json.loads(json.dumps(production_pages))

        build_paragraph_production_comparison_report(
            estimator,
            production_pages=production_pages,
            enabled=True)

        self.assertEqual(estimator, before_estimator)
        self.assertEqual(production_pages, before_production)

    def test_paragraph_production_comparison_disabled_mode_is_clear(self):
        estimator = build_paragraph_reconstruction_validation_report(
            build_layout_analysis_report([
                _page(0, [
                    _block('Body line one', 50, 300, 520, 320),
                    _block('Body line two.', 50, 326, 520, 346),
                ]),
            ])['pages'],
            enabled=True)
        production_pages = [
            _production_page(0, [
                _production_text_block(
                    ['Body line one', 'Body line two.'],
                    [50, 300, 520, 346]),
            ]),
        ]

        comparison = build_paragraph_production_comparison_report(
            estimator,
            production_pages=production_pages)

        self.assertFalse(comparison['enabled'])
        self.assertFalse(comparison['production_observed']['available'])
        self.assertFalse(comparison['mismatch']['available'])

    def test_paragraph_mismatch_analysis_identifies_estimator_over_splitting(self):
        estimator = _estimator_report_for_mismatch([
            _estimator_page_for_mismatch(
                0,
                group_count=5,
                body_blocks=10,
                split_reasons=['indentation_change'] * 4),
        ])
        production_pages = [
            _production_page(0, [
                _production_text_block(['Line one', 'Line two', 'Line three'], [50, 300, 520, 360]),
            ]),
        ]

        analysis = build_paragraph_mismatch_analysis_report(
            estimator,
            production_pages=production_pages,
            enabled=True)

        self.assertEqual(
            analysis['summary']['dominant_mismatch_cause'],
            'estimator_over_split_by_indentation')
        self.assertEqual(
            analysis['pages'][0]['likely_cause'],
            'estimator_over_split_by_indentation')
        self.assertIn('indentation_change', analysis['pages'][0]['estimator_split_reason_counts'])

    def test_paragraph_mismatch_analysis_identifies_possible_production_over_merge(self):
        estimator = _estimator_report_for_mismatch([
            _estimator_page_for_mismatch(
                0,
                group_count=4,
                body_blocks=8,
                split_reasons=[]),
        ])
        production_pages = [
            _production_page(0, [
                _production_text_block(
                    ['Line one', 'Line two', 'Line three', 'Line four', 'Line five', 'Line six'],
                    [50, 300, 520, 460]),
            ]),
        ]

        analysis = build_paragraph_mismatch_analysis_report(
            estimator,
            production_pages=production_pages,
            enabled=True)

        self.assertEqual(
            analysis['pages'][0]['likely_cause'],
            'production_possible_over_merge')
        self.assertTrue(analysis['summary']['mostly_production_over_merging'])

    def test_paragraph_mismatch_analysis_lists_worst_pages(self):
        estimator = _estimator_report_for_mismatch([
            _estimator_page_for_mismatch(0, group_count=2, body_blocks=4, split_reasons=[]),
            _estimator_page_for_mismatch(1, group_count=8, body_blocks=16, split_reasons=['style_change'] * 7),
        ])
        production_pages = [
            _production_page(0, [
                _production_text_block(['One'], [50, 300, 520, 320]),
            ]),
            _production_page(1, [
                _production_text_block(['One', 'Two'], [50, 300, 520, 350]),
            ]),
        ]

        analysis = build_paragraph_mismatch_analysis_report(
            estimator,
            production_pages=production_pages,
            enabled=True)

        self.assertEqual(analysis['pages'][0]['page_number'], 2)
        self.assertEqual(analysis['pages'][0]['absolute_group_count_delta'], 7)
        self.assertEqual(
            analysis['pages'][0]['likely_cause'],
            'estimator_over_split_by_style_change')

    def test_paragraph_mismatch_analysis_handles_missing_production_metrics(self):
        estimator = _estimator_report_for_mismatch([
            _estimator_page_for_mismatch(0, group_count=2, body_blocks=4, split_reasons=[]),
        ])

        analysis = build_paragraph_mismatch_analysis_report(
            estimator,
            enabled=True)

        self.assertFalse(analysis['summary']['available'])
        self.assertEqual(
            analysis['warnings'][0]['type'],
            'production_metrics_unavailable')

    def test_paragraph_mismatch_analysis_does_not_mutate_inputs(self):
        estimator = _estimator_report_for_mismatch([
            _estimator_page_for_mismatch(
                0,
                group_count=3,
                body_blocks=6,
                split_reasons=['sentence_end_with_trailing_space'] * 2),
        ])
        production_pages = [
            _production_page(0, [
                _production_text_block(['One', 'Two'], [50, 300, 520, 350]),
            ]),
        ]
        before_estimator = json.loads(json.dumps(estimator))
        before_production = json.loads(json.dumps(production_pages))

        build_paragraph_mismatch_analysis_report(
            estimator,
            production_pages=production_pages,
            enabled=True)

        self.assertEqual(estimator, before_estimator)
        self.assertEqual(production_pages, before_production)

    def test_paragraph_mismatch_analysis_disabled_mode_is_clear(self):
        estimator = _estimator_report_for_mismatch([
            _estimator_page_for_mismatch(0, group_count=2, body_blocks=4, split_reasons=[]),
        ])
        production_pages = [
            _production_page(0, [
                _production_text_block(['One'], [50, 300, 520, 320]),
            ]),
        ]

        analysis = build_paragraph_mismatch_analysis_report(
            estimator,
            production_pages=production_pages)

        self.assertFalse(analysis['enabled'])
        self.assertFalse(analysis['summary']['available'])
        self.assertEqual(analysis['pages'], [])

    def test_indentation_rule_comparison_marks_small_indent_as_mergeable(self):
        report = _estimator_report_for_indentation([
            _indentation_boundary(
                page_index=0,
                left_delta=12.0,
                previous_sentence_end=False,
                width_similar=True,
                previous_width_ratio=0.95,
                previous_right_gap_ratio=0.02),
        ])

        comparison = build_indentation_rule_comparison_report(
            report,
            enabled=True)

        self.assertEqual(
            comparison['summary']['total_indentation_split_boundaries'], 1)
        self.assertEqual(
            comparison['boundaries'][0]['recommendation'],
            'estimator_should_merge')
        self.assertEqual(
            comparison['boundaries'][0]['production_like_expected_behavior'],
            'keep_together')

    def test_indentation_rule_comparison_marks_clear_new_paragraph_as_split(self):
        report = _estimator_report_for_indentation([
            _indentation_boundary(
                page_index=0,
                left_delta=40.0,
                previous_sentence_end=True,
                width_similar=False,
                previous_width_ratio=0.1,
                previous_right_gap_ratio=0.9),
        ])

        comparison = build_indentation_rule_comparison_report(
            report,
            enabled=True)

        self.assertEqual(
            comparison['boundaries'][0]['recommendation'],
            'estimator_should_split')
        self.assertEqual(
            comparison['summary']['estimator_should_split_count'], 1)

    def test_indentation_rule_comparison_keeps_heading_list_boundary_split(self):
        report = _estimator_report_for_indentation([
            _indentation_boundary(
                page_index=0,
                left_delta=35.0,
                previous_sentence_end=False,
                width_similar=True,
                previous_width_ratio=0.95,
                previous_right_gap_ratio=0.01,
                extra_reasons=['list_marker'],
                current_list_marker=True),
        ])

        comparison = build_indentation_rule_comparison_report(
            report,
            enabled=True)

        self.assertEqual(
            comparison['boundaries'][0]['production_like_expected_behavior'],
            'treat_as_heading_list_table_boundary')
        self.assertEqual(
            comparison['boundaries'][0]['recommendation'],
            'estimator_should_split')

    def test_indentation_rule_comparison_reports_missing_metadata(self):
        report = _estimator_report_for_indentation([
            _indentation_boundary(
                page_index=0,
                left_delta=0.0,
                previous_sentence_end=False,
                width_similar=False,
                previous_width_ratio=0.0,
                previous_right_gap_ratio=0.0,
                insufficient_metadata=True),
        ])

        comparison = build_indentation_rule_comparison_report(
            report,
            enabled=True)

        self.assertEqual(
            comparison['boundaries'][0]['recommendation'],
            'needs_more_metadata')
        self.assertEqual(
            comparison['summary']['needs_more_metadata_count'], 1)

    def test_indentation_rule_comparison_produces_summary_counts(self):
        report = _estimator_report_for_indentation([
            _indentation_boundary(0, left_delta=12.0, previous_sentence_end=False, width_similar=True),
            _indentation_boundary(0, left_delta=40.0, previous_sentence_end=True, previous_width_ratio=0.1, previous_right_gap_ratio=0.9),
            _indentation_boundary(1, left_delta=35.0, previous_sentence_end=False, extra_reasons=['heading_like'], current_heading_like=True),
        ])

        comparison = build_indentation_rule_comparison_report(
            report,
            enabled=True)

        self.assertEqual(
            comparison['summary']['total_indentation_split_boundaries'], 3)
        self.assertEqual(
            comparison['summary']['estimator_should_merge_count'], 1)
        self.assertEqual(
            comparison['summary']['estimator_should_split_count'], 2)
        self.assertEqual(comparison['pages'][0]['page_number'], 1)

    def test_indentation_rule_comparison_does_not_mutate_inputs(self):
        report = _estimator_report_for_indentation([
            _indentation_boundary(0, left_delta=12.0, previous_sentence_end=False, width_similar=True),
        ])
        before = json.loads(json.dumps(report))

        build_indentation_rule_comparison_report(
            report,
            enabled=True)

        self.assertEqual(report, before)

    def test_indentation_rule_comparison_disabled_mode_is_clear(self):
        report = _estimator_report_for_indentation([
            _indentation_boundary(0, left_delta=12.0, previous_sentence_end=False, width_similar=True),
        ])

        comparison = build_indentation_rule_comparison_report(report)

        self.assertFalse(comparison['enabled'])
        self.assertEqual(
            comparison['summary']['total_indentation_split_boundaries'], 0)
        self.assertEqual(comparison['boundaries'], [])

    def test_filter_insertion_analysis_includes_all_candidate_stages(self):
        analysis = build_filter_insertion_point_analysis_report(
            **_filter_insertion_reports(),
            enabled=True)

        candidate_ids = {
            point['candidate_id']
            for point in analysis['insertion_points']
        }
        self.assertEqual(analysis['summary']['evaluated_insertion_point_count'], 6)
        self.assertEqual(candidate_ids, {
            'raw_page_cleanup',
            'document_parse',
            'before_page_parse',
            'before_blocks_cleanup_or_grouping',
            'after_textblock_grouping',
            'docx_generation',
        })

    def test_filter_insertion_analysis_stage_fields_are_complete(self):
        analysis = build_filter_insertion_point_analysis_report(
            **_filter_insertion_reports(),
            enabled=True)

        for point in analysis['insertion_points']:
            self.assertIn(point['risk_level'], {'low', 'medium', 'high'})
            self.assertIn(point['implementation_complexity'], {'low', 'medium', 'high'})
            self.assertIn(point['recommendation'], {'preferred', 'possible', 'avoid'})
            self.assertIn('paragraph_grouping_impact', point)
            self.assertIn('table_detection_impact', point)
            self.assertIn('image_shape_extraction_impact', point)

    def test_filter_insertion_analysis_prefers_document_level_stage(self):
        analysis = build_filter_insertion_point_analysis_report(
            **_filter_insertion_reports(),
            enabled=True)

        by_id = _insertion_points_by_id(analysis)
        self.assertEqual(
            analysis['summary']['preferred_insertion_point'],
            'document_parse')
        self.assertEqual(by_id['document_parse']['recommendation'], 'preferred')
        self.assertEqual(by_id['document_parse']['risk_level'], 'low')

    def test_filter_insertion_analysis_marks_docx_stage_incomplete(self):
        analysis = build_filter_insertion_point_analysis_report(
            **_filter_insertion_reports(),
            enabled=True)

        docx_point = _insertion_points_by_id(analysis)['docx_generation']
        self.assertEqual(docx_point['recommendation'], 'avoid')
        self.assertEqual(docx_point['risk_level'], 'high')
        self.assertIn('Incomplete', docx_point['paragraph_grouping_impact'])

    def test_filter_insertion_analysis_marks_post_textblock_filtering_risky(self):
        analysis = build_filter_insertion_point_analysis_report(
            **_filter_insertion_reports(),
            enabled=True)

        point = _insertion_points_by_id(analysis)['after_textblock_grouping']
        self.assertEqual(point['recommendation'], 'avoid')
        self.assertEqual(point['risk_level'], 'high')
        self.assertTrue(any('merged' in signal for signal in point['negative_signals']))

    def test_filter_insertion_analysis_handles_missing_reports(self):
        analysis = build_filter_insertion_point_analysis_report(enabled=True)

        self.assertTrue(analysis['enabled'])
        self.assertIn(
            'layout_analysis_report',
            analysis['summary']['missing_inputs'])
        self.assertIn(
            'production_grouping_metrics_unavailable',
            {warning['type'] for warning in analysis['warnings']})

    def test_filter_insertion_analysis_does_not_mutate_inputs(self):
        reports = _filter_insertion_reports()
        before = json.loads(json.dumps(reports))

        build_filter_insertion_point_analysis_report(
            **reports,
            enabled=True)

        self.assertEqual(reports, before)

    def test_filter_insertion_analysis_disabled_mode_is_clear(self):
        analysis = build_filter_insertion_point_analysis_report(
            **_filter_insertion_reports())

        self.assertFalse(analysis['enabled'])
        self.assertEqual(analysis['insertion_points'], [])
        self.assertEqual(
            analysis['summary']['evaluated_insertion_point_count'],
            0)

    def test_document_parse_simulation_removes_only_approved_candidates(self):
        inputs = _document_parse_simulation_inputs()

        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=True,
            expected_removed_count=2,
            expected_kept_count=4)

        self.assertEqual(simulation['insertion_point'], 'document_parse')
        self.assertEqual(simulation['summary']['would_remove_block_count'], 2)
        self.assertEqual(simulation['summary']['simulated_removed_count'], 2)
        self.assertEqual(simulation['removed_counts_by_role'], {
            ROLE_FOOTER: 1,
            ROLE_HEADER: 1,
        })

    def test_document_parse_simulation_keeps_rejected_unsure_and_placeholders(self):
        inputs = _document_parse_simulation_inputs()

        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=True,
            expected_removed_count=2,
            expected_kept_count=4)
        kept_fingerprints = {
            block['fingerprint']
            for page in simulation['simulated_apply']['filtered_pages']
            for block in page['text_blocks']
        }

        self.assertIn('reject-header', kept_fingerprints)
        self.assertIn('unsure-footer', kept_fingerprints)
        self.assertIn('image-placeholder', kept_fingerprints)
        self.assertEqual(simulation['summary']['rejected_removed_count'], 0)
        self.assertEqual(simulation['summary']['unsure_removed_count'], 0)
        self.assertEqual(simulation['summary']['layout_placeholder_removed_count'], 0)

    def test_document_parse_simulation_preserves_body_region_blocks(self):
        inputs = _document_parse_simulation_inputs()

        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=True,
            expected_removed_count=2,
            expected_kept_count=4)

        self.assertEqual(simulation['summary']['body_region_removed_count'], 0)
        self.assertTrue(
            simulation['downstream_availability']['body_region_blocks_preserved'])
        self.assertEqual(
            simulation['downstream_availability']['paragraph_grouping_body_block_count'],
            1)

    def test_document_parse_simulation_dry_run_removes_zero_blocks(self):
        inputs = _document_parse_simulation_inputs()

        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=False,
            expected_removed_count=2,
            expected_kept_count=4)

        self.assertEqual(simulation['dry_run']['would_remove_block_count'], 2)
        self.assertEqual(simulation['dry_run']['removed_block_count'], 0)
        self.assertFalse(simulation['simulated_apply']['applied'])
        self.assertEqual(simulation['simulated_apply']['removed_block_count'], 0)
        self.assertEqual(simulation['summary']['simulated_kept_count'], 6)

    def test_document_parse_simulation_apply_uses_copied_data(self):
        inputs = _document_parse_simulation_inputs()
        before_pages = json.loads(json.dumps(inputs['page_summaries']))

        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=True,
            expected_removed_count=2,
            expected_kept_count=4)

        self.assertEqual(inputs['page_summaries'], before_pages)
        self.assertIsNot(
            simulation['simulated_apply']['filtered_pages'][0],
            inputs['page_summaries'][0])
        self.assertEqual(
            len(simulation['simulated_apply']['filtered_pages'][0]['text_blocks']),
            4)

    def test_document_parse_simulation_counts_match_expected_reviewed_filtering(self):
        inputs = _document_parse_simulation_inputs()

        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=True,
            expected_removed_count=2,
            expected_kept_count=4)

        checks = simulation['consistency_checks']
        self.assertTrue(checks['phase_2b_removed_match'])
        self.assertTrue(checks['phase_2b_kept_match'])
        self.assertTrue(checks['phase_2c_body_region_removed_match'])
        self.assertEqual(simulation['summary']['simulated_kept_count'], 4)

    def test_document_parse_simulation_reports_missing_review_decisions(self):
        inputs = _document_parse_simulation_inputs()
        inputs['review_decisions'] = {'decisions': [], 'summary': {'decision_counts': {}}}

        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=True,
            expected_removed_count=0,
            expected_kept_count=6)

        warning_types = {warning['type'] for warning in simulation['safety_warnings']}
        self.assertIn('missing_review_decisions', warning_types)
        self.assertIn('no_approved_candidates', warning_types)
        self.assertEqual(simulation['summary']['would_remove_block_count'], 0)

    def test_document_parse_simulation_disabled_mode_is_clear(self):
        inputs = _document_parse_simulation_inputs()

        simulation = build_document_parse_filtering_simulation_report(**inputs)

        self.assertFalse(simulation['enabled'])
        self.assertFalse(simulation['applied'])
        self.assertEqual(simulation['summary']['would_remove_block_count'], 0)
        self.assertEqual(simulation['simulated_apply']['kept_block_count'], 6)

    def test_document_parse_hook_dry_run_reports_without_production_removal(self):
        inputs = _document_parse_simulation_inputs()

        hook = build_document_parse_filtering_hook_report(
            **inputs,
            enabled=True,
            apply=False,
            expected_removed_count=2,
            expected_kept_count=4)

        self.assertEqual(hook['hook_location'], 'Pages._parse_document()')
        self.assertEqual(hook['mode'], 'dry_run_report_only')
        self.assertEqual(hook['summary']['original_block_count'], 6)
        self.assertEqual(hook['summary']['would_remove_block_count'], 2)
        self.assertEqual(hook['summary']['simulated_removed_count'], 0)
        self.assertEqual(hook['summary']['production_removed_count'], 0)
        self.assertFalse(hook['summary']['production_objects_mutated'])
        self.assertFalse(hook['summary']['default_behavior_changed'])

    def test_document_parse_hook_uses_only_explicit_approval(self):
        inputs = _document_parse_simulation_inputs()

        hook = build_document_parse_filtering_hook_report(
            **inputs,
            enabled=True,
            apply=False,
            expected_removed_count=2,
            expected_kept_count=4)

        self.assertEqual(hook['summary']['approved_candidate_count'], 2)
        self.assertEqual(hook['summary']['blocked_candidate_count'], 3)
        self.assertEqual(hook['summary']['rejected_removed_count'], 0)
        self.assertEqual(hook['summary']['unsure_removed_count'], 0)
        self.assertEqual(hook['summary']['layout_placeholder_removed_count'], 0)

    def test_document_parse_hook_counts_match_simulation_helper(self):
        inputs = _document_parse_simulation_inputs()

        hook = build_document_parse_filtering_hook_report(
            **inputs,
            enabled=True,
            apply=False,
            expected_removed_count=2,
            expected_kept_count=4)
        simulation = build_document_parse_filtering_simulation_report(
            **inputs,
            enabled=True,
            apply=False,
            expected_removed_count=2,
            expected_kept_count=4)

        self.assertEqual(
            hook['summary']['would_remove_block_count'],
            simulation['summary']['would_remove_block_count'])
        self.assertEqual(
            hook['phase_2k_consistency']['would_remove_count_matches_phase_2k'],
            True)
        self.assertEqual(
            hook['phase_2k_consistency']['would_keep_count_matches_phase_2k'],
            True)

    def test_document_parse_hook_reports_missing_review_decisions(self):
        inputs = _document_parse_simulation_inputs()
        inputs['review_decisions'] = {'decisions': [], 'summary': {'decision_counts': {}}}

        hook = build_document_parse_filtering_hook_report(
            **inputs,
            enabled=True,
            apply=False,
            expected_removed_count=0,
            expected_kept_count=6)

        warning_types = {warning['type'] for warning in hook['safety_warnings']}
        self.assertIn('missing_review_decisions', warning_types)
        self.assertIn('no_approved_candidates', warning_types)
        self.assertFalse(hook['recommendation']['safe_to_attempt_phase_2m'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_parse_document_default_behavior_remains_unchanged(self):
        pages = Pages()

        self.assertEqual(Pages._parse_document([object()]), ('', ''))
        self.assertIsNone(pages._document_parse_filtering_hook_report)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_document_parse_hook_stores_report_without_mutating_input(self):
        inputs = _document_parse_simulation_inputs()
        layout_report = _document_parse_hook_layout_report(inputs)
        before = json.loads(json.dumps(layout_report))
        pages = Pages()

        report = pages._run_document_parse_filtering_hook(
            layout_report,
            _document_parse_filtering_hook_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_filtering_body_diff_report=inputs['body_filtering_diff_report'],
            _document_parse_filtering_paragraph_integrity_report=inputs['paragraph_integrity_report'],
            _document_parse_filtering_expected_removed_count=2,
            _document_parse_filtering_expected_kept_count=4)

        self.assertEqual(layout_report, before)
        self.assertIs(pages._document_parse_filtering_hook_report, report)
        self.assertEqual(report['summary']['would_remove_block_count'], 2)
        self.assertEqual(report['summary']['simulated_removed_count'], 0)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_document_parse_hook_disabled_leaves_report_empty(self):
        inputs = _document_parse_simulation_inputs()
        layout_report = _document_parse_hook_layout_report(inputs)
        pages = Pages()

        report = pages._run_document_parse_filtering_hook(layout_report)

        self.assertIsNone(report)
        self.assertIsNone(pages._document_parse_filtering_hook_report)

    def test_raw_object_mapping_maps_approved_summary_to_one_raw_object(self):
        inputs = _document_parse_raw_mapping_inputs()

        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        self.assertEqual(mapping['summary']['approved_candidate_count'], 2)
        self.assertEqual(mapping['summary']['expected_would_remove_count'], 2)
        self.assertEqual(mapping['summary']['mapped_raw_object_count'], 2)
        self.assertEqual(mapping['summary']['exact_match_count'], 2)
        self.assertEqual(mapping['summary']['ambiguous_match_count'], 0)
        self.assertEqual(mapping['summary']['missing_match_count'], 0)
        self.assertTrue(mapping['summary']['all_expected_blocks_mapped_once'])

    def test_raw_object_mapping_does_not_map_rejected_unsure_or_placeholder(self):
        inputs = _document_parse_raw_mapping_inputs()

        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)
        mapped_fingerprints = {
            raw_object['fingerprint']
            for item in mapping['mappings']
            for raw_object in item['selected_raw_objects']
        }

        self.assertNotIn('reject-header', mapped_fingerprints)
        self.assertNotIn('unsure-footer', mapped_fingerprints)
        self.assertNotIn('image-placeholder', mapped_fingerprints)
        self.assertEqual(
            mapping['summary']['rejected_unsure_layout_placeholder_matched_for_removal_count'],
            0)

    def test_raw_object_mapping_does_not_map_body_region_for_removal(self):
        inputs = _document_parse_raw_mapping_inputs()
        inputs['raw_object_pages'][0]['raw_objects'].append(
            _raw_object(6, 'Approved Header', REGION_BODY))

        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        self.assertEqual(mapping['summary']['mapped_raw_object_count'], 2)
        self.assertEqual(mapping['summary']['body_region_matched_for_removal_count'], 0)
        self.assertEqual(mapping['summary']['unsafe_match_count'], 0)

    def test_raw_object_mapping_warns_when_raw_object_is_missing(self):
        inputs = _document_parse_raw_mapping_inputs()
        inputs['raw_object_pages'][0]['raw_objects'] = [
            raw_object
            for raw_object in inputs['raw_object_pages'][0]['raw_objects']
            if raw_object['fingerprint'] != _summary_fingerprint('Approved Footer', REGION_BOTTOM)
        ]

        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        warning_types = {warning['type'] for warning in mapping['safety_warnings']}
        self.assertEqual(mapping['summary']['missing_match_count'], 1)
        self.assertIn('missing_raw_object_match', warning_types)
        self.assertFalse(mapping['recommendation']['safe_to_attempt_phase_2n'])

    def test_raw_object_mapping_warns_when_raw_match_is_ambiguous(self):
        inputs = _document_parse_raw_mapping_inputs()
        inputs['raw_object_pages'][0]['raw_objects'].append(
            _raw_object(6, 'Approved Header', REGION_TOP))

        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        warning_types = {warning['type'] for warning in mapping['safety_warnings']}
        self.assertEqual(mapping['summary']['ambiguous_match_count'], 1)
        self.assertIn('ambiguous_raw_object_match', warning_types)

    def test_raw_object_mapping_reports_fuzzy_match_separately(self):
        inputs = _document_parse_raw_mapping_inputs()
        inputs['raw_object_pages'][0]['raw_objects'][0]['bbox'] = [52.0, 22.0, 302.0, 42.0]

        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        self.assertEqual(mapping['summary']['exact_match_count'], 1)
        self.assertEqual(mapping['summary']['fuzzy_match_count'], 1)
        self.assertEqual(mapping['summary']['mapped_raw_object_count'], 2)

    def test_raw_object_mapping_does_not_mutate_inputs(self):
        inputs = _document_parse_raw_mapping_inputs()
        before_pages = json.loads(json.dumps(inputs['page_summaries']))
        before_raw = json.loads(json.dumps(inputs['raw_object_pages']))

        build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        self.assertEqual(inputs['page_summaries'], before_pages)
        self.assertEqual(inputs['raw_object_pages'], before_raw)

    def test_raw_object_mapping_disabled_mode_is_clear(self):
        inputs = _document_parse_raw_mapping_inputs()

        mapping = build_document_parse_raw_object_mapping_report(**inputs)

        self.assertFalse(mapping['enabled'])
        self.assertEqual(mapping['summary']['mapped_raw_object_count'], 0)
        self.assertFalse(mapping['recommendation']['safe_to_attempt_phase_2n'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_raw_object_mapping_hook_stores_report_without_mutating_raw_pages(self):
        inputs = _document_parse_raw_mapping_inputs()
        layout_report = _document_parse_hook_layout_report(inputs)
        fake_pages = [_FakePage(0)]
        fake_raw_pages = [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])]
        before_raw_texts = [block.text for block in fake_raw_pages[0].blocks]
        pages = Pages()

        report = pages._run_document_parse_raw_object_mapping_validation(
            layout_report,
            fake_pages,
            fake_raw_pages,
            _document_parse_raw_object_mapping_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_mapping_expected_would_remove_count=2)

        self.assertIs(pages._document_parse_raw_object_mapping_report, report)
        self.assertEqual(report['summary']['mapped_raw_object_count'], 2)
        self.assertEqual([block.text for block in fake_raw_pages[0].blocks], before_raw_texts)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_raw_object_mapping_disabled_leaves_report_empty(self):
        inputs = _document_parse_raw_mapping_inputs()
        layout_report = _document_parse_hook_layout_report(inputs)
        pages = Pages()

        report = pages._run_document_parse_raw_object_mapping_validation(
            layout_report,
            [_FakePage(0)],
            [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])])

        self.assertIsNone(report)
        self.assertIsNone(pages._document_parse_raw_object_mapping_report)

    def test_copied_raw_apply_removes_only_approved_mapped_objects(self):
        inputs = _document_parse_raw_mapping_inputs()

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(
            **inputs,
            enabled=True,
            expected_mapping_count=2)

        self.assertEqual(apply_report['summary']['removed_copied_block_count'], 2)
        self.assertEqual(apply_report['summary']['copied_filtered_block_count'], 4)
        self.assertEqual(apply_report['removed_counts_by_role'], {
            ROLE_FOOTER: 1,
            ROLE_HEADER: 1,
        })
        self.assertTrue(apply_report['summary']['removed_count_matches_phase_2m'])

    def test_copied_raw_apply_does_not_mutate_original_raw_objects(self):
        inputs = _document_parse_raw_mapping_inputs()
        before_raw = json.loads(json.dumps(inputs['raw_object_pages']))

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(
            **inputs,
            enabled=True,
            expected_mapping_count=2)

        self.assertEqual(inputs['raw_object_pages'], before_raw)
        self.assertFalse(apply_report['summary']['original_objects_mutated'])
        self.assertIsNot(
            apply_report['copied_filtered_pages'][0],
            inputs['raw_object_pages'][0])

    def test_copied_raw_apply_filters_copied_objects_as_expected(self):
        inputs = _document_parse_raw_mapping_inputs()

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(
            **inputs,
            enabled=True,
            expected_mapping_count=2)
        remaining = {
            raw_object['fingerprint']
            for page in apply_report['copied_filtered_pages']
            for raw_object in page['raw_objects']
        }

        self.assertNotIn(_summary_fingerprint('Approved Header', REGION_TOP), remaining)
        self.assertNotIn(_summary_fingerprint('Approved Footer', REGION_BOTTOM), remaining)
        self.assertIn(_summary_fingerprint('Body paragraph', REGION_BODY), remaining)

    def test_copied_raw_apply_keeps_rejected_unsure_placeholder_and_body(self):
        inputs = _document_parse_raw_mapping_inputs()

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(
            **inputs,
            enabled=True,
            expected_mapping_count=2)
        remaining = {
            raw_object['fingerprint']
            for page in apply_report['copied_filtered_pages']
            for raw_object in page['raw_objects']
        }

        self.assertIn(_summary_fingerprint('Rejected Header', REGION_TOP), remaining)
        self.assertIn(_summary_fingerprint('Unsure Footer', REGION_BOTTOM), remaining)
        self.assertIn(_summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP), remaining)
        self.assertIn(_summary_fingerprint('Body paragraph', REGION_BODY), remaining)
        self.assertEqual(apply_report['summary']['body_region_removed_count'], 0)
        self.assertEqual(
            apply_report['summary']['rejected_unsure_layout_placeholder_removed_count'],
            0)

    def test_copied_raw_apply_warns_for_missing_mapping(self):
        inputs = _document_parse_raw_mapping_inputs()
        inputs['raw_object_pages'][0]['raw_objects'] = [
            raw_object
            for raw_object in inputs['raw_object_pages'][0]['raw_objects']
            if raw_object['fingerprint'] != _summary_fingerprint('Approved Footer', REGION_BOTTOM)
        ]

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(
            **inputs,
            enabled=True,
            expected_mapping_count=2)

        warning_types = {warning['type'] for warning in apply_report['safety_warnings']}
        self.assertEqual(apply_report['summary']['removed_copied_block_count'], 1)
        self.assertIn('mapping_missing_raw_object_match', warning_types)
        self.assertIn('expected_mapping_count_mismatch_phase_2m', warning_types)

    def test_copied_raw_apply_warns_for_ambiguous_mapping(self):
        inputs = _document_parse_raw_mapping_inputs()
        inputs['raw_object_pages'][0]['raw_objects'].append(
            _raw_object(6, 'Approved Header', REGION_TOP))

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(
            **inputs,
            enabled=True,
            expected_mapping_count=2)

        warning_types = {warning['type'] for warning in apply_report['safety_warnings']}
        self.assertEqual(apply_report['summary']['removed_copied_block_count'], 1)
        self.assertIn('mapping_ambiguous_raw_object_match', warning_types)
        self.assertIn('expected_mapping_count_mismatch_phase_2m', warning_types)

    def test_copied_raw_apply_count_matches_phase_2m_mapping_count(self):
        inputs = _document_parse_raw_mapping_inputs()
        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(
            **inputs,
            raw_object_mapping_report=mapping,
            enabled=True,
            expected_mapping_count=2)

        consistency = apply_report['consistency_with_phase_2m']
        self.assertEqual(consistency['phase_2m_mapped_raw_object_count'], 2)
        self.assertTrue(consistency['removed_count_matches_phase_2m'])
        self.assertTrue(apply_report['recommendation']['safe_to_attempt_phase_2o'])

    def test_copied_raw_apply_disabled_mode_is_clear(self):
        inputs = _document_parse_raw_mapping_inputs()

        apply_report = build_document_parse_copied_raw_page_filtering_apply_report(**inputs)

        self.assertFalse(apply_report['enabled'])
        self.assertFalse(apply_report['applied_to_copy'])
        self.assertEqual(apply_report['summary']['removed_copied_block_count'], 0)
        self.assertFalse(apply_report['recommendation']['safe_to_attempt_phase_2o'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_copied_raw_apply_path_stores_report_without_mutating_fake_raw_pages(self):
        inputs = _document_parse_raw_mapping_inputs()
        layout_report = _document_parse_hook_layout_report(inputs)
        fake_pages = [_FakePage(0)]
        fake_raw_pages = [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])]
        before_raw_texts = [block.text for block in fake_raw_pages[0].blocks]
        pages = Pages()

        report = pages._run_document_parse_copied_raw_filtering_apply(
            layout_report,
            fake_pages,
            fake_raw_pages,
            _document_parse_copied_raw_filtering_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_copied_raw_filtering_expected_mapping_count=2)

        self.assertIs(pages._document_parse_copied_raw_filtering_apply_report, report)
        self.assertEqual(report['summary']['removed_copied_block_count'], 2)
        self.assertEqual([block.text for block in fake_raw_pages[0].blocks], before_raw_texts)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_copied_raw_apply_disabled_leaves_report_empty(self):
        inputs = _document_parse_raw_mapping_inputs()
        layout_report = _document_parse_hook_layout_report(inputs)
        pages = Pages()

        report = pages._run_document_parse_copied_raw_filtering_apply(
            layout_report,
            [_FakePage(0)],
            [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])])

        self.assertIsNone(report)
        self.assertIsNone(pages._document_parse_copied_raw_filtering_apply_report)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_guarded_raw_apply_removes_only_approved_during_apply_and_restores(self):
        inputs = _document_parse_raw_mapping_inputs()
        fake_pages = [_FakePage(0)]
        fake_raw_pages = [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])]
        before_raw_texts = [block.text for block in fake_raw_pages[0].blocks]

        report = Pages()._run_document_parse_guarded_raw_apply_restore(
            _document_parse_hook_layout_report(inputs),
            fake_pages,
            fake_raw_pages,
            _document_parse_guarded_raw_apply_restore_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_guarded_raw_apply_restore_expected_mapping_count=2)

        self.assertEqual(report['summary']['removed_during_apply_count'], 2)
        self.assertEqual(report['summary']['filtered_raw_block_count_during_apply'], 4)
        self.assertEqual(report['summary']['restored_raw_block_count_after_restore'], 6)
        self.assertTrue(report['summary']['restore_exact_count_match'])
        self.assertTrue(report['summary']['restore_fingerprint_match'])
        self.assertFalse(report['summary']['original_raw_pages_left_mutated'])
        self.assertEqual([block.text for block in fake_raw_pages[0].blocks], before_raw_texts)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_guarded_raw_apply_keeps_rejected_unsure_placeholder_and_body(self):
        inputs = _document_parse_raw_mapping_inputs()
        report = Pages()._run_document_parse_guarded_raw_apply_restore(
            _document_parse_hook_layout_report(inputs),
            [_FakePage(0)],
            [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])],
            _document_parse_guarded_raw_apply_restore_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_guarded_raw_apply_restore_expected_mapping_count=2)

        self.assertEqual(report['summary']['body_region_removed_count'], 0)
        self.assertEqual(
            report['summary']['rejected_unsure_layout_placeholder_removed_count'],
            0)
        self.assertEqual(
            report['downstream_risk_notes']['body_block_count_before'],
            report['downstream_risk_notes']['body_block_count_during_apply'])
        self.assertEqual(
            report['downstream_risk_notes']['image_shape_placeholder_count_before'],
            report['downstream_risk_notes']['image_shape_placeholder_count_during_apply'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_guarded_raw_apply_warns_and_skips_when_mapping_missing(self):
        inputs = _document_parse_raw_mapping_inputs()
        raw_objects = [
            raw_object
            for raw_object in inputs['raw_object_pages'][0]['raw_objects']
            if raw_object['fingerprint'] != _summary_fingerprint('Approved Footer', REGION_BOTTOM)
        ]
        fake_raw_pages = [_FakeRawPage(raw_objects)]

        report = Pages()._run_document_parse_guarded_raw_apply_restore(
            _document_parse_hook_layout_report(inputs),
            [_FakePage(0)],
            fake_raw_pages,
            _document_parse_guarded_raw_apply_restore_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_guarded_raw_apply_restore_expected_mapping_count=2)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertEqual(report['summary']['removed_during_apply_count'], 0)
        self.assertIn('mapping_missing_raw_object_match', warning_types)
        self.assertIn('guarded_apply_skipped', warning_types)
        self.assertTrue(report['summary']['restore_exact_count_match'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_guarded_raw_apply_warns_and_skips_when_mapping_ambiguous(self):
        inputs = _document_parse_raw_mapping_inputs()
        raw_objects = list(inputs['raw_object_pages'][0]['raw_objects'])
        raw_objects.append(_raw_object(6, 'Approved Header', REGION_TOP))

        report = Pages()._run_document_parse_guarded_raw_apply_restore(
            _document_parse_hook_layout_report(inputs),
            [_FakePage(0)],
            [_FakeRawPage(raw_objects)],
            _document_parse_guarded_raw_apply_restore_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_guarded_raw_apply_restore_expected_mapping_count=2)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertEqual(report['summary']['removed_during_apply_count'], 0)
        self.assertIn('mapping_ambiguous_raw_object_match', warning_types)
        self.assertIn('guarded_apply_skipped', warning_types)
        self.assertTrue(report['summary']['restore_fingerprint_match'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_guarded_raw_apply_count_matches_phase_2m_and_2n(self):
        inputs = _document_parse_raw_mapping_inputs()
        report = Pages()._run_document_parse_guarded_raw_apply_restore(
            _document_parse_hook_layout_report(inputs),
            [_FakePage(0)],
            [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])],
            _document_parse_guarded_raw_apply_restore_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_guarded_raw_apply_restore_expected_mapping_count=2)

        self.assertTrue(report['consistency']['removed_count_matches_phase_2m'])
        self.assertTrue(report['consistency']['removed_count_matches_phase_2n'])
        self.assertTrue(report['recommendation']['safe_to_attempt_phase_2p'])

    def test_guarded_raw_apply_disabled_mode_is_clear(self):
        inputs = _document_parse_raw_mapping_inputs()

        report = build_document_parse_guarded_raw_page_apply_restore_report(
            raw_object_pages_before=inputs['raw_object_pages'])

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['removed_during_apply_count'], 0)
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2p'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_guarded_raw_apply_disabled_leaves_report_empty(self):
        inputs = _document_parse_raw_mapping_inputs()
        pages = Pages()

        report = pages._run_document_parse_guarded_raw_apply_restore(
            _document_parse_hook_layout_report(inputs),
            [_FakePage(0)],
            [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])])

        self.assertIsNone(report)
        self.assertIsNone(pages._document_parse_guarded_raw_apply_restore_report)

    def test_filtered_parse_experiment_reports_baseline_and_filtered_metrics(self):
        inputs = _document_parse_raw_mapping_inputs()
        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)
        report = build_document_parse_filtered_parse_experiment_report(
            raw_object_pages_before=inputs['raw_object_pages'],
            raw_object_pages_filtered=[{
                'page_index': 0,
                'page_number': 1,
                'width': 600,
                'height': 1000,
                'raw_objects': inputs['raw_object_pages'][0]['raw_objects'][1:5],
            }],
            raw_object_pages_after=inputs['raw_object_pages'],
            removed_objects_by_page=[{
                'page_index': 0,
                'page_number': 1,
                'removed_count': 2,
                'objects': [
                    dict(inputs['raw_object_pages'][0]['raw_objects'][0], proposed_role=ROLE_HEADER),
                    dict(inputs['raw_object_pages'][0]['raw_objects'][2], proposed_role=ROLE_FOOTER),
                ],
            }],
            baseline_parse_metrics=_parse_metrics(raw_count=6, text_blocks=5, body_text_blocks=1),
            filtered_parse_metrics=_parse_metrics(raw_count=4, text_blocks=3, body_text_blocks=1),
            raw_object_mapping_report=mapping,
            enabled=True,
            restore_completed=True,
            restore_fingerprint_match=True,
            expected_mapping_count=2)

        self.assertEqual(report['summary']['baseline_raw_block_count'], 6)
        self.assertEqual(report['summary']['filtered_raw_block_count'], 4)
        self.assertEqual(report['summary']['removed_raw_block_count'], 2)
        self.assertEqual(report['summary']['baseline_parsed_text_block_count'], 5)
        self.assertEqual(report['summary']['filtered_parsed_text_block_count'], 3)
        self.assertEqual(report['summary']['baseline_body_text_block_count'], 1)
        self.assertEqual(report['summary']['filtered_body_text_block_count'], 1)
        self.assertTrue(report['summary']['raw_pages_restored_or_reloaded'])
        self.assertTrue(report['summary']['restore_fingerprint_match'])
        self.assertFalse(report['summary']['production_default_changed'])

    def test_filtered_parse_experiment_keeps_rejected_unsure_placeholder_and_body(self):
        inputs = _document_parse_raw_mapping_inputs()
        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)
        report = build_document_parse_filtered_parse_experiment_report(
            raw_object_pages_before=inputs['raw_object_pages'],
            raw_object_pages_filtered=inputs['raw_object_pages'],
            raw_object_pages_after=inputs['raw_object_pages'],
            removed_objects_by_page=[],
            baseline_parse_metrics=_parse_metrics(raw_count=6, body_text_blocks=1),
            filtered_parse_metrics=_parse_metrics(raw_count=6, body_text_blocks=1),
            raw_object_mapping_report=mapping,
            enabled=True,
            restore_completed=True,
            restore_fingerprint_match=True,
            expected_mapping_count=2)

        self.assertEqual(report['summary']['body_region_removed_count'], 0)
        self.assertEqual(
            report['summary']['rejected_unsure_layout_placeholder_removed_count'],
            0)

    def test_filtered_parse_experiment_warns_on_body_drop_and_table_changes(self):
        inputs = _document_parse_raw_mapping_inputs()
        mapping = build_document_parse_raw_object_mapping_report(
            **inputs,
            enabled=True,
            expected_would_remove_count=2)

        report = build_document_parse_filtered_parse_experiment_report(
            raw_object_pages_before=inputs['raw_object_pages'],
            raw_object_pages_filtered=inputs['raw_object_pages'],
            raw_object_pages_after=inputs['raw_object_pages'],
            removed_objects_by_page=[],
            baseline_parse_metrics=_parse_metrics(
                raw_count=6,
                body_text_blocks=4,
                tables=1,
                paragraph_like=3),
            filtered_parse_metrics=_parse_metrics(
                raw_count=6,
                body_text_blocks=3,
                tables=0,
                paragraph_like=5),
            raw_object_mapping_report=mapping,
            enabled=True,
            restore_completed=True,
            restore_fingerprint_match=True,
            expected_mapping_count=2)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertIn('body_text_block_count_dropped', warning_types)
        self.assertIn('table_count_changed', warning_types)
        self.assertIn('paragraph_fragmentation_increased', warning_types)

    def test_filtered_parse_experiment_disabled_mode_is_clear(self):
        inputs = _document_parse_raw_mapping_inputs()

        report = build_document_parse_filtered_parse_experiment_report(
            raw_object_pages_before=inputs['raw_object_pages'])

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['removed_raw_block_count'], 0)
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2q'])

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_filtered_parse_experiment_removes_only_during_experiment_and_restores(self):
        inputs = _document_parse_raw_mapping_inputs()
        fake_pages = [_FakePage(0)]
        fake_raw_pages = [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])]
        before_raw_texts = [block.text for block in fake_raw_pages[0].blocks]

        report = Pages()._run_document_parse_filtered_parse_experiment(
            _document_parse_hook_layout_report(inputs),
            fake_pages,
            fake_raw_pages,
            _document_parse_filtered_parse_experiment_enabled=True,
            _document_parse_filtering_review_decisions=inputs['review_decisions'],
            _document_parse_filtered_parse_expected_mapping_count=2)

        self.assertEqual(report['summary']['removed_raw_block_count'], 2)
        self.assertEqual(report['summary']['filtered_raw_block_count'], 4)
        self.assertTrue(report['summary']['raw_pages_restored_or_reloaded'])
        self.assertTrue(report['summary']['restore_fingerprint_match'])
        self.assertEqual([block.text for block in fake_raw_pages[0].blocks], before_raw_texts)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_pages_filtered_parse_experiment_disabled_leaves_report_empty(self):
        inputs = _document_parse_raw_mapping_inputs()
        pages = Pages()

        report = pages._run_document_parse_filtered_parse_experiment(
            _document_parse_hook_layout_report(inputs),
            [_FakePage(0)],
            [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])])

        self.assertIsNone(report)
        self.assertIsNone(pages._document_parse_filtered_parse_experiment_report)

    def test_table_delta_report_detects_baseline_only_tables(self):
        baseline = _parse_metrics_with_tables([
            _table_record('base-top', 0, REGION_TOP, [50, 20, 300, 44]),
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460]),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460]),
        ])

        report = build_table_delta_investigation_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[_removed_objects_page([
                _removed_raw_object('Approved Header', REGION_TOP, [50, 20, 300, 44], ROLE_HEADER),
            ])],
            enabled=True)

        self.assertEqual(report['summary']['baseline_only_table_count'], 1)
        self.assertEqual(report['baseline_only_tables'][0]['classification'], 'likely_header_footer_false_positive')

    def test_table_delta_classifies_boundary_overlap_as_likely_pollution_removed(self):
        baseline = _parse_metrics_with_tables([
            _table_record('base-bottom', 0, REGION_BOTTOM, [50, 940, 300, 980]),
        ])
        filtered = _parse_metrics_with_tables([])

        report = build_table_delta_investigation_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[_removed_objects_page([
                _removed_raw_object('Approved Footer', REGION_BOTTOM, [50, 942, 300, 978], ROLE_FOOTER),
            ])],
            enabled=True)

        self.assertEqual(report['summary']['likely_header_footer_false_positive_table_count'], 1)
        self.assertEqual(report['summary']['suspicious_body_table_loss_count'], 0)
        self.assertTrue(report['recommendation']['safe_to_attempt_phase_2r'])

    def test_table_delta_body_region_loss_triggers_warning(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460]),
        ])
        filtered = _parse_metrics_with_tables([])

        report = build_table_delta_investigation_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[],
            enabled=True)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertIn('body_region_table_disappeared', warning_types)
        self.assertEqual(report['summary']['suspicious_body_table_loss_count'], 1)
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2r'])

    def test_table_delta_changed_common_geometry_triggers_warning(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=2),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [51, 400, 520, 470], rows=3, cols=2),
        ])

        report = build_table_delta_investigation_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertEqual(report['summary']['changed_common_table_count'], 1)
        self.assertIn('common_table_changed', warning_types)

    def test_table_delta_unchanged_table_counts_have_no_warning(self):
        table = _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460])
        report = build_table_delta_investigation_report(
            baseline_parse_metrics=_parse_metrics_with_tables([table]),
            filtered_parse_metrics=_parse_metrics_with_tables([table]),
            enabled=True)

        self.assertEqual(report['summary']['table_count_delta'], 0)
        self.assertEqual(report['safety_warnings'], [])
        self.assertEqual(report['summary']['classification'], 'no_delta')

    def test_table_delta_filtered_only_tables_are_reported(self):
        report = build_table_delta_investigation_report(
            baseline_parse_metrics=_parse_metrics_with_tables([]),
            filtered_parse_metrics=_parse_metrics_with_tables([
                _table_record('new-table', 0, REGION_BODY, [50, 400, 520, 460]),
            ]),
            enabled=True)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertEqual(report['summary']['filtered_only_table_count'], 1)
        self.assertIn('filtered_only_table_detected', warning_types)

    def test_table_delta_report_does_not_mutate_inputs(self):
        baseline = _parse_metrics_with_tables([
            _table_record('base-top', 0, REGION_TOP, [50, 20, 300, 44]),
        ])
        filtered = _parse_metrics_with_tables([])
        removed = [_removed_objects_page([
            _removed_raw_object('Approved Header', REGION_TOP, [50, 20, 300, 44], ROLE_HEADER),
        ])]
        before = json.dumps({
            'baseline': baseline,
            'filtered': filtered,
            'removed': removed,
        }, sort_keys=True)

        build_table_delta_investigation_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=removed,
            enabled=True)

        after = json.dumps({
            'baseline': baseline,
            'filtered': filtered,
            'removed': removed,
        }, sort_keys=True)
        self.assertEqual(before, after)

    def test_table_delta_disabled_mode_is_clear(self):
        report = build_table_delta_investigation_report(
            baseline_parse_metrics=_parse_metrics_with_tables([
                _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460]),
            ]))

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['baseline_only_table_count'], 0)
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2r'])

    def test_body_table_root_cause_body_baseline_only_defaults_to_unsafe(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], rows=3, cols=3),
        ])
        filtered = _parse_metrics_with_tables([])

        report = build_body_table_delta_root_cause_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        finding = report['baseline_only_findings'][0]
        self.assertEqual(finding['likely_cause'], 'possible_real_body_table_loss')
        self.assertEqual(finding['severity'], 'unsafe')
        self.assertEqual(report['summary']['possible_real_body_table_loss_count'], 1)

    def test_body_table_root_cause_boundary_overlap_is_likely_pollution(self):
        baseline = _parse_metrics_with_tables([
            _table_record('footer-table', 0, REGION_BOTTOM, [50, 940, 300, 980], rows=1, cols=3),
        ])
        filtered = _parse_metrics_with_tables([])

        report = build_body_table_delta_root_cause_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[_removed_objects_page([
                _removed_raw_object('Approved Footer', REGION_BOTTOM, [50, 942, 300, 978], ROLE_FOOTER),
            ])],
            enabled=True)

        finding = report['baseline_only_findings'][0]
        self.assertEqual(finding['likely_cause'], 'header_footer_pollution_removed')
        self.assertEqual(finding['severity'], 'safe')
        self.assertEqual(report['summary']['likely_header_footer_pollution_table_count'], 1)

    def test_body_table_root_cause_changed_boundary_near_artifact_can_be_safe(self):
        baseline = _parse_metrics_with_tables([
            _table_record('footer-table', 0, REGION_BOTTOM, [50, 940, 300, 980], rows=1, cols=3),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('footer-table-filtered', 0, REGION_BOTTOM, [50, 940, 300, 970], rows=1, cols=3),
        ])

        report = build_body_table_delta_root_cause_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[_removed_objects_page([
                _removed_raw_object('Approved Footer', REGION_BOTTOM, [50, 972, 300, 980], ROLE_FOOTER),
            ])],
            enabled=True)

        finding = report['changed_common_findings'][0]
        self.assertEqual(finding['likely_cause'], 'table_geometry_changed_near_removed_artifact')
        self.assertEqual(finding['severity'], 'safe')

    def test_body_table_root_cause_changed_body_cell_loss_is_unsafe(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=3),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=2),
        ])

        report = build_body_table_delta_root_cause_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        finding = report['changed_common_findings'][0]
        self.assertEqual(finding['likely_cause'], 'possible_real_body_table_loss')
        self.assertEqual(finding['severity'], 'unsafe')

    def test_body_table_root_cause_body_false_positive_can_be_review(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-artifact-table', 0, REGION_BODY, [50, 900, 300, 940], rows=1, cols=3),
        ])
        filtered = _parse_metrics_with_tables([])

        report = build_body_table_delta_root_cause_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[_removed_objects_page([
                _removed_raw_object('Approved Footer', REGION_BOTTOM, [50, 902, 300, 938], ROLE_FOOTER),
            ])],
            enabled=True)

        finding = report['baseline_only_findings'][0]
        self.assertEqual(finding['likely_cause'], 'baseline_false_positive_table')
        self.assertEqual(finding['severity'], 'review')
        self.assertEqual(report['summary']['likely_false_positive_table_count'], 1)

    def test_body_table_root_cause_reports_overlap_and_distance(self):
        baseline = _parse_metrics_with_tables([
            _table_record('footer-table', 0, REGION_BOTTOM, [50, 940, 300, 980], rows=1, cols=3),
        ])
        filtered = _parse_metrics_with_tables([])

        report = build_body_table_delta_root_cause_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[_removed_objects_page([
                _removed_raw_object('Approved Footer', REGION_BOTTOM, [50, 942, 300, 978], ROLE_FOOTER),
            ])],
            enabled=True)
        proximity = report['baseline_only_findings'][0]['removed_candidate_proximity']

        self.assertEqual(proximity['overlap_count'], 1)
        self.assertEqual(proximity['nearest_distance'], 0.0)
        self.assertEqual(report['overlap_proximity_summary']['tables_overlapping_removed_candidates'], 1)

    def test_body_table_root_cause_does_not_mutate_inputs(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460]),
        ])
        filtered = _parse_metrics_with_tables([])
        removed = [_removed_objects_page([
            _removed_raw_object('Approved Header', REGION_TOP, [50, 20, 300, 44], ROLE_HEADER),
        ])]
        before = json.dumps({
            'baseline': baseline,
            'filtered': filtered,
            'removed': removed,
        }, sort_keys=True)

        build_body_table_delta_root_cause_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=removed,
            enabled=True)

        after = json.dumps({
            'baseline': baseline,
            'filtered': filtered,
            'removed': removed,
        }, sort_keys=True)
        self.assertEqual(before, after)

    def test_body_table_root_cause_disabled_mode_is_clear(self):
        report = build_body_table_delta_root_cause_report(
            baseline_parse_metrics=_parse_metrics_with_tables([
                _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460]),
            ]))

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['classification'], 'disabled')
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2s'])

    def test_body_table_geometry_bbox_only_shift_is_review_safe(self):
        cell_bboxes = [[50, 400, 285, 430], [285, 400, 520, 430]]
        baseline = _parse_metrics_with_tables([
            _table_record(
                'body-table', 0, REGION_BODY, [50, 400, 520, 460],
                rows=1, cols=2, cells=2, cell_bboxes=cell_bboxes),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record(
                'body-table-filtered', 0, REGION_BODY, [50, 400, 520, 450],
                rows=1, cols=2, cells=2, cell_bboxes=cell_bboxes),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)
        finding = report['findings'][0]

        self.assertEqual(finding['likely_cause'], 'harmless_bbox_boundary_shift')
        self.assertEqual(finding['severity'], 'review')
        self.assertTrue(finding['text_cell_signature_preserved'])

    def test_body_table_geometry_row_count_change_is_unsafe(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=2, cells=4),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 400, 520, 460], rows=3, cols=2, cells=4),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        self.assertEqual(report['findings'][0]['likely_cause'], 'possible_body_table_structure_change')
        self.assertEqual(report['findings'][0]['severity'], 'unsafe')

    def test_body_table_geometry_column_count_change_is_unsafe(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=2, cells=4),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=3, cells=4),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        self.assertEqual(report['findings'][0]['likely_cause'], 'possible_body_table_structure_change')
        self.assertEqual(report['findings'][0]['severity'], 'unsafe')

    def test_body_table_geometry_cell_count_change_is_unsafe(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=2, cells=4),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=2, cells=3),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        self.assertEqual(report['findings'][0]['likely_cause'], 'possible_cell_loss')
        self.assertEqual(report['findings'][0]['severity'], 'unsafe')

    def test_body_table_geometry_cell_text_signature_change_is_unsafe(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], cell_texts=['A', 'B', 'C', 'D']),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 400, 520, 450], cell_texts=['A', 'B', 'C', 'X']),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        self.assertEqual(report['findings'][0]['likely_cause'], 'possible_body_table_structure_change')
        self.assertEqual(report['findings'][0]['severity'], 'unsafe')

    def test_body_table_geometry_bbox_edge_shift_near_removed_candidate_is_detected(self):
        cell_bboxes = [[50, 400, 285, 430], [285, 400, 520, 430]]
        baseline = _parse_metrics_with_tables([
            _table_record(
                'body-table', 0, REGION_BODY, [50, 400, 520, 460],
                rows=1, cols=2, cells=2, cell_bboxes=cell_bboxes),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record(
                'body-table-filtered', 0, REGION_BODY, [50, 400, 520, 450],
                rows=1, cols=2, cells=2, cell_bboxes=cell_bboxes),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            removed_objects_by_page=[_removed_objects_page([
                _removed_raw_object('Approved Footer', REGION_BOTTOM, [50, 455, 520, 470], ROLE_FOOTER),
            ])],
            enabled=True)
        finding = report['findings'][0]

        self.assertTrue(finding['changed_bbox_edge_near_removed_candidate'])
        self.assertEqual(finding['likely_cause'], 'header_footer_boundary_cleanup')
        self.assertEqual(finding['severity'], 'safe')

    def test_body_table_geometry_body_shrink_without_text_loss_is_review(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 470]),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 410, 520, 460]),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)
        finding = report['findings'][0]

        self.assertEqual(finding['likely_cause'], 'stream_table_boundary_adjustment')
        self.assertEqual(finding['severity'], 'review')
        self.assertTrue(finding['changed_area_intersects_body_text'])

    def test_body_table_geometry_insufficient_evidence_remains_review(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 470], include_cells=False),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 410, 520, 460], include_cells=False),
        ])

        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        self.assertEqual(report['findings'][0]['likely_cause'], 'insufficient_evidence')
        self.assertEqual(report['findings'][0]['severity'], 'review')

    def test_body_table_geometry_report_does_not_mutate_inputs(self):
        baseline = _parse_metrics_with_tables([
            _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 470]),
        ])
        filtered = _parse_metrics_with_tables([
            _table_record('body-table-filtered', 0, REGION_BODY, [50, 410, 520, 460]),
        ])
        before = json.dumps({'baseline': baseline, 'filtered': filtered}, sort_keys=True)

        build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=baseline,
            filtered_parse_metrics=filtered,
            enabled=True)

        after = json.dumps({'baseline': baseline, 'filtered': filtered}, sort_keys=True)
        self.assertEqual(before, after)

    def test_body_table_geometry_disabled_mode_is_clear(self):
        report = build_body_table_geometry_delta_safety_report(
            baseline_parse_metrics=_parse_metrics_with_tables([
                _table_record('body-table', 0, REGION_BODY, [50, 400, 520, 470]),
            ]))

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['classification'], 'disabled')
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2t'])

    def test_table_geometry_visual_review_pack_includes_all_changed_items(self):
        safety_report = _geometry_safety_report_with_preserved_items(8)

        report = build_table_geometry_visual_review_pack(
            safety_report,
            visual_rendering={'supported': True, 'output_directory': 'local_reports/table_geometry_review'},
            enabled=True)

        self.assertTrue(report['enabled'])
        self.assertEqual(report['summary']['review_item_count'], 8)
        self.assertEqual(report['summary']['affected_pages'], [5, 8, 10])
        self.assertEqual(len(report['review_items']), 8)
        for item in report['review_items']:
            self.assertIn('baseline_bbox', item)
            self.assertIn('filtered_bbox', item)
            self.assertIn('row_count_before', item)
            self.assertIn('row_count_after', item)
            self.assertIn('human_decision_fields', item)

    def test_table_geometry_visual_review_preserved_items_need_human_approval(self):
        safety_report = _geometry_safety_report_with_preserved_items(1)

        report = build_table_geometry_visual_review_pack(
            safety_report,
            visual_rendering={'supported': True},
            enabled=True)
        item = report['review_items'][0]

        self.assertEqual(item['review_classification'], 'likely_safe_but_needs_human_approval')
        self.assertTrue(item['human_approval_required'])
        self.assertEqual(report['summary']['requiring_human_approval_count'], 1)
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2u'])

    def test_table_geometry_visual_review_unsafe_geometry_change_is_marked_unsafe(self):
        safety_report = _geometry_safety_report_from_tables(
            [_table_record('body-table', 0, REGION_BODY, [50, 400, 520, 460], rows=2, cols=2, cells=4)],
            [_table_record('body-table-filtered', 0, REGION_BODY, [50, 400, 520, 455], rows=3, cols=2, cells=4)])

        report = build_table_geometry_visual_review_pack(
            safety_report,
            visual_rendering={'supported': True},
            enabled=True)
        item = report['review_items'][0]

        self.assertEqual(item['review_classification'], 'unsafe_do_not_integrate')
        self.assertFalse(item['human_approval_required'])
        self.assertEqual(report['summary']['automatically_unsafe_count'], 1)

    def test_table_geometry_visual_review_reports_missing_visual_rendering_support(self):
        safety_report = _geometry_safety_report_with_preserved_items(1)

        report = build_table_geometry_visual_review_pack(
            safety_report,
            visual_rendering={'supported': False, 'skipped_reason': 'fitz_unavailable'},
            enabled=True)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertIn('visual_rendering_unavailable', warning_types)
        self.assertEqual(report['visual_rendering']['skipped_reason'], 'fitz_unavailable')

    def test_table_geometry_visual_review_report_does_not_mutate_inputs(self):
        safety_report = _geometry_safety_report_with_preserved_items(1)
        before = json.dumps(safety_report, sort_keys=True)

        build_table_geometry_visual_review_pack(
            safety_report,
            visual_rendering={'supported': True},
            enabled=True)

        after = json.dumps(safety_report, sort_keys=True)
        self.assertEqual(before, after)

    def test_table_geometry_visual_review_disabled_mode_is_clear(self):
        safety_report = _geometry_safety_report_with_preserved_items(1)

        report = build_table_geometry_visual_review_pack(
            safety_report,
            enabled=False)

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['classification'], 'disabled')
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2u'])

    def test_table_visual_approval_gate_passes_when_all_expected_items_approved(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(item_count=8))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        self.assertEqual(report['gate_status'], 'passed')
        self.assertEqual(report['summary']['approve_count'], 8)
        self.assertEqual(report['summary']['reject_count'], 0)
        self.assertEqual(report['summary']['unsure_count'], 0)
        self.assertEqual(report['summary']['missing_decision_count'], 0)
        self.assertTrue(report['recommendation']['safe_to_attempt_phase_2v'])

    def test_table_visual_approval_gate_blocks_rejected_item(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(
                item_count=8,
                decisions={3: 'reject_unsafe_table_change'}))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        self.assertEqual(report['gate_status'], 'blocked')
        self.assertEqual(report['summary']['reject_count'], 1)
        self.assertIn(
            'rejected_items_present',
            {reason['type'] for reason in report['blocking_reasons']})

    def test_table_visual_approval_gate_blocks_unsure_item(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(
                item_count=8,
                decisions={2: 'unsure'}))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        self.assertEqual(report['gate_status'], 'blocked')
        self.assertEqual(report['summary']['unsure_count'], 1)
        self.assertIn(
            'unsure_items_present',
            {reason['type'] for reason in report['blocking_reasons']})

    def test_table_visual_approval_gate_blocks_missing_decision(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(
                item_count=8,
                decisions={4: None}))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        self.assertEqual(report['gate_status'], 'blocked')
        self.assertEqual(report['summary']['missing_decision_count'], 1)

    def test_table_visual_approval_gate_blocks_item_count_mismatch(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(item_count=7))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        self.assertEqual(report['gate_status'], 'blocked')
        self.assertIn(
            'parsed_review_item_count_mismatch',
            {reason['type'] for reason in report['blocking_reasons']})

    def test_table_visual_approval_gate_blocks_row_column_cell_inconsistency(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(
                item_count=8,
                changed_counts={5}))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        self.assertEqual(report['gate_status'], 'blocked')
        self.assertEqual(report['summary']['row_column_cell_preservation_count'], 7)
        self.assertIn(
            'row_column_cell_counts_not_fully_preserved',
            {reason['type'] for reason in report['blocking_reasons']})

    def test_table_visual_approval_gate_blocks_text_signature_inconsistency(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(
                item_count=8,
                changed_text={6}))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        self.assertEqual(report['gate_status'], 'blocked')
        self.assertEqual(report['summary']['text_cell_signature_preservation_count'], 7)
        self.assertIn(
            'text_cell_signatures_not_fully_preserved',
            {reason['type'] for reason in report['blocking_reasons']})

    def test_table_visual_review_markdown_parsing_is_whitespace_tolerant(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(item_count=1, extra_whitespace=True))

        self.assertEqual(decisions['summary']['review_item_count'], 1)
        self.assertEqual(decisions['items'][0]['manual_decision'], 'approve_safe_boundary_shift')
        self.assertTrue(decisions['items'][0]['row_column_cell_counts_preserved'])

    def test_table_visual_approval_gate_report_does_not_mutate_inputs(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(item_count=8))
        before = json.dumps(decisions, sort_keys=True)

        build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=True)

        after = json.dumps(decisions, sort_keys=True)
        self.assertEqual(before, after)

    def test_table_visual_approval_gate_disabled_mode_is_clear(self):
        decisions = parse_table_geometry_visual_review_markdown(
            _table_visual_review_markdown(item_count=8))

        report = build_table_geometry_visual_approval_gate_report(
            decisions,
            expected_review_item_count=8,
            enabled=False)

        self.assertFalse(report['enabled'])
        self.assertEqual(report['gate_status'], 'blocked')
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2v'])

    def test_filtered_docx_experiment_requires_explicit_enablement(self):
        report = build_filtered_docx_generation_comparison_report(
            filtered_parse_experiment_report=_filtered_docx_experiment_report(),
            table_visual_approval_gate_report=_passed_table_visual_gate_report(),
            baseline_docx_path='local_reports/docx_compare/baseline.docx',
            filtered_docx_path='local_reports/docx_compare/filtered.docx')

        self.assertFalse(report['enabled'])
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2w'])

    def test_filtered_docx_experiment_blocks_missing_gate(self):
        report = build_filtered_docx_generation_comparison_report(
            filtered_parse_experiment_report=_filtered_docx_experiment_report(),
            baseline_docx_path='local_reports/docx_compare/baseline.docx',
            filtered_docx_path='local_reports/docx_compare/filtered.docx',
            baseline_docx_metrics=_docx_metrics(),
            filtered_docx_metrics=_docx_metrics(),
            normal_conversion_check=_normal_conversion_check(),
            enabled=True)

        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2w'])
        self.assertIn(
            'table_visual_approval_gate_not_passed',
            {warning['type'] for warning in report['safety_warnings']})

    def test_filtered_docx_experiment_blocks_failed_gate(self):
        gate = _passed_table_visual_gate_report()
        gate['gate_status'] = 'blocked'
        gate['summary']['gate_status'] = 'blocked'

        report = build_filtered_docx_generation_comparison_report(
            filtered_parse_experiment_report=_filtered_docx_experiment_report(),
            table_visual_approval_gate_report=gate,
            baseline_docx_path='local_reports/docx_compare/baseline.docx',
            filtered_docx_path='local_reports/docx_compare/filtered.docx',
            baseline_docx_metrics=_docx_metrics(),
            filtered_docx_metrics=_docx_metrics(),
            normal_conversion_check=_normal_conversion_check(),
            enabled=True)

        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2w'])
        self.assertEqual(report['summary']['table_visual_approval_gate_status'], 'blocked')

    def test_filtered_docx_experiment_reports_local_only_paths(self):
        report = build_filtered_docx_generation_comparison_report(
            filtered_parse_experiment_report=_filtered_docx_experiment_report(),
            table_visual_approval_gate_report=_passed_table_visual_gate_report(),
            baseline_docx_path='local_reports/docx_compare/baseline.docx',
            filtered_docx_path='local_reports/docx_compare/filtered.docx',
            baseline_docx_metrics=_docx_metrics(),
            filtered_docx_metrics=_docx_metrics(),
            normal_conversion_check=_normal_conversion_check(),
            enabled=True)

        self.assertTrue(report['docx_files']['baseline']['local_only_path'])
        self.assertTrue(report['docx_files']['filtered']['local_only_path'])

    def test_filtered_docx_experiment_reports_reviewed_removal_and_metrics(self):
        report = build_filtered_docx_generation_comparison_report(
            filtered_parse_experiment_report=_filtered_docx_experiment_report(),
            table_visual_approval_gate_report=_passed_table_visual_gate_report(),
            baseline_docx_path='local_reports/docx_compare/baseline.docx',
            filtered_docx_path='local_reports/docx_compare/filtered.docx',
            baseline_docx_metrics=_docx_metrics(paragraphs=20, tables=3),
            filtered_docx_metrics=_docx_metrics(paragraphs=17, tables=2),
            normal_conversion_check=_normal_conversion_check(),
            enabled=True)

        summary = report['summary']
        self.assertEqual(summary['removed_approved_header_footer_page_number_count'], 48)
        self.assertEqual(summary['body_region_removed_count'], 0)
        self.assertEqual(summary['baseline_parsed_text_block_count'], 523)
        self.assertEqual(summary['filtered_parsed_text_block_count'], 486)
        self.assertEqual(summary['baseline_body_text_block_count'], 393)
        self.assertEqual(summary['filtered_body_text_block_count'], 393)
        self.assertEqual(summary['baseline_table_count'], 139)
        self.assertEqual(summary['filtered_table_count'], 127)

    def test_filtered_docx_experiment_reports_missing_and_empty_files(self):
        report = build_filtered_docx_generation_comparison_report(
            filtered_parse_experiment_report=_filtered_docx_experiment_report(),
            table_visual_approval_gate_report=_passed_table_visual_gate_report(),
            baseline_docx_path='local_reports/docx_compare/missing-baseline.docx',
            filtered_docx_path='local_reports/docx_compare/empty-filtered.docx',
            filtered_docx_metrics=_docx_metrics(size=0),
            normal_conversion_check=_normal_conversion_check(),
            enabled=True)

        warning_types = {warning['type'] for warning in report['safety_warnings']}
        self.assertIn('baseline_docx_missing', warning_types)
        self.assertIn('filtered_docx_empty', warning_types)

    def test_filtered_docx_experiment_requires_state_reload_confirmation(self):
        report = build_filtered_docx_generation_comparison_report(
            filtered_parse_experiment_report=_filtered_docx_experiment_report(),
            table_visual_approval_gate_report=_passed_table_visual_gate_report(),
            baseline_docx_path='local_reports/docx_compare/baseline.docx',
            filtered_docx_path='local_reports/docx_compare/filtered.docx',
            baseline_docx_metrics=_docx_metrics(),
            filtered_docx_metrics=_docx_metrics(),
            normal_conversion_check={'passed': True, 'state_restored_or_reloaded': False},
            enabled=True)

        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2w'])
        self.assertIn(
            'state_restore_or_reload_not_confirmed',
            {warning['type'] for warning in report['safety_warnings']})

    def test_filtered_docx_residual_body_paragraph_location_is_reported(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            baseline = Path(tmpdir) / 'baseline.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            _write_docx_fixture(baseline, body_paragraphs=['Residual text', 'Residual text'])
            _write_docx_fixture(filtered, body_paragraphs=['Residual text'])

            report = build_filtered_docx_residual_structure_report(
                str(baseline),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

        self.assertEqual(report['summary']['residual_removed_string_count'], 1)
        self.assertEqual(report['residuals'][0]['locations'][0]['location_type'], 'body_paragraph')
        self.assertEqual(report['residuals'][0]['classification'], 'legitimate_body_duplicate')

    def test_filtered_docx_residual_table_cell_location_is_reported(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            baseline = Path(tmpdir) / 'baseline.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            _write_docx_fixture(baseline, table_cells=['Residual text'])
            _write_docx_fixture(filtered, table_cells=['Residual text'])

            report = build_filtered_docx_residual_structure_report(
                str(baseline),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

        self.assertEqual(report['residuals'][0]['locations'][0]['location_type'], 'table_cell')
        self.assertEqual(report['residuals'][0]['classification'], 'legitimate_body_or_table_content')

    def test_filtered_docx_residual_header_footer_parts_are_reported_separately(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            baseline = Path(tmpdir) / 'baseline.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            _write_docx_fixture(baseline, body_paragraphs=['Body'])
            _write_docx_fixture(
                filtered,
                body_paragraphs=['Body'],
                header_texts=['Residual text'],
                footer_texts=['Residual text'])

            report = build_filtered_docx_residual_structure_report(
                str(baseline),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

        location_types = {
            location['location_type']
            for location in report['residuals'][0]['locations']
        }
        self.assertIn('header_part', location_types)
        self.assertIn('footer_part', location_types)
        self.assertEqual(report['residuals'][0]['classification'], 'docx_header_footer_part_content')

    def test_filtered_docx_true_repeated_header_footer_residual_warns(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            baseline = Path(tmpdir) / 'baseline.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            _write_docx_fixture(baseline, body_paragraphs=['Residual text', 'Residual text'])
            _write_docx_fixture(filtered, body_paragraphs=['Residual text', 'Residual text'])

            report = build_filtered_docx_residual_structure_report(
                str(baseline),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

        self.assertEqual(report['residuals'][0]['classification'], 'true_residual_header_footer_pollution')
        self.assertIn(
            'true_residual_header_footer_pollution',
            {warning['type'] for warning in report['safety_warnings']})

    def test_filtered_docx_residual_reports_paragraph_and_table_deltas(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            baseline = Path(tmpdir) / 'baseline.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            _write_docx_fixture(baseline, body_paragraphs=['A', 'B'], table_cells=['C'])
            _write_docx_fixture(filtered, body_paragraphs=['A'], table_cells=[])

            report = build_filtered_docx_residual_structure_report(
                str(baseline),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

        self.assertEqual(report['summary']['baseline_body_paragraph_count'], 2)
        self.assertEqual(report['summary']['filtered_body_paragraph_count'], 1)
        self.assertEqual(report['summary']['paragraph_delta'], -1)
        self.assertEqual(report['summary']['baseline_table_count'], 1)
        self.assertEqual(report['summary']['filtered_table_count'], 0)
        self.assertEqual(report['summary']['table_delta'], -1)

    def test_filtered_docx_residual_missing_docx_is_reported(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            missing = Path(tmpdir) / 'missing.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            _write_docx_fixture(filtered, body_paragraphs=['Body'])

            report = build_filtered_docx_residual_structure_report(
                str(missing),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

        self.assertIn('docx_missing', {warning['type'] for warning in report['safety_warnings']})

    def test_filtered_docx_residual_empty_docx_is_reported(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            baseline = Path(tmpdir) / 'baseline.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            baseline.write_bytes(b'')
            _write_docx_fixture(filtered, body_paragraphs=['Body'])

            report = build_filtered_docx_residual_structure_report(
                str(baseline),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

        self.assertIn('docx_empty', {warning['type'] for warning in report['safety_warnings']})

    def test_filtered_docx_residual_report_does_not_mutate_docx_files(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            baseline = Path(tmpdir) / 'baseline.docx'
            filtered = Path(tmpdir) / 'filtered.docx'
            _write_docx_fixture(baseline, body_paragraphs=['Residual text'])
            _write_docx_fixture(filtered, body_paragraphs=['Residual text'])
            before = {
                'baseline': baseline.read_bytes(),
                'filtered': filtered.read_bytes(),
            }

            build_filtered_docx_residual_structure_report(
                str(baseline),
                str(filtered),
                removed_strings=['Residual text'],
                enabled=True)

            after = {
                'baseline': baseline.read_bytes(),
                'filtered': filtered.read_bytes(),
            }

        self.assertEqual(before, after)

    def test_filtered_docx_residual_disabled_mode_is_clear(self):
        report = build_filtered_docx_residual_structure_report(
            baseline_docx_path='missing-baseline.docx',
            filtered_docx_path='missing-filtered.docx',
            removed_strings=['Residual text'])

        self.assertFalse(report['enabled'])
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2x'])

    def test_readiness_gate_passes_when_all_evidence_is_safe(self):
        report = build_reviewed_filtering_feature_readiness_report(
            **_safe_readiness_inputs(),
            enabled=True)

        self.assertEqual(
            report['readiness_status'],
            'ready_for_internal_opt_in_integration_experiment')
        self.assertEqual(report['blocking_reasons'], [])
        self.assertTrue(report['recommendation']['safe_to_attempt_phase_2y'])

    def test_readiness_gate_blocks_when_header_footer_approval_missing(self):
        inputs = _safe_readiness_inputs()
        inputs['header_footer_review_report']['approved_candidate_count'] = 0
        inputs['raw_object_mapping_report']['summary']['approved_candidate_count'] = 0

        report = build_reviewed_filtering_feature_readiness_report(
            **inputs,
            enabled=True)

        self.assertEqual(report['readiness_status'], 'blocked')
        self.assertIn(
            'header_footer_review_approval_missing',
            _readiness_reason_types(report))

    def test_readiness_gate_blocks_when_table_visual_gate_is_missing_or_failed(self):
        inputs = _safe_readiness_inputs()
        inputs['table_visual_approval_gate_report'] = {
            'gate_status': 'blocked',
            'summary': {
                'gate_status': 'blocked',
                'expected_review_item_count': 8,
                'parsed_review_item_count': 8,
                'approve_count': 7,
                'reject_count': 1,
                'unsure_count': 0,
                'missing_decision_count': 0,
            },
        }

        report = build_reviewed_filtering_feature_readiness_report(
            **inputs,
            enabled=True)

        self.assertEqual(report['readiness_status'], 'blocked')
        self.assertIn(
            'table_visual_approval_gate_not_passed',
            _readiness_reason_types(report))

    def test_readiness_gate_blocks_when_body_region_removal_is_nonzero(self):
        inputs = _safe_readiness_inputs()
        inputs['filtered_parse_experiment_report']['summary']['body_region_removed_count'] = 1

        report = build_reviewed_filtering_feature_readiness_report(
            **inputs,
            enabled=True)

        self.assertIn('body_region_removed', _readiness_reason_types(report))

    def test_readiness_gate_blocks_true_residual_header_footer_pollution(self):
        inputs = _safe_readiness_inputs()
        inputs['docx_residual_structure_report']['summary'][
            'true_residual_header_footer_pollution_count'] = 1
        inputs['docx_residual_structure_report']['summary']['classification'] = 'unsafe'

        report = build_reviewed_filtering_feature_readiness_report(
            **inputs,
            enabled=True)

        self.assertIn(
            'true_residual_header_footer_pollution_present',
            _readiness_reason_types(report))

    def test_readiness_gate_blocks_body_text_loss_warnings(self):
        inputs = _safe_readiness_inputs()
        inputs['docx_residual_structure_report']['summary']['body_text_loss_warning_count'] = 1
        inputs['docx_residual_structure_report']['summary']['classification'] = 'unsafe'

        report = build_reviewed_filtering_feature_readiness_report(
            **inputs,
            enabled=True)

        self.assertIn(
            'body_text_loss_warnings_present',
            _readiness_reason_types(report))

    def test_readiness_gate_blocks_table_text_loss_warnings(self):
        inputs = _safe_readiness_inputs()
        inputs['docx_residual_structure_report']['summary']['table_text_loss_warning_count'] = 1
        inputs['docx_residual_structure_report']['summary']['classification'] = 'unsafe'

        report = build_reviewed_filtering_feature_readiness_report(
            **inputs,
            enabled=True)

        self.assertIn(
            'table_text_loss_warnings_present',
            _readiness_reason_types(report))

    def test_readiness_gate_records_local_sample_dependency_as_non_blocking_risk(self):
        report = build_reviewed_filtering_feature_readiness_report(
            **_safe_readiness_inputs(),
            enabled=True)

        self.assertEqual(
            report['readiness_status'],
            'ready_for_internal_opt_in_integration_experiment')
        self.assertIn(
            'local_sample_dependency',
            {risk['type'] for risk in report['non_blocking_risks']})

    def test_readiness_gate_report_does_not_mutate_inputs(self):
        inputs = _safe_readiness_inputs()
        before = json.loads(json.dumps(inputs))

        build_reviewed_filtering_feature_readiness_report(
            **inputs,
            enabled=True)

        self.assertEqual(inputs, before)

    def test_readiness_gate_disabled_mode_is_clear(self):
        report = build_reviewed_filtering_feature_readiness_report(
            **_safe_readiness_inputs())

        self.assertFalse(report['enabled'])
        self.assertEqual(report['readiness_status'], 'blocked')
        self.assertIn('readiness_gate_disabled', _readiness_reason_types(report))

    def test_corpus_summary_builder_handles_multiple_sample_results(self):
        report = build_local_corpus_validation_summary_report(
            [
                _corpus_sample_result('input2.pdf', _corpus_layout_report()),
                _corpus_sample_result('input3.pdf', _corpus_layout_report()),
            ],
            enabled=True)

        self.assertEqual(report['summary']['sample_count'], 2)
        self.assertEqual(report['summary']['samples_analyzed_successfully'], 2)
        self.assertEqual(
            report['summary']['samples_with_likely_valid_header_footer_candidates'],
            2)
        self.assertEqual(report['summary']['samples_needing_manual_review'], 2)

    def test_corpus_summary_reports_failed_sample_analysis_clearly(self):
        report = build_local_corpus_validation_summary_report(
            [_corpus_sample_result(
                'broken.pdf',
                {},
                parsing_succeeded=False,
                analysis_succeeded=False,
                error='cannot open')],
            enabled=True)

        self.assertEqual(report['summary']['samples_failed_analysis'], 1)
        self.assertIn('parsing_failed', _corpus_warning_types(report))
        self.assertEqual(report['samples'][0]['recommendation']['label'], 'analysis_failed')

    def test_corpus_summary_marks_large_sample_analysis_only(self):
        report = build_local_corpus_validation_summary_report(
            [_corpus_sample_result(
                'input6_large.pdf',
                _corpus_layout_report(),
                page_count=756,
                pages_analyzed=15,
                analysis_mode='analysis_only_bounded_subset')],
            enabled=True,
            large_page_threshold=100)

        sample = report['samples'][0]
        self.assertTrue(sample['basic_file_summary']['large_sample'])
        self.assertTrue(sample['basic_file_summary']['partial_or_bounded'])
        self.assertIn('large_sample_analysis_only', _corpus_warning_types(report))
        self.assertEqual(report['summary']['samples_too_large_for_full_pipeline'], 1)

    def test_corpus_summary_does_not_mutate_input_reports(self):
        sample_results = [_corpus_sample_result('input2.pdf', _corpus_layout_report())]
        before = json.loads(json.dumps(sample_results))

        build_local_corpus_validation_summary_report(sample_results, enabled=True)

        self.assertEqual(sample_results, before)

    def test_corpus_summary_disabled_mode_is_clear(self):
        report = build_local_corpus_validation_summary_report(
            [_corpus_sample_result('input2.pdf', _corpus_layout_report())])

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['sample_count'], 0)
        self.assertFalse(report['recommendation']['safe_to_attempt_phase_2y1'])

    def test_corpus_manual_review_summary_handles_multiple_selected_samples(self):
        packs = [
            build_local_corpus_manual_review_pack(
                _corpus_sample_result('input3.pdf', _corpus_layout_report()),
                enabled=True),
            build_local_corpus_manual_review_pack(
                _corpus_sample_result(
                    'input6_large.pdf',
                    _corpus_layout_report(),
                    page_count=756,
                    pages_analyzed=15,
                    analysis_mode='analysis_only_bounded_subset'),
                enabled=True),
        ]

        report = build_local_corpus_manual_review_summary_report(
            packs,
            enabled=True)

        self.assertEqual(report['summary']['selected_sample_count'], 2)
        self.assertEqual(report['summary']['review_packs_ready_count'], 2)
        self.assertEqual(report['summary']['total_would_exclude_candidate_count'], 4)
        self.assertEqual(report['summary']['total_would_remove_block_count'], 12)

    def test_corpus_manual_review_pack_marks_candidates_ready_for_review(self):
        pack = build_local_corpus_manual_review_pack(
            _corpus_sample_result('input3.pdf', _corpus_layout_report()),
            enabled=True)

        self.assertTrue(pack['summary']['ready_for_human_approval'])
        self.assertEqual(
            pack['summary']['recommended_next_action'],
            'manual_approve_then_full_local_pipeline')
        self.assertTrue(pack['summary']['manual_approval_required'])

    def test_corpus_manual_review_pack_marks_large_sample_bounded_only(self):
        pack = build_local_corpus_manual_review_pack(
            _corpus_sample_result(
                'input6_large.pdf',
                _corpus_layout_report(),
                page_count=756,
                pages_analyzed=15,
                analysis_mode='analysis_only_bounded_subset'),
            enabled=True)

        self.assertTrue(pack['summary']['bounded_analysis_only'])
        self.assertEqual(
            pack['summary']['recommended_next_action'],
            'analysis_only_large_sample')
        self.assertIn(
            'bounded_large_sample_review',
            {warning['type'] for warning in pack['warnings']})

    def test_corpus_manual_review_pack_creates_no_auto_approval(self):
        pack = build_local_corpus_manual_review_pack(
            _corpus_sample_result('input3.pdf', _corpus_layout_report()),
            enabled=True)

        self.assertEqual(pack['summary']['auto_approved_decision_count'], 0)
        self.assertTrue(all(
            not item['manual_decision_fields']['approve_exclude']
            for item in pack['review_items']))
        self.assertTrue(all(
            item['auto_approved'] is False
            for item in pack['review_items']))

    def test_corpus_manual_review_summary_reports_missing_pack_clearly(self):
        report = build_local_corpus_manual_review_summary_report(
            [None],
            enabled=True)

        self.assertEqual(report['summary']['missing_review_pack_count'], 1)
        self.assertIn(
            'missing_corpus_manual_review_pack',
            {warning['type'] for warning in report['warnings']})

    def test_corpus_manual_review_pack_generation_does_not_mutate_input(self):
        sample = _corpus_sample_result('input3.pdf', _corpus_layout_report())
        before = json.loads(json.dumps(sample))

        build_local_corpus_manual_review_pack(sample, enabled=True)

        self.assertEqual(sample, before)

    def test_corpus_manual_review_disabled_mode_is_clear(self):
        pack = build_local_corpus_manual_review_pack(
            _corpus_sample_result('input3.pdf', _corpus_layout_report()))
        summary = build_local_corpus_manual_review_summary_report([pack])

        self.assertFalse(pack['enabled'])
        self.assertFalse(summary['enabled'])
        self.assertEqual(pack['summary']['candidate_count'], 0)
        self.assertEqual(summary['summary']['selected_sample_count'], 0)

    def test_corpus_approval_validation_passes_with_explicit_decisions(self):
        report = build_local_corpus_approval_validation_report(
            _corpus_approval_layout_report(),
            _corpus_review_decisions(
                ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'approve_exclude'),
                ('repeated-2', 'page-number||bottom', 'page_number', ACTION_WOULD_EXCLUDE, 'reject_exclude'),
                ('repeated-3', 'body repeat||body', 'review_only', ACTION_REVIEW, 'unsure'),
            ),
            sample_name='input3.pdf',
            enabled=True)

        self.assertTrue(report['summary']['explicit_decisions_complete'])
        self.assertEqual(report['summary']['approve_count'], 1)
        self.assertEqual(report['summary']['reject_count'], 1)
        self.assertEqual(report['summary']['unsure_count'], 1)
        self.assertEqual(report['summary']['eligible_approved_candidate_count'], 1)
        self.assertEqual(report['summary']['reviewed_removed_block_count'], 2)
        self.assertEqual(report['summary']['unsafe_removed_count'], 0)

    def test_corpus_approval_validation_blocks_missing_decisions(self):
        report = build_local_corpus_approval_validation_report(
            _corpus_approval_layout_report(),
            _corpus_review_decisions(
                ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'approve_exclude')),
            sample_name='input3.pdf',
            enabled=True)

        self.assertFalse(report['summary']['explicit_decisions_complete'])
        self.assertEqual(report['summary']['missing_decision_count'], 2)
        self.assertIn('missing_review_decisions', _corpus_warning_types(report))
        self.assertFalse(
            report['recommendation']['safe_to_run_approved_only_validation'])

    def test_corpus_approval_validation_reports_unsure_candidates(self):
        report = build_local_corpus_approval_validation_report(
            _corpus_approval_layout_report(),
            _corpus_review_decisions(
                ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'approve_exclude'),
                ('repeated-2', 'page-number||bottom', 'page_number', ACTION_WOULD_EXCLUDE, 'approve_exclude'),
                ('repeated-3', 'body repeat||body', 'review_only', ACTION_REVIEW, 'unsure'),
            ),
            sample_name='input3.pdf',
            enabled=True)

        self.assertEqual(report['summary']['unsure_count'], 1)
        self.assertEqual(report['summary']['unsure_removed_count'], 0)
        self.assertIn('unsure_candidates_present', _corpus_warning_types(report))

    def test_corpus_approval_validation_removes_only_approved_candidates(self):
        report = build_local_corpus_approval_validation_report(
            _corpus_approval_layout_report(),
            _corpus_review_decisions(
                ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'approve_exclude'),
                ('repeated-2', 'page-number||bottom', 'page_number', ACTION_WOULD_EXCLUDE, 'reject_exclude'),
                ('repeated-3', 'body repeat||body', 'review_only', ACTION_REVIEW, 'reject_exclude'),
            ),
            sample_name='input3.pdf',
            enabled=True)

        self.assertEqual(report['summary']['reviewed_removed_block_count'], 2)
        self.assertEqual(report['summary']['rejected_removed_count'], 0)
        self.assertEqual(
            report['summary']['raw_would_exclude_without_approval_removed_count'],
            0)

    def test_corpus_approval_validation_keeps_rejected_candidates_blocked(self):
        report = build_local_corpus_approval_validation_report(
            _corpus_approval_layout_report(),
            _corpus_review_decisions(
                ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'reject_exclude'),
                ('repeated-2', 'page-number||bottom', 'page_number', ACTION_WOULD_EXCLUDE, 'reject_exclude'),
                ('repeated-3', 'body repeat||body', 'review_only', ACTION_REVIEW, 'reject_exclude'),
            ),
            sample_name='input3.pdf',
            enabled=True)

        self.assertEqual(report['summary']['eligible_approved_candidate_count'], 0)
        self.assertEqual(report['summary']['reviewed_removed_block_count'], 0)
        self.assertEqual(report['summary']['blocked_candidate_count'], 3)

    def test_corpus_approval_validation_keeps_bounded_large_sample_from_full_docx(self):
        report = build_local_corpus_approval_validation_report(
            _corpus_approval_layout_report(),
            _corpus_review_decisions(
                ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'approve_exclude'),
                ('repeated-2', 'page-number||bottom', 'page_number', ACTION_WOULD_EXCLUDE, 'approve_exclude'),
                ('repeated-3', 'body repeat||body', 'review_only', ACTION_REVIEW, 'reject_exclude'),
            ),
            sample_name='input6_large.pdf',
            bounded_analysis_only=True,
            full_docx_validation_allowed=False,
            enabled=True)

        self.assertTrue(report['summary']['full_docx_validation_blocked'])
        self.assertIn(
            'bounded_large_sample_full_docx_blocked',
            _corpus_warning_types(report))

    def test_corpus_approval_validation_does_not_mutate_inputs(self):
        layout_report = _corpus_approval_layout_report()
        decisions = _corpus_review_decisions(
            ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'approve_exclude'),
            ('repeated-2', 'page-number||bottom', 'page_number', ACTION_WOULD_EXCLUDE, 'reject_exclude'),
            ('repeated-3', 'body repeat||body', 'review_only', ACTION_REVIEW, 'unsure'),
        )
        before = json.loads(json.dumps({
            'layout_report': layout_report,
            'decisions': decisions,
        }))

        build_local_corpus_approval_validation_report(
            layout_report,
            decisions,
            sample_name='input3.pdf',
            enabled=True)

        self.assertEqual({'layout_report': layout_report, 'decisions': decisions}, before)

    def test_corpus_approval_validation_disabled_mode_is_clear(self):
        report = build_local_corpus_approval_validation_report(
            _corpus_approval_layout_report(),
            _corpus_review_decisions(
                ('repeated-1', 'annual report||top', 'header', ACTION_WOULD_EXCLUDE, 'approve_exclude')))
        summary = build_local_corpus_approval_validation_summary_report([report])

        self.assertFalse(report['enabled'])
        self.assertFalse(summary['enabled'])
        self.assertEqual(report['summary']['candidate_count'], 0)
        self.assertEqual(summary['summary']['sample_count'], 0)

    def test_reviewed_filtering_internal_config_default_is_disabled(self):
        config = build_reviewed_filtering_internal_config()
        report = build_reviewed_filtering_internal_config_report(
            config,
            _internal_config_dry_run_report(),
            _internal_config_decisions(
                ('header-1', 'synthetic header||top', 'approve_exclude')))

        self.assertFalse(config['enabled'])
        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['activation_status'], 'disabled')
        self.assertEqual(report['document_parse_settings'], {})

    def test_reviewed_filtering_missing_config_preserves_disabled_behavior(self):
        report = build_reviewed_filtering_internal_config_report(
            dry_run_report=_internal_config_dry_run_report(),
            review_decisions=_internal_config_decisions(
                ('header-1', 'synthetic header||top', 'approve_exclude')))

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['activation_status'], 'disabled')
        self.assertIn(
            'reviewed_filtering_config_disabled',
            _config_warning_types(report))

    def test_reviewed_filtering_enabled_false_preserves_disabled_behavior(self):
        report = build_reviewed_filtering_internal_config_report(
            {'enabled': False, 'mode': 'guarded_apply_restore'},
            _internal_config_dry_run_report(),
            _internal_config_decisions(
                ('header-1', 'synthetic header||top', 'approve_exclude')))

        self.assertFalse(report['enabled'])
        self.assertEqual(report['summary']['activation_status'], 'disabled')
        self.assertEqual(report['document_parse_settings'], {})

    def test_reviewed_filtering_enabled_without_review_decisions_is_blocked(self):
        report = build_reviewed_filtering_internal_config_report(
            {'enabled': True, 'mode': 'dry_run'},
            _internal_config_dry_run_report())

        self.assertEqual(report['summary']['activation_status'], 'blocked')
        self.assertIn('missing_review_decisions', _config_warning_types(report))
        self.assertFalse(report['recommendation']['safe_for_internal_experiment'])

    def test_reviewed_filtering_raw_would_exclude_without_approval_is_blocked(self):
        report = build_reviewed_filtering_internal_config_report(
            {'enabled': True, 'mode': 'dry_run'},
            _internal_config_dry_run_report(),
            {'decisions': [], 'summary': {'candidate_count': 0}})

        self.assertEqual(report['summary']['eligible_candidate_count'], 0)
        self.assertIn(
            'raw_would_exclude_without_approval_blocked',
            _config_warning_types(report))
        self.assertEqual(report['summary']['activation_status'], 'blocked')

    def test_reviewed_filtering_rejected_and_unsure_decisions_remain_blocked(self):
        report = build_reviewed_filtering_internal_config_report(
            {'enabled': True, 'mode': 'dry_run'},
            _internal_config_dry_run_report(),
            _internal_config_decisions(
                ('header-1', 'synthetic header||top', 'reject_exclude'),
                ('footer-1', 'synthetic footer||bottom', 'unsure')))

        blocked = {
            item['candidate_id']: item['blocked_reason']
            for item in report['candidates']
        }

        self.assertEqual(blocked['header-1'], 'rejected_candidate_blocked')
        self.assertEqual(blocked['footer-1'], 'unsure_candidate_blocked')
        self.assertEqual(report['summary']['eligible_candidate_count'], 0)

    def test_reviewed_filtering_body_region_candidates_remain_protected(self):
        report = build_reviewed_filtering_internal_config_report(
            {'enabled': True, 'mode': 'dry_run'},
            _internal_config_dry_run_report(),
            _internal_config_decisions(
                ('body-1', 'body repeat||body', 'approve_exclude')))

        body_row = [
            row for row in report['candidates']
            if row['candidate_id'] == 'body-1'
        ][0]

        self.assertFalse(body_row['eligible_for_reviewed_filtering'])
        self.assertEqual(body_row['blocked_reason'], 'body_region_protected')
        self.assertIn('body_region_candidates_protected', _config_warning_types(report))

    def test_reviewed_filtering_layout_placeholder_candidates_remain_protected(self):
        report = build_reviewed_filtering_internal_config_report(
            {'enabled': True, 'mode': 'dry_run'},
            _internal_config_dry_run_report(),
            _internal_config_decisions(
                ('placeholder-1', '<image>||top', 'approve_exclude')))

        placeholder_row = [
            row for row in report['candidates']
            if row['candidate_id'] == 'placeholder-1'
        ][0]

        self.assertFalse(placeholder_row['eligible_for_reviewed_filtering'])
        self.assertEqual(placeholder_row['blocked_reason'], 'layout_placeholder_protected')
        self.assertIn(
            'layout_placeholder_candidates_protected',
            _config_warning_types(report))

    def test_reviewed_filtering_config_summary_is_json_serializable(self):
        report = build_reviewed_filtering_internal_config_report(
            {
                'enabled': True,
                'mode': 'guarded_apply_restore',
                'page_subset': (0, 2, 4),
                'max_pages': '5',
            },
            _internal_config_safe_dry_run_report(),
            _internal_config_decisions(
                ('header-1', 'synthetic header||top', 'approve_exclude'),
                ('footer-1', 'synthetic footer||bottom', 'approve_exclude')))

        encoded = json.dumps(report)
        decoded = json.loads(encoded)

        self.assertEqual(decoded['config']['page_subset'], [0, 2, 4])
        self.assertEqual(decoded['config']['max_pages'], 5)
        self.assertEqual(
            decoded['summary']['activation_status'],
            'ready_for_internal_experiment')
        self.assertTrue(decoded['document_parse_settings']['layout_analysis'])
        self.assertTrue(
            decoded['document_parse_settings']['_document_parse_guarded_raw_apply_restore_enabled'])

    def test_reviewed_filtering_internal_config_does_not_change_public_defaults(self):
        _require_synthetic_pdf_support(self)
        doc = fitz.open()
        try:
            doc.new_page(width=72, height=72)
            stream = doc.tobytes()
        finally:
            doc.close()
        converter = Converter(stream=stream)
        try:
            settings = converter.default_settings
        finally:
            converter.close()

        self.assertNotIn('reviewed_header_footer_filtering', settings)
        self.assertNotIn('_document_parse_filtering_review_decisions', settings)
        self.assertEqual(
            reviewed_filtering_config_to_document_parse_settings(
                build_reviewed_filtering_internal_config()),
            {})

    def test_internal_config_missing_config_does_not_run_filtered_parse_integration(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'default-path.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')

            converter = Converter(str(pdf_path))
            settings = converter.default_settings.copy()
            try:
                converter.load_pages().parse_document(**settings)
                self.assertIsNone(converter.pages._reviewed_filtering_internal_config_report)
                self.assertIsNone(converter.pages._reviewed_filtering_internal_filtered_parse_report)
            finally:
                converter.close()

    def test_internal_config_enabled_false_does_not_apply_filtered_parse(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'disabled-config.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')

            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})
            report = _run_synthetic_internal_filtered_parse_integration(
                pdf_path,
                decisions,
                enabled=False)

            self.assertEqual(
                report['config']['summary']['activation_status'],
                'disabled')
            self.assertFalse(report['integration']['applied_to_parse'])
            self.assertEqual(report['integration']['summary']['removed_raw_block_count'], 0)

    def test_internal_config_enabled_without_review_decisions_blocks_integration(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'missing-decisions.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')

            report = _run_synthetic_internal_filtered_parse_integration(
                pdf_path,
                review_decisions=None)

            self.assertEqual(
                report['config']['summary']['activation_status'],
                'blocked')
            self.assertFalse(report['integration']['applied_to_parse'])
            self.assertIn(
                'config_not_ready',
                {warning['type'] for warning in report['integration']['safety_warnings']})

    def test_internal_config_raw_would_exclude_alone_blocks_filtered_parse(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'raw-would-exclude.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')

            report = _run_synthetic_internal_filtered_parse_integration(
                pdf_path,
                review_decisions={'decisions': [], 'summary': {'candidate_count': 0}})

            self.assertFalse(report['integration']['applied_to_parse'])
            self.assertIn(
                'config_raw_would_exclude_without_approval_blocked',
                {warning['type'] for warning in report['integration']['safety_warnings']})

    def test_internal_config_filtered_parse_applies_approved_synthetic_candidates(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'approved-filtered-parse.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')

            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})
            filtering = _synthetic_filtering_report(layout, decisions)
            report = _run_synthetic_internal_filtered_parse_integration(
                pdf_path,
                decisions,
                run_parse_pages=True)
            summary = report['integration']['summary']

            self.assertEqual(
                report['config']['summary']['activation_status'],
                'ready_for_internal_experiment')
            self.assertTrue(report['integration']['applied_to_parse'])
            self.assertEqual(
                summary['removed_raw_block_count'],
                filtering['summary']['removed_block_count'])
            self.assertEqual(summary['body_region_removed_count'], 0)
            self.assertEqual(summary['rejected_unsure_layout_placeholder_removed_count'], 0)
            self.assertTrue(summary['body_text_signature_preserved'])
            self.assertEqual(
                summary['filtered_raw_block_count'],
                summary['original_raw_block_count'] - summary['removed_raw_block_count'])

    def test_internal_config_fail_closed_blocks_rejected_and_unsure_decisions(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'fail-closed-rejected.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')

            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER})
            report = _run_synthetic_internal_filtered_parse_integration(
                pdf_path,
                decisions)
            warning_types = {
                warning['type']
                for warning in report['integration']['safety_warnings']
            }

            self.assertEqual(report['config']['summary']['activation_status'], 'blocked')
            self.assertFalse(report['integration']['applied_to_parse'])
            self.assertIn('config_rejected_candidates_blocked', warning_types)

    @unittest.skipIf(Pages is None, 'Pages import unavailable')
    def test_internal_config_filtered_parse_reports_textblock_delta_warning(self):
        inputs = _document_parse_raw_mapping_inputs()
        safe_decisions = {
            'decisions': [
                _review_decision(
                    'c-approved-header',
                    _summary_fingerprint('Approved Header', REGION_TOP),
                    'approve_exclude'),
                _review_decision(
                    'c-approved-footer',
                    _summary_fingerprint('Approved Footer', REGION_BOTTOM),
                    'approve_exclude'),
            ],
            'summary': {'decision_counts': {'approve_exclude': 2}},
        }
        layout_report = _document_parse_hook_layout_report({
            **inputs,
            'dry_run_report': {
                'summary': {'candidate_count': 2},
                'candidates': inputs['dry_run_report']['candidates'][:2],
            },
        })
        config_report = build_reviewed_filtering_internal_config_report(
            {
                'enabled': True,
                'mode': 'filtered_parse_experiment',
                'review_decisions': safe_decisions,
            },
            layout_report['header_footer_exclusion_dry_run'],
            safe_decisions)
        filtered_experiment_report = {
            'summary': {
                'baseline_body_text_block_count': 3,
                'filtered_body_text_block_count': 2,
            },
            'safety_warnings': [{
                'type': 'body_text_block_count_dropped',
                'baseline': 3,
                'filtered': 2,
            }],
        }

        pages = [_FakePage(0)]
        raw_pages = [_FakeRawPage(inputs['raw_object_pages'][0]['raw_objects'])]
        report = Pages._build_reviewed_filtering_internal_filtered_parse_report(
            layout_report,
            pages,
            raw_pages,
            config_report,
            filtered_parse_experiment_report=filtered_experiment_report,
            _reviewed_filtering_internal_config_enabled=True,
            _document_parse_filtering_review_decisions=safe_decisions)
        warning_types = {warning['type'] for warning in report['safety_warnings']}

        self.assertTrue(report['applied_to_parse'])
        self.assertEqual(report['summary']['body_text_block_delta'], -1)
        self.assertIn('body_text_block_count_changed', warning_types)
        self.assertIn('parse_metric_body_text_block_count_dropped', warning_types)
        self.assertTrue(report['summary']['body_text_signature_preserved'])

    def test_internal_filtered_docx_generation_removes_approved_synthetic_residuals(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-repeated.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                decisions,
                layout,
                Path(tmp))
            summary = report['summary']

            self.assertTrue(summary['baseline_docx_exists'])
            self.assertTrue(summary['filtered_docx_exists'])
            self.assertGreater(summary['baseline_docx_size'], 0)
            self.assertGreater(summary['filtered_docx_size'], 0)
            self.assertTrue(summary['internal_filtered_parse_applied'])
            self.assertGreater(summary['removed_approved_header_footer_page_number_count'], 0)
            self.assertEqual(summary['true_residual_header_footer_pollution_count'], 0)
            self.assertEqual(summary['body_text_loss_warning_count'], 0)
            self.assertEqual(summary['table_text_loss_warning_count'], 0)
            self.assertTrue(summary['body_text_signature_preserved'])
            self.assertLessEqual(
                summary['filtered_docx_paragraph_count'],
                summary['baseline_docx_paragraph_count'])
            self.assertTrue(summary['generated_docx_artifacts_temp_only'])
            self.assertTrue(report['recommendation']['safe_for_internal_filtered_docx'])

    def test_internal_filtered_docx_generation_preserves_body_table_like_text(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-body-table.pdf'
            _write_synthetic_pdf(pdf_path, 'body_table_near_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_FOOTER, ROLE_PAGE_NUMBER})

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                decisions,
                layout,
                Path(tmp))
            filtered_text = report['filtered_docx_metrics']['all_text']

            self.assertIn('body table header', filtered_text)
            self.assertIn('body table cell alpha', filtered_text)
            self.assertIn('body table cell beta', filtered_text)
            self.assertEqual(report['summary']['body_text_loss_warning_count'], 0)
            self.assertEqual(report['summary']['table_text_loss_warning_count'], 0)
            self.assertTrue(report['summary']['body_text_signature_preserved'])
            self.assertTrue(report['recommendation']['safe_for_internal_filtered_docx'])

    def test_internal_filtered_docx_preserves_callout_text_box_content(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-callout.pdf'
            _write_synthetic_pdf(pdf_path, 'callout_text_box_near_edges')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                decisions,
                layout,
                Path(tmp))
            filtered_text = report['filtered_docx_metrics']['all_text']

            self.assertIn('callout panel body content', filtered_text)
            self.assertIn('table like callout row alpha', filtered_text)
            self.assertEqual(_removed_text_match_count(
                _synthetic_filtering_report(layout, decisions),
                'callout panel body content'), 0)
            self.assertEqual(report['summary']['body_text_loss_warning_count'], 0)
            self.assertTrue(report['summary']['body_text_signature_preserved'])
            self.assertTrue(report['recommendation']['safe_for_internal_filtered_docx'])

    def test_internal_filtered_docx_preserves_list_items_and_headings(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-list-heading.pdf'
            _write_synthetic_pdf(pdf_path, 'list_heading_boundaries')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                decisions,
                layout,
                Path(tmp))
            filtered_text = report['filtered_docx_metrics']['all_text']

            self.assertIn('synthetic section heading 1 remains body text', filtered_text)
            self.assertIn('- list item alpha 1 remains in body', filtered_text)
            self.assertIn('1. numbered item beta 1 remains in body', filtered_text)
            self.assertEqual(report['summary']['body_text_loss_warning_count'], 0)
            self.assertTrue(report['summary']['body_text_signature_preserved'])
            self.assertTrue(report['recommendation']['safe_for_internal_filtered_docx'])

    def test_internal_filtered_docx_reports_table_geometry_stress_without_text_loss(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-table-geometry.pdf'
            _write_synthetic_pdf(pdf_path, 'table_geometry_delta_stress')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_FOOTER, ROLE_PAGE_NUMBER})

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                decisions,
                layout,
                Path(tmp))
            summary = report['summary']
            filtered_text = report['filtered_docx_metrics']['all_text']

            self.assertIn('geometry table header', filtered_text)
            self.assertIn('geometry table cell alpha', filtered_text)
            self.assertIn('geometry table cell beta', filtered_text)
            self.assertEqual(summary['table_text_loss_warning_count'], 0)
            self.assertTrue(summary['table_text_signature_preserved'])
            self.assertIn(
                summary['table_count_delta_classification'],
                {'unchanged', 'reported_no_table_text_loss'})
            self.assertTrue(report['recommendation']['safe_for_internal_filtered_docx'])

    def test_internal_filtered_docx_negative_control_preserves_body_content(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-negative.pdf'
            _write_synthetic_pdf(pdf_path, 'no_header_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = {'decisions': [], 'summary': {'candidate_count': 0}}

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                decisions,
                layout,
                Path(tmp))

            self.assertFalse(report['summary']['internal_filtered_parse_applied'])
            self.assertEqual(report['summary']['removed_approved_header_footer_page_number_count'], 0)
            self.assertEqual(report['summary']['body_text_loss_warning_count'], 0)
            self.assertTrue(report['summary']['body_text_signature_preserved'])

    def test_internal_filtered_docx_raw_would_exclude_without_approval_removes_nothing(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-no-approval.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')
            layout = _parse_synthetic_layout(pdf_path)

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                {'decisions': [], 'summary': {'candidate_count': 0}},
                layout,
                Path(tmp))
            filtered_text = report['filtered_docx_metrics']['all_text']

            self.assertFalse(report['summary']['internal_filtered_parse_applied'])
            self.assertEqual(report['summary']['removed_approved_header_footer_page_number_count'], 0)
            self.assertIn('synthetic report header', filtered_text)
            self.assertIn('synthetic report footer', filtered_text)

    def test_internal_filtered_docx_rejected_unsure_decisions_remain_blocked(self):
        _require_synthetic_docx_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-docx-rejected.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER})

            report = _run_synthetic_filtered_docx_comparison(
                pdf_path,
                decisions,
                layout,
                Path(tmp))
            warning_types = {
                warning['type']
                for warning in report['internal_filtered_parse_report']['safety_warnings']
            }

            self.assertFalse(report['summary']['internal_filtered_parse_applied'])
            self.assertIn('config_rejected_candidates_blocked', warning_types)
            self.assertEqual(report['summary']['removed_approved_header_footer_page_number_count'], 0)

    def test_synthetic_docx_comparison_fail_closes_on_body_text_loss(self):
        baseline = {
            'exists': True,
            'size': 100,
            'paragraph_count': 2,
            'table_count': 0,
            'all_text': 'body text remains synthetic report header',
        }
        filtered = {
            'exists': True,
            'size': 100,
            'paragraph_count': 1,
            'table_count': 0,
            'all_text': '',
        }
        report = _synthetic_docx_comparison_report_from_metrics(
            baseline,
            filtered,
            expected_body_texts=['body text remains'],
            removed_texts=['synthetic report header'],
            baseline_path='/tmp/baseline.docx',
            filtered_path='/tmp/filtered.docx',
            temp_root='/tmp',
            internal_filtered_parse_report={
                'applied_to_parse': True,
                'summary': {'body_text_block_delta': -1},
                'safety_warnings': [],
            },
            default_after_metrics=baseline)

        self.assertFalse(report['recommendation']['safe_for_internal_filtered_docx'])
        self.assertEqual(report['summary']['body_text_loss_warning_count'], 1)
        self.assertIn(
            'body_text_loss',
            {warning['type'] for warning in report['safety_warnings']})

    def test_synthetic_docx_comparison_reports_table_count_delta_without_text_loss(self):
        baseline = {
            'exists': True,
            'size': 100,
            'paragraph_count': 2,
            'table_count': 2,
            'table_texts': ['Geometry table cell alpha'],
            'all_text': 'body text remains geometry table cell alpha synthetic footer',
        }
        filtered = {
            'exists': True,
            'size': 100,
            'paragraph_count': 2,
            'table_count': 1,
            'table_texts': ['Geometry table cell alpha'],
            'all_text': 'body text remains geometry table cell alpha',
        }
        report = _synthetic_docx_comparison_report_from_metrics(
            baseline,
            filtered,
            expected_body_texts=['body text remains', 'geometry table cell alpha'],
            removed_texts=['synthetic footer'],
            baseline_path='/tmp/baseline.docx',
            filtered_path='/tmp/filtered.docx',
            temp_root='/tmp',
            internal_filtered_parse_report={
                'applied_to_parse': True,
                'summary': {'removed_raw_block_count': 1},
                'safety_warnings': [],
            },
            default_after_metrics=baseline)

        self.assertTrue(report['summary']['table_text_signature_preserved'])
        self.assertEqual(report['summary']['docx_table_count_delta'], -1)
        self.assertEqual(
            report['summary']['table_count_delta_classification'],
            'reported_no_table_text_loss')
        self.assertTrue(report['recommendation']['safe_for_internal_filtered_docx'])

    def test_synthetic_docx_comparison_fail_closes_on_table_text_loss(self):
        baseline = {
            'exists': True,
            'size': 100,
            'paragraph_count': 2,
            'table_count': 1,
            'table_texts': ['Geometry table cell alpha'],
            'all_text': 'body text remains geometry table cell alpha',
        }
        filtered = {
            'exists': True,
            'size': 100,
            'paragraph_count': 2,
            'table_count': 1,
            'table_texts': [],
            'all_text': 'body text remains',
        }
        report = _synthetic_docx_comparison_report_from_metrics(
            baseline,
            filtered,
            expected_body_texts=['body text remains', 'geometry table cell alpha'],
            removed_texts=[],
            baseline_path='/tmp/baseline.docx',
            filtered_path='/tmp/filtered.docx',
            temp_root='/tmp',
            internal_filtered_parse_report={
                'applied_to_parse': True,
                'summary': {'removed_raw_block_count': 1},
                'safety_warnings': [],
            },
            default_after_metrics=baseline)

        self.assertFalse(report['summary']['table_text_signature_preserved'])
        self.assertEqual(report['summary']['table_text_loss_warning_count'], 1)
        self.assertFalse(report['recommendation']['safe_for_internal_filtered_docx'])
        self.assertIn(
            'table_text_loss',
            {warning['type'] for warning in report['safety_warnings']})

    def test_docx_header_footer_generation_plan_is_serializable(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'docx-header-footer-plan.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})

            plan = build_docx_header_footer_generation_plan(
                layout.get('pages', []),
                layout.get('header_footer_exclusion_dry_run', {}),
                decisions,
                enabled=True)
            decoded = json.loads(json.dumps(plan))

            self.assertTrue(decoded['enabled'])
            self.assertEqual(decoded['summary']['header_text_count'], 1)
            self.assertEqual(decoded['summary']['footer_text_count'], 1)
            self.assertEqual(decoded['summary']['page_number_placeholder_count'], 1)
            self.assertEqual(decoded['summary']['page_number_field_generation'], 'deferred_placeholder_only')
            self.assertFalse(decoded['summary']['public_cli_exposed'])
            self.assertFalse(decoded['summary']['production_default_enabled'])

    def test_docx_header_footer_generation_plan_uses_only_explicit_approvals(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'docx-header-footer-plan-review-gated.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER})

            plan = build_docx_header_footer_generation_plan(
                layout.get('pages', []),
                layout.get('header_footer_exclusion_dry_run', {}),
                decisions,
                enabled=True)
            section = plan['sections'][0]

            self.assertEqual(plan['summary']['header_text_count'], 1)
            self.assertEqual(plan['summary']['footer_text_count'], 0)
            self.assertEqual(plan['summary']['page_number_placeholder_count'], 0)
            self.assertTrue(section['header_texts'])
            self.assertFalse(section['footer_texts'])
            self.assertFalse(section['page_number_placeholders'])
            self.assertTrue(all(
                candidate['manual_decision'] != 'approve_exclude'
                for candidate in plan['blocked_candidates']))

    def test_docx_header_footer_generation_plan_blocks_body_and_layout_placeholders(self):
        body_fingerprint = _summary_fingerprint('Approved Body Heading', REGION_BODY)
        placeholder_fingerprint = _summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP)
        page_summaries = [{
            'page_index': 0,
            'text_blocks': [
                _mapping_summary_block(0, body_fingerprint, REGION_BODY, 'Approved Body Heading'),
                _mapping_summary_block(1, placeholder_fingerprint, REGION_TOP, IMAGE_PLACEHOLDER),
            ],
        }]
        dry_run = {
            'candidates': [
                _dry_run_candidate(
                    'body-candidate',
                    body_fingerprint,
                    ROLE_HEADER,
                    ACTION_WOULD_EXCLUDE,
                    [0],
                    [REGION_BODY]),
                _dry_run_candidate(
                    'placeholder-candidate',
                    placeholder_fingerprint,
                    ROLE_LAYOUT_PLACEHOLDER,
                    ACTION_WOULD_EXCLUDE,
                    [0],
                    [REGION_TOP]),
            ],
        }
        decisions = {
            'decisions': [
                _review_decision('body-candidate', body_fingerprint, 'approve_exclude'),
                _review_decision('placeholder-candidate', placeholder_fingerprint, 'approve_exclude'),
            ],
        }

        plan = build_docx_header_footer_generation_plan(
            page_summaries,
            dry_run,
            decisions,
            enabled=True)
        warning_types = {warning['type'] for warning in plan['safety_warnings']}

        self.assertEqual(plan['summary']['representable_entry_count'], 0)
        self.assertEqual(plan['summary']['unrepresentable_approved_candidate_count'], 1)
        self.assertEqual(plan['summary']['blocked_candidate_count'], 1)
        self.assertIn('body_region_candidate_not_represented', warning_types)
        self.assertEqual(
            plan['blocked_candidates'][0]['reason'],
            'layout_placeholder_not_filterable')
        self.assertFalse(
            plan['recommendation']['safe_for_internal_docx_header_footer_experiment'])

    def test_internal_docx_header_footer_text_plan_writes_temp_docx_parts(self):
        _require_docx_header_footer_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'docx-header-footer-write.pdf'
            docx_path = Path(tmp) / 'header-footer.docx'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})
            plan = build_docx_header_footer_generation_plan(
                layout.get('pages', []),
                layout.get('header_footer_exclusion_dry_run', {}),
                decisions,
                enabled=True)

            document = DocxDocument()
            report = docx_utils.apply_header_footer_text_plan(
                document,
                plan,
                enabled=True)
            document.save(str(docx_path))
            with zipfile.ZipFile(docx_path) as archive:
                names = set(archive.namelist())
                header_xml = archive.read('word/header1.xml').decode('utf-8')
                footer_xml = archive.read('word/footer1.xml').decode('utf-8')

            self.assertTrue(report['applied'])
            self.assertIn('word/header1.xml', names)
            self.assertIn('word/footer1.xml', names)
            self.assertIn('SYNTHETIC REPORT HEADER', header_xml)
            self.assertIn('SYNTHETIC REPORT FOOTER', footer_xml)
            self.assertIn('&lt;PAGE_NUMBER&gt;', footer_xml)
            self.assertEqual(
                report['summary']['page_number_field_generation'],
                'deferred_placeholder_only')

    def test_internal_docx_header_footer_text_plan_disabled_by_default(self):
        _require_docx_header_footer_support(self)
        document = DocxDocument()
        report = docx_utils.apply_header_footer_text_plan(
            document,
            {'sections': [{'header_texts': ['Internal Header']}]},
            enabled=False)

        self.assertFalse(report['enabled'])
        self.assertFalse(report['applied'])
        self.assertEqual(report['summary']['header_paragraphs_written'], 0)

    def test_internal_filtered_body_docx_header_footer_output_repeated_fixture(self):
        _require_docx_header_footer_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'filtered-body-with-header-footer.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})

            report = _run_synthetic_filtered_docx_with_header_footer_parts(
                pdf_path,
                decisions,
                layout,
                tmp)

            self.assertTrue(report['plan']['enabled'])
            self.assertTrue(report['apply_report']['applied'])
            self.assertTrue(report['comparison']['recommendation']['safe_for_internal_filtered_docx'])
            self.assertEqual(report['plan']['summary']['header_text_count'], 1)
            self.assertEqual(report['plan']['summary']['footer_text_count'], 1)
            self.assertEqual(report['plan']['summary']['page_number_placeholder_count'], 1)
            self.assertIn('SYNTHETIC REPORT HEADER', report['openxml']['header_xml'])
            self.assertIn('SYNTHETIC REPORT FOOTER', report['openxml']['footer_xml'])
            self.assertIn('&lt;PAGE_NUMBER&gt;', report['openxml']['footer_xml'])
            self.assertNotIn('SYNTHETIC REPORT HEADER', report['openxml']['body_xml'])
            self.assertNotIn('SYNTHETIC REPORT FOOTER', report['openxml']['body_xml'])
            self.assertNotIn('Page 1 of 4', report['openxml']['body_xml'])
            self.assertEqual(
                report['comparison']['summary']['true_residual_header_footer_pollution_count'],
                0)
            self.assertTrue(
                report['comparison']['summary']['body_text_signature_preserved'])
            self.assertTrue(
                report['comparison']['summary']['generated_docx_artifacts_temp_only'])
            self.assertEqual(
                report['apply_report']['summary']['page_number_field_generation'],
                'deferred_placeholder_only')
            self.assertFalse(report['default_metrics']['header_footer_xml']['has_header_text'])

    def test_internal_docx_header_footer_output_preserves_body_callout_content(self):
        _require_docx_header_footer_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'callout-preserved-with-header-footer.pdf'
            _write_synthetic_pdf(pdf_path, 'callout_text_box_near_edges')
            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})

            report = _run_synthetic_filtered_docx_with_header_footer_parts(
                pdf_path,
                decisions,
                layout,
                tmp)
            body_xml = report['openxml']['body_xml']
            header_footer_xml = (
                report['openxml']['header_xml'] +
                report['openxml']['footer_xml'])

            self.assertIn('Callout panel body content', body_xml)
            self.assertIn('Table like callout row alpha', body_xml)
            self.assertNotIn('Callout panel body content', header_footer_xml)
            self.assertNotIn('Table like callout row alpha', header_footer_xml)
            self.assertEqual(_removed_region_count(report['filtering_report'], REGION_BODY), 0)
            self.assertTrue(
                report['comparison']['summary']['body_text_signature_preserved'])
            self.assertEqual(
                report['comparison']['summary']['body_text_loss_warning_count'],
                0)

    def test_internal_docx_header_footer_output_excludes_blocked_candidates(self):
        _require_docx_header_footer_support(self)
        page_summaries = [{
            'page_index': 0,
            'text_blocks': [
                _mapping_summary_block(
                    0,
                    _summary_fingerprint('Approved Header', REGION_TOP),
                    REGION_TOP,
                    'Approved Header'),
                _mapping_summary_block(
                    1,
                    _summary_fingerprint('Rejected Footer', REGION_BOTTOM),
                    REGION_BOTTOM,
                    'Rejected Footer'),
                _mapping_summary_block(
                    2,
                    _summary_fingerprint('Unsure Footer', REGION_BOTTOM),
                    REGION_BOTTOM,
                    'Unsure Footer'),
                _mapping_summary_block(
                    3,
                    _summary_fingerprint('Body Heading', REGION_BODY),
                    REGION_BODY,
                    'Body Heading'),
                _mapping_summary_block(
                    4,
                    _summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP),
                    REGION_TOP,
                    IMAGE_PLACEHOLDER),
            ],
        }]
        dry_run = {
            'candidates': [
                _dry_run_candidate(
                    'approved-header',
                    _summary_fingerprint('Approved Header', REGION_TOP),
                    ROLE_HEADER,
                    ACTION_WOULD_EXCLUDE,
                    [0],
                    [REGION_TOP]),
                _dry_run_candidate(
                    'rejected-footer',
                    _summary_fingerprint('Rejected Footer', REGION_BOTTOM),
                    ROLE_FOOTER,
                    ACTION_WOULD_EXCLUDE,
                    [0],
                    [REGION_BOTTOM]),
                _dry_run_candidate(
                    'unsure-footer',
                    _summary_fingerprint('Unsure Footer', REGION_BOTTOM),
                    ROLE_FOOTER,
                    ACTION_WOULD_EXCLUDE,
                    [0],
                    [REGION_BOTTOM]),
                _dry_run_candidate(
                    'body-heading',
                    _summary_fingerprint('Body Heading', REGION_BODY),
                    ROLE_HEADER,
                    ACTION_WOULD_EXCLUDE,
                    [0],
                    [REGION_BODY]),
                _dry_run_candidate(
                    'placeholder',
                    _summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP),
                    ROLE_LAYOUT_PLACEHOLDER,
                    ACTION_WOULD_EXCLUDE,
                    [0],
                    [REGION_TOP]),
            ],
        }
        decisions = {
            'decisions': [
                _review_decision(
                    'approved-header',
                    _summary_fingerprint('Approved Header', REGION_TOP),
                    'approve_exclude'),
                _review_decision(
                    'rejected-footer',
                    _summary_fingerprint('Rejected Footer', REGION_BOTTOM),
                    'reject_exclude'),
                _review_decision(
                    'unsure-footer',
                    _summary_fingerprint('Unsure Footer', REGION_BOTTOM),
                    'unsure'),
                _review_decision(
                    'body-heading',
                    _summary_fingerprint('Body Heading', REGION_BODY),
                    'approve_exclude'),
                _review_decision(
                    'placeholder',
                    _summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP),
                    'approve_exclude'),
            ],
        }
        plan = build_docx_header_footer_generation_plan(
            page_summaries,
            dry_run,
            decisions,
            enabled=True)

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / 'blocked-candidates.docx'
            document = DocxDocument()
            apply_report = docx_utils.apply_header_footer_text_plan(
                document,
                plan,
                enabled=True)
            document.save(str(docx_path))
            openxml = _read_docx_openxml_parts(docx_path)
            header_footer_xml = openxml['header_xml'] + openxml['footer_xml']

        self.assertFalse(apply_report['applied'])
        self.assertEqual(apply_report['summary']['header_paragraphs_written'], 0)
        self.assertNotIn('Approved Header', header_footer_xml)
        self.assertNotIn('Rejected Footer', header_footer_xml)
        self.assertNotIn('Unsure Footer', header_footer_xml)
        self.assertNotIn('Body Heading', header_footer_xml)
        self.assertNotIn(IMAGE_PLACEHOLDER, header_footer_xml)
        self.assertIn(
            'header_footer_plan_not_safe_to_apply',
            {warning['type'] for warning in apply_report['safety_warnings']})
        self.assertFalse(
            plan['recommendation']['safe_for_internal_docx_header_footer_experiment'])
        self.assertIn(
            'body_region_candidate_not_represented',
            {warning['type'] for warning in plan['safety_warnings']})

    def test_local_corpus_smoke_summary_handles_multiple_samples(self):
        reports = [
            _local_corpus_docx_smoke_sample_report(
                'input.pdf',
                _local_docx_metric(paragraphs=20, tables=2),
                _local_docx_metric(paragraphs=16, tables=2),
                _local_internal_report(removed=4),
                body_text_signature_preserved=True,
                generated_paths=['local_reports/phase3d/input/baseline.docx']),
            _local_corpus_docx_smoke_sample_report(
                'input3.pdf',
                _local_docx_metric(paragraphs=18, tables=1),
                _local_docx_metric(paragraphs=17, tables=1),
                _local_internal_report(removed=1),
                body_text_signature_preserved=True,
                generated_paths=['local_reports/phase3d/input3/filtered.docx']),
        ]
        summary = _local_corpus_docx_smoke_summary_report(reports)

        self.assertEqual(summary['summary']['sample_count'], 2)
        self.assertEqual(summary['summary']['passed_count'], 2)
        self.assertEqual(summary['summary']['blocked_count'], 0)
        self.assertEqual(summary['summary']['total_removed_approved_count'], 5)

    def test_local_corpus_smoke_marks_missing_approval_artifacts_blocked(self):
        report = _local_corpus_docx_smoke_sample_report(
            'missing.pdf',
            _local_docx_metric(),
            _local_docx_metric(),
            {},
            approval_artifacts_available=False)

        self.assertEqual(report['status'], 'blocked')
        self.assertIn(
            'missing_approval_artifacts',
            {warning['type'] for warning in report['safety_warnings']})

    def test_local_corpus_smoke_keeps_large_sample_bounded_only(self):
        report = _local_corpus_docx_smoke_sample_report(
            'input6_large.pdf',
            _local_docx_metric(),
            _local_docx_metric(),
            _local_internal_report(removed=30),
            bounded_subset_only=True,
            full_document_skipped=True,
            body_text_signature_preserved=True,
            generated_paths=['local_reports/phase3d/input6_large_subset/filtered.docx'])

        self.assertEqual(report['status'], 'bounded_subset_passed')
        self.assertTrue(report['bounded_subset_only'])
        self.assertTrue(report['full_document_skipped'])

    def test_local_corpus_smoke_requires_body_text_signature_preservation(self):
        report = _local_corpus_docx_smoke_sample_report(
            'input3.pdf',
            _local_docx_metric(),
            _local_docx_metric(),
            _local_internal_report(removed=1),
            body_text_signature_preserved=False)

        self.assertEqual(report['status'], 'blocked')
        self.assertIn(
            'body_text_signature_not_preserved',
            {warning['type'] for warning in report['safety_warnings']})

    def test_local_corpus_smoke_fail_closes_on_residual_body_or_table_loss(self):
        report = _local_corpus_docx_smoke_sample_report(
            'input3.pdf',
            _local_docx_metric(),
            _local_docx_metric(),
            _local_internal_report(removed=1),
            body_text_signature_preserved=True,
            true_residual_header_footer_pollution_count=1,
            body_text_loss_warning_count=1,
            table_text_loss_warning_count=1)
        warning_types = {warning['type'] for warning in report['safety_warnings']}

        self.assertEqual(report['status'], 'blocked')
        self.assertIn('true_residual_header_footer_pollution', warning_types)
        self.assertIn('body_text_loss', warning_types)
        self.assertIn('table_text_loss', warning_types)

    def test_local_corpus_smoke_reports_body_textblock_delta(self):
        report = _local_corpus_docx_smoke_sample_report(
            'input3.pdf',
            _local_docx_metric(),
            _local_docx_metric(),
            _local_internal_report(removed=1, body_text_block_delta=-2),
            body_text_signature_preserved=True)

        self.assertEqual(report['summary']['body_text_block_delta'], -2)
        self.assertIn(
            'body_text_block_count_changed',
            {warning['type'] for warning in report['diagnostic_warnings']})
        self.assertEqual(report['status'], 'passed')

    def test_local_corpus_smoke_blocks_nonlocal_docx_paths(self):
        report = _local_corpus_docx_smoke_sample_report(
            'input.pdf',
            _local_docx_metric(),
            _local_docx_metric(),
            _local_internal_report(removed=1),
            body_text_signature_preserved=True,
            generated_paths=['/outside/generated.docx'])

        self.assertEqual(report['status'], 'blocked')
        self.assertIn(
            'generated_docx_path_not_local_only',
            {warning['type'] for warning in report['safety_warnings']})

    def test_synthetic_repeated_header_footer_fixture_supports_reviewed_filtering(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'repeated-header-footer.pdf'
            _write_synthetic_pdf(pdf_path, 'repeated_header_footer')

            layout = _parse_synthetic_layout(pdf_path)
            dry_run = layout['header_footer_exclusion_dry_run']
            roles = {candidate['proposed_role'] for candidate in dry_run['candidates']}

            self.assertIn(ROLE_HEADER, roles)
            self.assertIn(ROLE_FOOTER, roles)
            self.assertIn(ROLE_PAGE_NUMBER, roles)

            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})
            filtering = _synthetic_filtering_report(layout, decisions)
            body_validation = _synthetic_body_text_validation(layout, filtering)
            diagnostics = _run_synthetic_document_parse_diagnostics(
                pdf_path,
                decisions,
                expected_remove_count=filtering['summary']['removed_block_count'])

            self.assertEqual(filtering['approved_candidate_count'], 3)
            self.assertEqual(filtering['summary']['removed_block_count'], 12)
            self.assertEqual(_removed_region_count(filtering, REGION_BODY), 0)
            self.assertTrue(body_validation['body_text_signature_preserved'])
            self.assertEqual(
                diagnostics['mapping']['summary']['exact_match_count'],
                filtering['summary']['removed_block_count'])
            self.assertTrue(
                diagnostics['guarded']['summary']['restore_fingerprint_match'])
            self.assertFalse(
                diagnostics['guarded']['summary']['original_raw_pages_left_mutated'])

    def test_synthetic_body_table_near_footer_preserves_body_content(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'body-table-near-footer.pdf'
            _write_synthetic_pdf(pdf_path, 'body_table_near_footer')

            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_FOOTER, ROLE_PAGE_NUMBER})
            filtering = _synthetic_filtering_report(layout, decisions)
            body_validation = _synthetic_body_text_validation(layout, filtering)

            self.assertGreater(filtering['summary']['removed_block_count'], 0)
            self.assertEqual(_removed_region_count(filtering, REGION_BODY), 0)
            self.assertEqual(_removed_text_match_count(filtering, 'Body Table'), 0)
            self.assertIn('body table cell alpha', body_validation['filtered_signature'])
            self.assertTrue(body_validation['body_text_signature_preserved'])

    def test_synthetic_no_header_footer_negative_control_requires_manual_approval(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            negative_pdf = Path(tmp) / 'no-header-footer.pdf'
            repeated_pdf = Path(tmp) / 'raw-would-exclude-needs-review.pdf'
            _write_synthetic_pdf(negative_pdf, 'no_header_footer')
            _write_synthetic_pdf(repeated_pdf, 'repeated_header_footer')

            negative_layout = _parse_synthetic_layout(negative_pdf)
            repeated_layout = _parse_synthetic_layout(repeated_pdf)
            no_decision_filter = _synthetic_filtering_report(
                repeated_layout,
                {'decisions': [], 'summary': {'candidate_count': 0}})

            self.assertEqual(
                negative_layout['header_footer_exclusion_dry_run']['summary']['candidate_count'],
                0)
            self.assertGreater(
                repeated_layout['header_footer_exclusion_dry_run']['summary']['candidate_count'],
                0)
            self.assertEqual(no_decision_filter['approved_candidate_count'], 0)
            self.assertEqual(no_decision_filter['summary']['removed_block_count'], 0)

    def test_synthetic_first_page_odd_even_headers_remain_review_gated(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'first-page-odd-even.pdf'
            _write_synthetic_pdf(pdf_path, 'first_page_odd_even_headers')

            layout = _parse_synthetic_layout(pdf_path)
            top_candidates = [
                candidate
                for candidate in layout['header_footer_exclusion_dry_run']['candidates']
                if candidate.get('region') == REGION_TOP
            ]
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_REVIEW_ONLY})
            filtering = _synthetic_filtering_report(layout, decisions)

            self.assertTrue(top_candidates)
            self.assertTrue(all(
                candidate['action'] == ACTION_REVIEW
                for candidate in top_candidates))
            self.assertTrue(all(
                candidate['support_count'] < candidate['page_count']
                for candidate in top_candidates))
            self.assertEqual(filtering['approved_candidate_count'], 0)
            self.assertEqual(filtering['summary']['removed_block_count'], 0)

    def test_synthetic_odd_even_body_heading_interaction_remains_review_gated(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'odd-even-body-heading.pdf'
            _write_synthetic_pdf(pdf_path, 'odd_even_body_heading_interaction')

            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})
            filtering = _synthetic_filtering_report(layout, decisions)
            body_validation = _synthetic_body_text_validation(layout, filtering)

            self.assertEqual(_removed_region_count(filtering, REGION_BODY), 0)
            self.assertEqual(
                _removed_text_match_count(filtering, 'Synthetic odd header body heading'),
                0)
            self.assertTrue(body_validation['body_text_signature_preserved'])

    def test_synthetic_paragraph_continuity_preserves_body_text_signature(self):
        _require_synthetic_pdf_support(self)
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / 'paragraph-continuity.pdf'
            _write_synthetic_pdf(pdf_path, 'paragraph_continuity')

            layout = _parse_synthetic_layout(pdf_path)
            decisions = _synthetic_review_decisions(
                layout,
                approve_roles={ROLE_HEADER, ROLE_FOOTER, ROLE_PAGE_NUMBER})
            filtering = _synthetic_filtering_report(layout, decisions)
            body_validation = _synthetic_body_text_validation(
                layout,
                filtering,
                baseline_body_text_block_count=3,
                filtered_body_text_block_count=2)

            continuation_labels = {
                candidate.get('label')
                for candidate in layout.get('paragraph_continuation_candidates', []) or []
            }

            self.assertIn('candidate', continuation_labels)
            self.assertEqual(_removed_region_count(filtering, REGION_BODY), 0)
            self.assertTrue(body_validation['body_text_signature_preserved'])
            self.assertEqual(
                body_validation['body_text_block_delta_classification'],
                'acceptable_boundary_or_grouping_shift')


def _page(page_index, blocks):
    return {
        'page_index': page_index,
        'width': 600,
        'height': 1000,
        'blocks': blocks,
    }


def _block(text, x0, y0, x1, y1, style=None):
    return {
        'text': text,
        'bbox': [x0, y0, x1, y1],
        'style': style or {'font': 'Times New Roman', 'size': 11.0},
    }


def _production_page(page_index, text_blocks):
    return {
        'id': page_index,
        'width': 600,
        'height': 1000,
        'sections': [
            {
                'columns': [
                    {
                        'blocks': text_blocks,
                    },
                ],
            },
        ],
    }


def _production_text_block(lines, bbox):
    return {
        'type': 0,
        'bbox': bbox,
        'lines': [
            {
                'spans': [
                    {
                        'text': line,
                    },
                ],
            }
            for line in lines
        ],
    }


def _estimator_report_for_mismatch(pages):
    group_count = sum(page['estimated_paragraph_group_count'] for page in pages)
    body_blocks = sum(page['body_block_count_after_filtering'] for page in pages)
    return {
        'enabled': True,
        'summary': {
            'estimated_paragraph_group_count': group_count,
            'body_block_count_after_filtering': body_blocks,
            'average_blocks_per_estimated_paragraph': (
                round(body_blocks / group_count, 3)
                if group_count else 0.0),
            'suspicious_single_line_paragraph_count': 0,
            'suspicious_short_fragment_count': 0,
            'one_line_group_ratio': 0.0,
            'short_fragment_ratio': 0.0,
        },
        'pages': pages,
        'diagnostics': {},
    }


def _estimator_page_for_mismatch(page_index, group_count, body_blocks, split_reasons):
    groups = [
        {
            'group_index': index,
            'line_count': 1,
            'block_count': 1,
            'break_before_reasons': split_reasons[index-1:index],
            'text_preview': f'Estimator group {index}',
        }
        for index in range(group_count)
    ]
    return {
        'page_index': page_index,
        'page_number': page_index + 1,
        'estimated_paragraph_group_count': group_count,
        'body_block_count_after_filtering': body_blocks,
        'average_blocks_per_estimated_paragraph': (
            round(body_blocks / group_count, 3)
            if group_count else 0.0),
        'suspicious_single_line_paragraph_count': 0,
        'suspicious_short_fragment_count': 0,
        'estimated_paragraph_groups': groups,
        'split_boundaries': [
            {
                'boundary_index': index,
                'reasons': [reason],
                'previous_text_preview': f'Previous {index}',
                'next_text_preview': f'Next {index}',
            }
            for index, reason in enumerate(split_reasons)
        ],
    }


def _estimator_report_for_indentation(boundaries):
    pages = {}
    for boundary in boundaries:
        page_index = boundary['page_index']
        pages.setdefault(page_index, {
            'page_index': page_index,
            'page_number': page_index + 1,
            'estimated_paragraph_group_count': 1,
            'body_block_count_after_filtering': 2,
            'average_blocks_per_estimated_paragraph': 2.0,
            'suspicious_single_line_paragraph_count': 0,
            'suspicious_short_fragment_count': 0,
            'estimated_paragraph_groups': [],
            'split_boundaries': [],
        })
        pages[page_index]['split_boundaries'].append(boundary)

    return {
        'enabled': True,
        'summary': {
            'estimated_paragraph_group_count': len(pages),
            'body_block_count_after_filtering': len(boundaries) * 2,
            'average_blocks_per_estimated_paragraph': 2.0,
            'suspicious_single_line_paragraph_count': 0,
            'suspicious_short_fragment_count': 0,
            'one_line_group_ratio': 0.0,
            'short_fragment_ratio': 0.0,
        },
        'pages': list(pages.values()),
        'diagnostics': {},
    }


def _indentation_boundary(
        page_index,
        left_delta=24.0,
        previous_sentence_end=False,
        width_similar=True,
        previous_width_ratio=0.95,
        previous_right_gap_ratio=0.02,
        extra_reasons=None,
        current_heading_like=False,
        current_list_marker=False,
        insufficient_metadata=False):
    reasons = ['indentation_change']
    reasons.extend(extra_reasons or [])
    return {
        'page_index': page_index,
        'page_number': page_index + 1,
        'boundary_index': 0,
        'previous_line_index': 0,
        'next_line_index': 1,
        'previous_block_indexes': [0],
        'next_block_indexes': [1],
        'previous_bbox': [50.0, 300.0, 520.0, 320.0],
        'next_bbox': [50.0 + left_delta, 326.0, 520.0, 346.0],
        'previous_left': 50.0,
        'previous_right': 520.0,
        'next_left': 50.0 + left_delta,
        'next_right': 520.0,
        'reasons': reasons,
        'signals': {
            'left_delta': left_delta,
            'right_delta': 0.0,
            'width_delta_ratio': 0.02 if width_similar else 0.4,
            'width_similar': width_similar,
            'previous_sentence_end': previous_sentence_end,
            'previous_hyphenated': False,
            'current_heading_like': current_heading_like,
            'current_list_marker': current_list_marker,
            'style_change': False,
            'significant_style_change': False,
            'previous_width_ratio': previous_width_ratio,
            'previous_right_gap_ratio': previous_right_gap_ratio,
            'gap_ratio': 0.3,
            'insufficient_metadata': insufficient_metadata,
        },
        'previous_text_preview': 'Previous visual line',
        'next_text_preview': 'Next visual line',
    }


def _filter_insertion_reports(body_region_removed_count=0):
    return {
        'layout_analysis_report': {
            'page_count': 2,
            'pages': [{}, {}],
            'header_footer_exclusion_dry_run': {
                'summary': {'candidate_count': 5},
                'candidates': [{} for _ in range(5)],
            },
        },
        'review_decisions': {
            'summary': {
                'approve_exclude': 2,
                'reject_exclude': 1,
                'unsure': 1,
            },
            'decisions': [],
        },
        'body_filtering_diff_report': {
            'summary': {
                'approved_candidate_count': 2,
                'blocked_candidate_count': 3,
                'would_remove_block_count': 8,
            },
        },
        'paragraph_integrity_report': {
            'summary': {
                'body_region_removed_count': body_region_removed_count,
                'suspicious_warning_count': 0,
            },
        },
        'paragraph_grouping_report': {
            'summary': {
                'estimated_paragraph_group_count': 79,
            },
        },
        'production_comparison_report': {
            'estimator': {
                'paragraph_group_count': 79,
            },
            'production_observed': {
                'available': True,
                'paragraph_group_count': 52,
            },
            'mismatch': {
                'absolute_group_count_delta': 27,
                'estimator_to_production_group_ratio': 1.519,
                'group_count_delta_ratio': 0.519,
            },
            'warnings': [],
        },
        'paragraph_mismatch_report': {
            'summary': {
                'dominant_mismatch_cause': 'estimator_over_split_by_indentation',
            },
            'warnings': [],
        },
        'indentation_rule_report': {
            'summary': {
                'total_indentation_split_boundaries': 22,
                'estimator_should_merge_count': 0,
                'estimator_should_split_count': 22,
            },
        },
    }


def _document_parse_simulation_inputs():
    page_summaries = [
        {
            'page_index': 0,
            'page_number': 1,
            'text_block_count': 6,
            'region_counts': {
                REGION_TOP: 3,
                REGION_BODY: 1,
                REGION_BOTTOM: 2,
            },
            'text_blocks': [
                _summary_block(0, 'approved-header', REGION_TOP, 'Approved Header'),
                _summary_block(1, 'body-text', REGION_BODY, 'Body paragraph'),
                _summary_block(2, 'approved-footer', REGION_BOTTOM, 'Approved Footer'),
                _summary_block(3, 'reject-header', REGION_TOP, 'Rejected Header'),
                _summary_block(4, 'unsure-footer', REGION_BOTTOM, 'Unsure Footer'),
                _summary_block(5, 'image-placeholder', REGION_TOP, IMAGE_PLACEHOLDER),
            ],
        },
    ]
    dry_run_report = {
        'summary': {'candidate_count': 5},
        'candidates': [
            _dry_run_candidate('c-approved-header', 'approved-header', ROLE_HEADER, ACTION_WOULD_EXCLUDE, [0], [REGION_TOP]),
            _dry_run_candidate('c-approved-footer', 'approved-footer', ROLE_FOOTER, ACTION_WOULD_EXCLUDE, [0], [REGION_BOTTOM]),
            _dry_run_candidate('c-reject-header', 'reject-header', ROLE_HEADER, ACTION_WOULD_EXCLUDE, [0], [REGION_TOP]),
            _dry_run_candidate('c-unsure-footer', 'unsure-footer', ROLE_FOOTER, ACTION_WOULD_EXCLUDE, [0], [REGION_BOTTOM]),
            _dry_run_candidate('c-image-placeholder', 'image-placeholder', ROLE_LAYOUT_PLACEHOLDER, ACTION_WOULD_EXCLUDE, [0], [REGION_TOP]),
        ],
    }
    review_decisions = {
        'decisions': [
            _review_decision('c-approved-header', 'approved-header', 'approve_exclude'),
            _review_decision('c-approved-footer', 'approved-footer', 'approve_exclude'),
            _review_decision('c-reject-header', 'reject-header', 'reject_exclude'),
            _review_decision('c-unsure-footer', 'unsure-footer', 'unsure'),
            _review_decision('c-image-placeholder', 'image-placeholder', 'approve_exclude'),
        ],
        'summary': {
            'decision_counts': {
                'approve_exclude': 3,
                'reject_exclude': 1,
                'unsure': 1,
            },
        },
    }
    return {
        'page_summaries': page_summaries,
        'dry_run_report': dry_run_report,
        'review_decisions': review_decisions,
        'body_filtering_diff_report': {
            'summary': {
                'would_remove_block_count': 2,
                'kept_block_count': 4,
            },
        },
        'paragraph_integrity_report': {
            'summary': {
                'body_region_removed_count': 0,
            },
        },
    }


def _document_parse_hook_layout_report(inputs):
    return {
        'pages': inputs['page_summaries'],
        'header_footer_exclusion_dry_run': inputs['dry_run_report'],
    }


def _document_parse_raw_mapping_inputs():
    page_summaries = [
        {
            'page_index': 0,
            'page_number': 1,
            'text_block_count': 6,
            'region_counts': {
                REGION_TOP: 3,
                REGION_BODY: 1,
                REGION_BOTTOM: 2,
            },
            'text_blocks': [
                _mapping_summary_block(
                    0,
                    _summary_fingerprint('Approved Header', REGION_TOP),
                    REGION_TOP,
                    'Approved Header'),
                _mapping_summary_block(
                    1,
                    _summary_fingerprint('Body paragraph', REGION_BODY),
                    REGION_BODY,
                    'Body paragraph'),
                _mapping_summary_block(
                    2,
                    _summary_fingerprint('Approved Footer', REGION_BOTTOM),
                    REGION_BOTTOM,
                    'Approved Footer'),
                _mapping_summary_block(
                    3,
                    _summary_fingerprint('Rejected Header', REGION_TOP),
                    REGION_TOP,
                    'Rejected Header'),
                _mapping_summary_block(
                    4,
                    _summary_fingerprint('Unsure Footer', REGION_BOTTOM),
                    REGION_BOTTOM,
                    'Unsure Footer'),
                _mapping_summary_block(
                    5,
                    _summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP),
                    REGION_TOP,
                    IMAGE_PLACEHOLDER),
            ],
        },
    ]
    dry_run_report = {
        'summary': {'candidate_count': 5},
        'candidates': [
            _dry_run_candidate(
                'c-approved-header',
                _summary_fingerprint('Approved Header', REGION_TOP),
                ROLE_HEADER,
                ACTION_WOULD_EXCLUDE,
                [0],
                [REGION_TOP]),
            _dry_run_candidate(
                'c-approved-footer',
                _summary_fingerprint('Approved Footer', REGION_BOTTOM),
                ROLE_FOOTER,
                ACTION_WOULD_EXCLUDE,
                [0],
                [REGION_BOTTOM]),
            _dry_run_candidate(
                'c-reject-header',
                _summary_fingerprint('Rejected Header', REGION_TOP),
                ROLE_HEADER,
                ACTION_WOULD_EXCLUDE,
                [0],
                [REGION_TOP]),
            _dry_run_candidate(
                'c-unsure-footer',
                _summary_fingerprint('Unsure Footer', REGION_BOTTOM),
                ROLE_FOOTER,
                ACTION_WOULD_EXCLUDE,
                [0],
                [REGION_BOTTOM]),
            _dry_run_candidate(
                'c-image-placeholder',
                _summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP),
                ROLE_LAYOUT_PLACEHOLDER,
                ACTION_WOULD_EXCLUDE,
                [0],
                [REGION_TOP]),
        ],
    }
    review_decisions = {
        'decisions': [
            _review_decision(
                'c-approved-header',
                _summary_fingerprint('Approved Header', REGION_TOP),
                'approve_exclude'),
            _review_decision(
                'c-approved-footer',
                _summary_fingerprint('Approved Footer', REGION_BOTTOM),
                'approve_exclude'),
            _review_decision(
                'c-reject-header',
                _summary_fingerprint('Rejected Header', REGION_TOP),
                'reject_exclude'),
            _review_decision(
                'c-unsure-footer',
                _summary_fingerprint('Unsure Footer', REGION_BOTTOM),
                'unsure'),
            _review_decision(
                'c-image-placeholder',
                _summary_fingerprint(IMAGE_PLACEHOLDER, REGION_TOP),
                'approve_exclude'),
        ],
        'summary': {
            'decision_counts': {
                'approve_exclude': 3,
                'reject_exclude': 1,
                'unsure': 1,
            },
        },
    }
    return {
        'page_summaries': page_summaries,
        'raw_object_pages': [{
            'page_index': 0,
            'page_number': 1,
            'width': 600,
            'height': 1000,
            'raw_objects': [
                _raw_object(0, 'Approved Header', REGION_TOP),
                _raw_object(1, 'Body paragraph', REGION_BODY),
                _raw_object(2, 'Approved Footer', REGION_BOTTOM),
                _raw_object(3, 'Rejected Header', REGION_TOP),
                _raw_object(4, 'Unsure Footer', REGION_BOTTOM),
                _raw_object(5, IMAGE_PLACEHOLDER, REGION_TOP),
            ],
        }],
        'dry_run_report': dry_run_report,
        'review_decisions': review_decisions,
    }


def _parse_metrics(
        raw_count=0,
        text_blocks=0,
        body_text_blocks=0,
        paragraph_like=0,
        tables=0,
        images=0,
        sections=1):
    return {
        'parse_metrics_available': True,
        'raw_block_count': raw_count,
        'body_raw_block_count': body_text_blocks,
        'parsed_text_block_count': text_blocks,
        'body_text_block_count': body_text_blocks,
        'paragraph_like_text_block_count': paragraph_like or text_blocks,
        'table_count': tables,
        'image_count': images,
        'section_count': sections,
        'pages': [],
        'warnings': [],
    }


def _parse_metrics_with_tables(tables):
    return {
        'parse_metrics_available': True,
        'raw_block_count': 0,
        'body_raw_block_count': 0,
        'parsed_text_block_count': 0,
        'body_text_block_count': 0,
        'paragraph_like_text_block_count': 0,
        'table_count': len(tables),
        'image_count': 0,
        'section_count': 1,
        'tables': [dict(table) for table in tables],
        'pages': [{
            'page_index': 0,
            'page_number': 1,
            'tables': [dict(table) for table in tables],
        }],
        'warnings': [],
    }


def _geometry_safety_report_with_preserved_items(count):
    pages = [4, 7, 9]
    baseline_tables = []
    filtered_tables = []
    for index in range(count):
        page_index = pages[index % len(pages)]
        top = 300 + index * 12
        baseline_tables.append(_table_record(
            f'body-table-{index}',
            page_index,
            REGION_BODY,
            [50, top, 520, top + 40],
            rows=1,
            cols=2,
            cells=2,
            cell_bboxes=[
                [50, top, 285, top + 20],
                [285, top, 520, top + 20],
            ]))
        filtered_tables.append(_table_record(
            f'body-table-{index}-filtered',
            page_index,
            REGION_BODY,
            [50, top, 520, top + 38],
            rows=1,
            cols=2,
            cells=2,
            cell_bboxes=[
                [50, top, 285, top + 20],
                [285, top, 520, top + 20],
            ]))
    return _geometry_safety_report_from_tables(baseline_tables, filtered_tables)


def _geometry_safety_report_from_tables(baseline_tables, filtered_tables):
    return build_body_table_geometry_delta_safety_report(
        baseline_parse_metrics=_parse_metrics_with_tables(baseline_tables),
        filtered_parse_metrics=_parse_metrics_with_tables(filtered_tables),
        enabled=True)


def _table_visual_review_markdown(
        item_count=8,
        decisions=None,
        changed_counts=None,
        changed_text=None,
        extra_whitespace=False):
    decisions = decisions or {}
    changed_counts = changed_counts or set()
    changed_text = changed_text or set()
    lines = ['# Table Geometry Visual Review Pack', '']
    pages = [5, 8, 10]
    for index in range(1, item_count + 1):
        page_number = pages[(index - 1) % len(pages)]
        item_id = f'table-geometry-review-{index:03d}'
        cell_after = 2 if index not in changed_counts else 3
        text_preserved = 'False' if index in changed_text else 'True'
        decision = decisions.get(index, 'approve_safe_boundary_shift')
        markers = {
            'approve_safe_boundary_shift': '[ ]',
            'reject_unsafe_table_change': '[ ]',
            'unsure': '[ ]',
        }
        if decision:
            markers[decision] = '[x]'
        prefix = '-   ' if extra_whitespace else '- '
        lines.extend([
            f'### {item_id} | Page {page_number}',
            '',
            f'{prefix}baseline_table_id: baseline-{index}',
            f'{prefix}filtered_table_id: filtered-{index}',
            f'{prefix}row_count_before_after: 1 -> 1',
            f'{prefix}column_count_before_after: 2 -> 2',
            f'{prefix}cell_count_before_after: 2 -> {cell_after}',
            f'{prefix}text_cell_signature_preserved: {text_preserved}',
            (
                f'{prefix}human_decision:   approve_safe_boundary_shift: {markers["approve_safe_boundary_shift"]}    '
                f'reject_unsafe_table_change: {markers["reject_unsafe_table_change"]}    '
                f'unsure: {markers["unsure"]}'
            ),
            '',
        ])
    return '\n'.join(lines)


def _safe_readiness_inputs():
    return {
        'header_footer_review_report': {
            'approved_candidate_count': 4,
            'blocked_candidate_count': 5,
            'summary': {
                'would_remove_block_count': 48,
                'removed_block_count': 48,
            },
        },
        'raw_object_mapping_report': {
            'summary': {
                'approved_candidate_count': 4,
                'blocked_candidate_count': 5,
                'expected_would_remove_count': 48,
                'observed_would_remove_count': 48,
                'mapped_raw_object_count': 48,
                'exact_match_count': 48,
                'fuzzy_match_count': 0,
                'ambiguous_match_count': 0,
                'missing_match_count': 0,
                'unsafe_match_count': 0,
                'body_region_matched_for_removal_count': 0,
                'rejected_unsure_layout_placeholder_matched_for_removal_count': 0,
                'all_expected_blocks_mapped_once': True,
            },
        },
        'filtered_parse_experiment_report': _filtered_docx_experiment_report(),
        'table_visual_approval_gate_report': _passed_table_visual_gate_report(),
        'body_table_geometry_delta_safety_report': {
            'summary': {
                'changed_body_table_geometry_count': 8,
                'stream_table_boundary_adjustment_count': 8,
                'possible_body_table_structure_change_count': 0,
                'possible_cell_loss_count': 0,
                'unchanged_row_column_cell_count': 8,
                'changed_row_column_cell_count': 0,
                'text_cell_signature_preserved_count': 8,
                'text_cell_signature_changed_count': 0,
                'classification': 'review',
            },
        },
        'filtered_docx_comparison_report': {
            'summary': {
                'table_visual_approval_gate_status': 'passed',
                'baseline_raw_block_count': 790,
                'filtered_raw_block_count': 742,
                'removed_approved_header_footer_page_number_count': 48,
                'baseline_parsed_text_block_count': 523,
                'filtered_parsed_text_block_count': 486,
                'baseline_body_text_block_count': 393,
                'filtered_body_text_block_count': 393,
                'baseline_table_count': 139,
                'filtered_table_count': 127,
                'baseline_image_count': 0,
                'filtered_image_count': 0,
                'baseline_section_count': 50,
                'filtered_section_count': 50,
                'body_region_removed_count': 0,
                'rejected_unsure_layout_placeholder_removed_count': 0,
                'normal_conversion_still_works': True,
                'state_restored_or_reloaded': True,
            },
            'docx_files': {
                'baseline': {
                    'path': 'local_reports/docx_compare/baseline.docx',
                    'local_only_path': True,
                    'exists': True,
                    'size_bytes': 1200,
                    'empty': False,
                },
                'filtered': {
                    'path': 'local_reports/docx_compare/filtered.docx',
                    'local_only_path': True,
                    'exists': True,
                    'size_bytes': 1100,
                    'empty': False,
                },
            },
            'safety_warnings': [],
        },
        'docx_residual_structure_report': {
            'summary': {
                'classification': 'safe',
                'residual_removed_string_count': 1,
                'true_residual_header_footer_pollution_count': 0,
                'body_text_loss_warning_count': 0,
                'table_text_loss_warning_count': 0,
            },
            'safety_warnings': [],
        },
        'verification_status': {
            'layout_analyzer_tests_passed': True,
            'py_compile_passed': True,
            'unittest_passed': True,
            'conversion_tests_passed': True,
            'git_diff_check_passed': True,
            'local_artifacts_ignored': True,
            'local_sample_dependency': True,
            'committed_synthetic_fixture_available': False,
            'committed_end_to_end_regression_fixture_available': False,
            'production_default_integration_enabled': False,
            'public_cli_enabled': False,
        },
    }


def _readiness_reason_types(report):
    return {reason['type'] for reason in report['blocking_reasons']}


def _corpus_layout_report():
    return {
        'page_count': 3,
        'pages': [
            {
                'page_index': 0,
                'text_block_count': 3,
                'region_counts': {
                    REGION_TOP: 1,
                    REGION_BODY: 1,
                    REGION_BOTTOM: 1,
                },
            },
            {
                'page_index': 1,
                'text_block_count': 3,
                'region_counts': {
                    REGION_TOP: 1,
                    REGION_BODY: 1,
                    REGION_BOTTOM: 1,
                },
            },
            {
                'page_index': 2,
                'text_block_count': 3,
                'region_counts': {
                    REGION_TOP: 1,
                    REGION_BODY: 1,
                    REGION_BOTTOM: 1,
                },
            },
        ],
        'repeated_text_candidates': [
            {
                'fingerprint': 'annual report||top',
                'confidence_label': 'strong',
            },
            {
                'fingerprint': 'page-number||bottom',
                'confidence_label': 'placeholder',
            },
        ],
        'header_footer_exclusion_dry_run': {
            'summary': {
                'candidate_count': 2,
                'action_counts': {
                    ACTION_WOULD_EXCLUDE: 2,
                },
                'role_counts': {
                    ROLE_HEADER: 1,
                    ROLE_PAGE_NUMBER: 1,
                },
            },
            'candidates': [
                {
                    'candidate_id': 'repeated-1',
                    'proposed_role': ROLE_HEADER,
                    'action': ACTION_WOULD_EXCLUDE,
                    'region': REGION_TOP,
                    'regions': [REGION_TOP],
                    'support_count': 3,
                    'page_count': 3,
                },
                {
                    'candidate_id': 'repeated-2',
                    'proposed_role': ROLE_PAGE_NUMBER,
                    'action': ACTION_WOULD_EXCLUDE,
                    'region': REGION_BOTTOM,
                    'regions': [REGION_BOTTOM],
                    'support_count': 3,
                    'page_count': 3,
                },
            ],
        },
    }


def _corpus_sample_result(
        file_name,
        layout_report,
        parsing_succeeded=True,
        analysis_succeeded=True,
        page_count=3,
        pages_analyzed=3,
        analysis_mode='analysis_only',
        error=''):
    return {
        'sample_name': file_name,
        'file_name': file_name,
        'file_path': f'local_samples/{file_name}',
        'file_size_bytes': 1234,
        'page_count': page_count,
        'pages_analyzed': pages_analyzed,
        'parsing_succeeded': parsing_succeeded,
        'analysis_succeeded': analysis_succeeded,
        'runtime_seconds': 0.12,
        'analysis_mode': analysis_mode,
        'layout_analysis_report': layout_report,
        'review_pack_generated': analysis_succeeded,
        'review_pack_path': f'local_reports/corpus_validation/{file_name}-review-pack.md',
        'error': error,
    }


def _corpus_approval_layout_report():
    pages = []
    for page_index in range(2):
        pages.append({
            'page_index': page_index,
            'text_blocks': [
                {
                    'block_index': 0,
                    'text': 'Annual Report',
                    'fingerprint': 'annual report||top',
                    'region': REGION_TOP,
                },
                {
                    'block_index': 1,
                    'text': 'Body paragraph line',
                    'fingerprint': f'body paragraph {page_index}||body',
                    'region': REGION_BODY,
                },
                {
                    'block_index': 2,
                    'text': 'Page 1',
                    'fingerprint': 'page-number||bottom',
                    'region': REGION_BOTTOM,
                },
                {
                    'block_index': 3,
                    'text': 'Body Repeat',
                    'fingerprint': 'body repeat||body',
                    'region': REGION_BODY,
                },
            ],
        })
    return {
        'page_count': 2,
        'pages': pages,
        'header_footer_exclusion_dry_run': {
            'summary': {
                'candidate_count': 3,
            },
            'candidates': [
                {
                    'candidate_id': 'repeated-1',
                    'fingerprint': 'annual report||top',
                    'proposed_role': ROLE_HEADER,
                    'action': ACTION_WOULD_EXCLUDE,
                    'region': REGION_TOP,
                    'regions': [REGION_TOP],
                    'support_count': 2,
                    'page_count': 2,
                    'affected_pages': [0, 1],
                },
                {
                    'candidate_id': 'repeated-2',
                    'fingerprint': 'page-number||bottom',
                    'proposed_role': ROLE_PAGE_NUMBER,
                    'action': ACTION_WOULD_EXCLUDE,
                    'region': REGION_BOTTOM,
                    'regions': [REGION_BOTTOM],
                    'support_count': 2,
                    'page_count': 2,
                    'affected_pages': [0, 1],
                },
                {
                    'candidate_id': 'repeated-3',
                    'fingerprint': 'body repeat||body',
                    'proposed_role': ROLE_REVIEW_ONLY,
                    'action': ACTION_REVIEW,
                    'region': REGION_BODY,
                    'regions': [REGION_BODY],
                    'support_count': 2,
                    'page_count': 2,
                    'affected_pages': [0, 1],
                },
            ],
        },
    }


def _corpus_review_decisions(*items):
    decisions = [
        {
            'candidate_id': candidate_id,
            'fingerprint': fingerprint,
            'proposed_role': role,
            'action': action,
            'manual_decision': decision,
            'checked_decisions': [decision],
        }
        for candidate_id, fingerprint, role, action, decision in items
    ]
    return {
        'decisions': decisions,
        'summary': {
            'candidate_count': len(decisions),
            'decision_counts': dict(Counter(
                decision['manual_decision'] for decision in decisions)),
        },
    }


def _internal_config_dry_run_report():
    return {
        'candidates': [
            {
                'candidate_id': 'header-1',
                'fingerprint': 'synthetic header||top',
                'proposed_role': ROLE_HEADER,
                'action': ACTION_WOULD_EXCLUDE,
                'region': REGION_TOP,
                'regions': [REGION_TOP],
                'support_count': 3,
                'page_count': 3,
                'affected_pages': [0, 1, 2],
            },
            {
                'candidate_id': 'footer-1',
                'fingerprint': 'synthetic footer||bottom',
                'proposed_role': ROLE_FOOTER,
                'action': ACTION_WOULD_EXCLUDE,
                'region': REGION_BOTTOM,
                'regions': [REGION_BOTTOM],
                'support_count': 3,
                'page_count': 3,
                'affected_pages': [0, 1, 2],
            },
            {
                'candidate_id': 'body-1',
                'fingerprint': 'body repeat||body',
                'proposed_role': ROLE_HEADER,
                'action': ACTION_WOULD_EXCLUDE,
                'region': REGION_BODY,
                'regions': [REGION_BODY],
                'support_count': 3,
                'page_count': 3,
                'affected_pages': [0, 1, 2],
            },
            {
                'candidate_id': 'placeholder-1',
                'fingerprint': '<image>||top',
                'proposed_role': ROLE_LAYOUT_PLACEHOLDER,
                'action': ACTION_WOULD_EXCLUDE,
                'region': REGION_TOP,
                'regions': [REGION_TOP],
                'support_count': 3,
                'page_count': 3,
                'affected_pages': [0, 1, 2],
            },
        ],
    }


def _internal_config_safe_dry_run_report():
    report = _internal_config_dry_run_report()
    return {
        'candidates': [
            candidate for candidate in report['candidates']
            if candidate['candidate_id'] in {'header-1', 'footer-1'}
        ],
    }


def _internal_config_decisions(*items):
    decisions = [
        {
            'candidate_id': candidate_id,
            'fingerprint': fingerprint,
            'manual_decision': decision,
            'checked_decisions': [decision],
        }
        for candidate_id, fingerprint, decision in items
    ]
    return {
        'decisions': decisions,
        'summary': {
            'candidate_count': len(decisions),
            'decision_counts': dict(Counter(
                decision['manual_decision'] for decision in decisions)),
        },
    }


def _config_warning_types(report):
    return {warning['type'] for warning in report.get('warnings', [])}


def _require_synthetic_pdf_support(testcase):
    if fitz is None or Converter is None or Pages is None:
        testcase.skipTest('Synthetic PDF regression tests require PyMuPDF and pdf2docx.')


def _require_synthetic_docx_support(testcase):
    _require_synthetic_pdf_support(testcase)
    if DocxDocument is None:
        testcase.skipTest('Synthetic DOCX comparison tests require python-docx.')


def _require_docx_header_footer_support(testcase):
    if DocxDocument is None or docx_utils is None:
        testcase.skipTest('DOCX header/footer tests require python-docx helpers.')


def _write_synthetic_pdf(path, scenario):
    doc = fitz.open()
    try:
        if scenario == 'repeated_header_footer':
            for page_index in range(4):
                page = _synthetic_page(doc)
                _synthetic_header(page, 'SYNTHETIC REPORT HEADER')
                _synthetic_body_line(
                    page,
                    120,
                    f'Synthetic body paragraph page {page_index + 1} remains editable.')
                _synthetic_body_line(
                    page,
                    142,
                    'The body region text must remain after reviewed filtering.')
                _synthetic_footer(page, 'SYNTHETIC REPORT FOOTER')
                _synthetic_page_number(page, page_index + 1, 4)
        elif scenario == 'body_table_near_footer':
            for page_index in range(3):
                page = _synthetic_page(doc)
                _synthetic_body_line(
                    page,
                    110,
                    f'Synthetic body introduction {page_index + 1}.')
                _synthetic_body_table(page, page_index)
                _synthetic_footer(page, 'SYNTHETIC TABLE FOOTER')
                _synthetic_page_number(page, page_index + 1, 3)
        elif scenario == 'callout_text_box_near_edges':
            for page_index in range(3):
                page = _synthetic_page(doc)
                _synthetic_header(page, 'SYNTHETIC CALLOUT HEADER')
                _synthetic_body_line(
                    page,
                    132,
                    f'Callout scenario body introduction {page_index + 1}.')
                _synthetic_callout_box(page, page_index)
                _synthetic_footer(page, 'SYNTHETIC CALLOUT FOOTER')
                _synthetic_page_number(page, page_index + 1, 3)
        elif scenario == 'list_heading_boundaries':
            for page_index in range(4):
                page = _synthetic_page(doc)
                _synthetic_header(page, 'SYNTHETIC LIST HEADER')
                _synthetic_body_line(
                    page,
                    150,
                    f'Synthetic section heading {page_index + 1} remains body text')
                _synthetic_body_line(
                    page,
                    178,
                    f'- List item alpha {page_index + 1} remains in body')
                _synthetic_body_line(
                    page,
                    204,
                    f'1. Numbered item beta {page_index + 1} remains in body')
                _synthetic_body_line(
                    page,
                    232,
                    f'Closing list paragraph {page_index + 1} remains body text.')
                _synthetic_footer(page, 'SYNTHETIC LIST FOOTER')
                _synthetic_page_number(page, page_index + 1, 4)
        elif scenario == 'table_geometry_delta_stress':
            for page_index in range(3):
                page = _synthetic_page(doc)
                _synthetic_body_line(
                    page,
                    122,
                    f'Synthetic geometry body introduction {page_index + 1}.')
                _synthetic_geometry_stress_table(page, page_index)
                _synthetic_footer(page, 'SYNTHETIC GEOMETRY FOOTER')
                _synthetic_page_number(page, page_index + 1, 3)
        elif scenario == 'no_header_footer':
            for page_index in range(3):
                page = _synthetic_page(doc)
                _synthetic_body_line(
                    page,
                    160,
                    f'Unique synthetic body paragraph {page_index + 1}.')
                _synthetic_body_line(
                    page,
                    190,
                    f'No repeated header or footer appears on page {page_index + 1}.')
        elif scenario == 'first_page_odd_even_headers':
            for page_index in range(5):
                page = _synthetic_page(doc)
                if page_index == 0:
                    header = 'SYNTHETIC FIRST PAGE HEADER'
                elif (page_index + 1) % 2:
                    header = 'SYNTHETIC ODD HEADER'
                else:
                    header = 'SYNTHETIC EVEN HEADER'
                _synthetic_header(page, header)
                _synthetic_body_line(
                    page,
                    150,
                    f'Synthetic body page {page_index + 1} with varied headers.')
        elif scenario == 'odd_even_body_heading_interaction':
            for page_index in range(5):
                page = _synthetic_page(doc)
                if page_index == 0:
                    header = 'SYNTHETIC FIRST PAGE HEADER'
                elif (page_index + 1) % 2:
                    header = 'SYNTHETIC ODD HEADER'
                else:
                    header = 'SYNTHETIC EVEN HEADER'
                _synthetic_header(page, header)
                _synthetic_body_line(
                    page,
                    152,
                    f'{header.title()} body heading {page_index + 1} must remain')
                _synthetic_body_line(
                    page,
                    184,
                    f'Body paragraph below varied heading {page_index + 1} remains.')
                _synthetic_footer(page, 'SYNTHETIC VARIED FOOTER')
                _synthetic_page_number(page, page_index + 1, 5)
        elif scenario == 'paragraph_continuity':
            page = _synthetic_page(doc)
            _synthetic_header(page, 'SYNTHETIC CONTINUITY HEADER')
            _synthetic_body_line(
                page,
                620,
                'This synthetic paragraph continues across the page break and ends with a hyphen-')
            _synthetic_footer(page, 'SYNTHETIC CONTINUITY FOOTER')
            _synthetic_page_number(page, 1, 2)

            page = _synthetic_page(doc)
            _synthetic_header(page, 'SYNTHETIC CONTINUITY HEADER')
            _synthetic_body_line(
                page,
                135,
                'ated continuation text resumes in the same paragraph with matching style.')
            _synthetic_body_line(
                page,
                170,
                'A later body sentence remains available for paragraph grouping.')
            _synthetic_footer(page, 'SYNTHETIC CONTINUITY FOOTER')
            _synthetic_page_number(page, 2, 2)
        else:
            raise ValueError(f'Unknown synthetic scenario: {scenario}')

        doc.save(str(path))
    finally:
        doc.close()


def _synthetic_page(doc):
    return doc.new_page(width=612, height=792)


def _synthetic_header(page, text):
    page.insert_text((54, 36), text, fontsize=9)


def _synthetic_footer(page, text):
    page.insert_text((54, 732), text, fontsize=9)


def _synthetic_page_number(page, page_number, page_count):
    page.insert_text((276, 758), f'Page {page_number} of {page_count}', fontsize=9)


def _synthetic_body_line(page, y, text):
    page.insert_text((54, y), text, fontsize=11)


def _synthetic_body_table(page, page_index):
    x0, y0, x1, y1 = 54, 596, 420, 668
    page.draw_rect(fitz.Rect(x0, y0, x1, y1), width=0.5)
    page.draw_line((x0, y0 + 24), (x1, y0 + 24), width=0.5)
    page.draw_line((x0, y0 + 48), (x1, y0 + 48), width=0.5)
    page.draw_line((190, y0), (190, y1), width=0.5)
    _synthetic_body_line(page, 612, 'Body Table Header')
    _synthetic_body_line(page, 636, f'Body table cell alpha {page_index + 1}')
    _synthetic_body_line(page, 660, f'Body table cell beta {page_index + 1}')


def _synthetic_callout_box(page, page_index):
    x0, y0, x1, y1 = 72, 582, 470, 656
    page.draw_rect(fitz.Rect(x0, y0, x1, y1), width=0.5)
    page.draw_line((x0, y0 + 25), (x1, y0 + 25), width=0.5)
    _synthetic_body_line(
        page,
        602,
        f'Callout panel body content {page_index + 1} must remain.')
    _synthetic_body_line(
        page,
        630,
        f'Table like callout row alpha {page_index + 1} must remain.')


def _synthetic_geometry_stress_table(page, page_index):
    x0, y0, x1, y1 = 54, 574, 462, 646
    page.draw_rect(fitz.Rect(x0, y0, x1, y1), width=0.5)
    page.draw_line((x0, y0 + 24), (x1, y0 + 24), width=0.5)
    page.draw_line((x0, y0 + 48), (x1, y0 + 48), width=0.5)
    page.draw_line((210, y0), (210, y1), width=0.5)
    _synthetic_body_line(page, 590, 'Geometry Table Header')
    _synthetic_body_line(page, 614, f'Geometry table cell alpha {page_index + 1}')
    _synthetic_body_line(page, 638, f'Geometry table cell beta {page_index + 1}')


def _parse_synthetic_layout(pdf_path):
    converter = Converter(str(pdf_path))
    settings = converter.default_settings.copy()
    settings.update({'layout_analysis': True})
    try:
        converter.load_pages().parse_document(**settings)
        return converter.pages.layout_analysis_report
    finally:
        converter.close()


def _synthetic_review_decisions(layout, approve_roles):
    decisions = []
    for candidate in layout['header_footer_exclusion_dry_run']['candidates']:
        decision = (
            'approve_exclude'
            if candidate.get('proposed_role') in approve_roles else
            'reject_exclude')
        decisions.append({
            'candidate_id': candidate.get('candidate_id', ''),
            'fingerprint': candidate.get('fingerprint', ''),
            'proposed_role': candidate.get('proposed_role', ''),
            'action': candidate.get('action', ''),
            'manual_decision': decision,
            'checked_decisions': [decision],
        })
    return {
        'decisions': decisions,
        'summary': {
            'candidate_count': len(decisions),
            'decision_counts': dict(Counter(
                decision['manual_decision'] for decision in decisions)),
        },
    }


def _synthetic_filtering_report(layout, decisions):
    return build_reviewed_header_footer_filter_report(
        layout.get('pages', []),
        layout.get('header_footer_exclusion_dry_run', {}),
        decisions,
        enabled=True,
        apply=True)


def _run_synthetic_document_parse_diagnostics(
        pdf_path,
        decisions,
        expected_remove_count):
    converter = Converter(str(pdf_path))
    settings = converter.default_settings.copy()
    settings.update({
        'layout_analysis': True,
        '_document_parse_raw_object_mapping_enabled': True,
        '_document_parse_copied_raw_filtering_enabled': True,
        '_document_parse_guarded_raw_apply_restore_enabled': True,
        '_document_parse_filtering_review_decisions': decisions,
        '_document_parse_mapping_expected_would_remove_count': expected_remove_count,
        '_document_parse_copied_raw_filtering_expected_mapping_count': expected_remove_count,
        '_document_parse_guarded_raw_apply_restore_expected_mapping_count': expected_remove_count,
    })
    try:
        converter.load_pages().parse_document(**settings)
        return {
            'mapping': converter.pages._document_parse_raw_object_mapping_report,
            'copied': converter.pages._document_parse_copied_raw_filtering_apply_report,
            'guarded': converter.pages._document_parse_guarded_raw_apply_restore_report,
        }
    finally:
        converter.close()


def _run_synthetic_internal_filtered_parse_integration(
        pdf_path,
        review_decisions,
        enabled=True,
        run_parse_pages=False):
    converter = Converter(str(pdf_path))
    settings = converter.default_settings.copy()
    settings['_reviewed_header_footer_filtering_config'] = build_reviewed_filtering_internal_config({
        'enabled': enabled,
        'mode': 'filtered_parse_experiment',
        'review_decisions': review_decisions,
    })
    try:
        converter.load_pages().parse_document(**settings)
        if run_parse_pages:
            converter.parse_pages(**settings)
        return {
            'config': converter.pages._reviewed_filtering_internal_config_report,
            'integration': converter.pages._reviewed_filtering_internal_filtered_parse_report,
            'filtered_experiment': converter.pages._document_parse_filtered_parse_experiment_report,
            'finalized_pages': sum(
                1 for page in converter.pages
                if getattr(page, 'finalized', False)),
        }
    finally:
        converter.close()


def _run_synthetic_filtered_docx_comparison(
        pdf_path,
        review_decisions,
        layout,
        temp_root):
    baseline_path = Path(temp_root) / 'baseline.docx'
    filtered_path = Path(temp_root) / 'filtered.docx'
    default_after_path = Path(temp_root) / 'default-after.docx'

    baseline = _convert_synthetic_pdf_to_docx(pdf_path, baseline_path)
    filtered = _convert_synthetic_pdf_to_docx(
        pdf_path,
        filtered_path,
        review_decisions=review_decisions)
    default_after = _convert_synthetic_pdf_to_docx(pdf_path, default_after_path)
    filtering_report = _synthetic_filtering_report(layout, review_decisions)

    return _synthetic_docx_comparison_report_from_metrics(
        baseline['metrics'],
        filtered['metrics'],
        expected_body_texts=_body_text_fragments(layout),
        removed_texts=_removed_text_fragments(filtering_report),
        baseline_path=str(baseline_path),
        filtered_path=str(filtered_path),
        temp_root=str(temp_root),
        internal_filtered_parse_report=filtered['internal_filtered_parse_report'],
        default_after_metrics=default_after['metrics'])


def _run_synthetic_filtered_docx_with_header_footer_parts(
        pdf_path,
        review_decisions,
        layout,
        temp_root):
    baseline_path = Path(temp_root) / 'baseline-body.docx'
    filtered_body_path = Path(temp_root) / 'filtered-body.docx'
    final_path = Path(temp_root) / 'filtered-body-with-header-footer.docx'
    default_after_path = Path(temp_root) / 'default-after-body.docx'

    baseline = _convert_synthetic_pdf_to_docx(pdf_path, baseline_path)
    filtered_body = _convert_synthetic_pdf_to_docx(
        pdf_path,
        filtered_body_path,
        review_decisions=review_decisions)
    default_after = _convert_synthetic_pdf_to_docx(pdf_path, default_after_path)
    filtering_report = _synthetic_filtering_report(layout, review_decisions)
    plan = build_docx_header_footer_generation_plan(
        layout.get('pages', []),
        layout.get('header_footer_exclusion_dry_run', {}),
        review_decisions,
        enabled=True)

    document = DocxDocument(str(filtered_body_path))
    apply_report = docx_utils.apply_header_footer_text_plan(
        document,
        plan,
        enabled=True)
    document.save(str(final_path))

    final_metrics = _read_docx_file_metrics(final_path)
    openxml = _read_docx_openxml_parts(final_path)
    default_openxml = _read_docx_openxml_parts(default_after_path)
    default_metrics = dict(default_after['metrics'])
    default_metrics['header_footer_xml'] = {
        'has_header_text': bool(default_openxml['header_xml'].strip()),
        'has_footer_text': bool(default_openxml['footer_xml'].strip()),
    }
    comparison = _synthetic_docx_comparison_report_from_metrics(
        baseline['metrics'],
        final_metrics,
        expected_body_texts=_body_text_fragments(layout),
        removed_texts=_removed_text_fragments(filtering_report),
        baseline_path=str(baseline_path),
        filtered_path=str(final_path),
        temp_root=str(temp_root),
        internal_filtered_parse_report=filtered_body['internal_filtered_parse_report'],
        default_after_metrics=default_after['metrics'])

    return {
        'baseline_metrics': baseline['metrics'],
        'filtered_body_metrics': filtered_body['metrics'],
        'final_metrics': final_metrics,
        'default_metrics': default_metrics,
        'filtering_report': filtering_report,
        'plan': plan,
        'apply_report': apply_report,
        'openxml': openxml,
        'comparison': comparison,
        'paths': {
            'baseline_docx_path': str(baseline_path),
            'filtered_body_docx_path': str(filtered_body_path),
            'final_docx_path': str(final_path),
            'default_after_docx_path': str(default_after_path),
        },
    }


def _convert_synthetic_pdf_to_docx(
        pdf_path,
        docx_path,
        review_decisions=None,
        enabled=True):
    converter = Converter(str(pdf_path))
    settings = converter.default_settings.copy()
    if review_decisions is not None:
        settings['_reviewed_header_footer_filtering_config'] = build_reviewed_filtering_internal_config({
            'enabled': enabled,
            'mode': 'filtered_parse_experiment',
            'review_decisions': review_decisions,
        })
    try:
        converter.convert(str(docx_path), **settings)
        return {
            'metrics': _read_docx_file_metrics(docx_path),
            'internal_config_report': converter.pages._reviewed_filtering_internal_config_report,
            'internal_filtered_parse_report': (
                converter.pages._reviewed_filtering_internal_filtered_parse_report or {}),
        }
    finally:
        converter.close()


def _read_docx_file_metrics(docx_path):
    path = Path(docx_path)
    exists = path.exists()
    size = path.stat().st_size if exists else 0
    metrics = {
        'path': str(path),
        'exists': exists,
        'is_zipfile': zipfile.is_zipfile(str(path)) if exists else False,
        'size': size,
        'paragraph_count': 0,
        'non_empty_paragraph_count': 0,
        'table_count': 0,
        'paragraph_texts': [],
        'table_texts': [],
        'all_text': '',
    }
    if not exists or not size or DocxDocument is None:
        return metrics

    doc = DocxDocument(str(path))
    paragraph_texts = [
        normalize_text(paragraph.text)
        for paragraph in doc.paragraphs
        if normalize_text(paragraph.text)
    ]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = normalize_text(paragraph.text)
                    if text:
                        table_texts.append(text)

    metrics.update({
        'paragraph_count': len(doc.paragraphs),
        'non_empty_paragraph_count': len(paragraph_texts),
        'table_count': len(doc.tables),
        'paragraph_texts': paragraph_texts,
        'table_texts': table_texts,
        'all_text': normalize_text(' '.join(paragraph_texts + table_texts)).lower(),
    })
    return metrics


def _read_docx_openxml_parts(docx_path):
    parts = {
        'part_names': [],
        'body_xml': '',
        'header_xml': '',
        'footer_xml': '',
    }
    path = Path(docx_path)
    if not path.exists() or not zipfile.is_zipfile(str(path)):
        return parts

    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        parts['part_names'] = names
        if 'word/document.xml' in names:
            parts['body_xml'] = archive.read('word/document.xml').decode('utf-8')
        parts['header_xml'] = ''.join(
            archive.read(name).decode('utf-8')
            for name in names
            if name.startswith('word/header') and name.endswith('.xml'))
        parts['footer_xml'] = ''.join(
            archive.read(name).decode('utf-8')
            for name in names
            if name.startswith('word/footer') and name.endswith('.xml'))
    return parts


def _synthetic_docx_comparison_report_from_metrics(
        baseline_metrics,
        filtered_metrics,
        expected_body_texts,
        removed_texts,
        baseline_path,
        filtered_path,
        temp_root,
        internal_filtered_parse_report=None,
        default_after_metrics=None):
    baseline_metrics = baseline_metrics or {}
    filtered_metrics = filtered_metrics or {}
    internal_report = internal_filtered_parse_report or {}
    internal_summary = internal_report.get('summary') or {}
    default_after_metrics = default_after_metrics or baseline_metrics
    filtered_text = filtered_metrics.get('all_text', '')
    baseline_text = baseline_metrics.get('all_text', '')
    body_fragments = [text for text in expected_body_texts or [] if text]
    body_fragment_set = set(body_fragments)
    missing_body = [
        text for text in body_fragments
        if text not in filtered_text
    ]
    removed_fragments = [text for text in removed_texts or [] if text]
    residual_removed = [
        text for text in removed_fragments
        if text in filtered_text
    ]
    true_residual = [
        text for text in residual_removed
        if text not in body_fragment_set
    ]
    baseline_table_texts = [
        normalize_text(text).lower()
        for text in baseline_metrics.get('table_texts', []) or []
        if normalize_text(text)
    ]
    filtered_table_texts = [
        normalize_text(text).lower()
        for text in filtered_metrics.get('table_texts', []) or []
        if normalize_text(text)
    ]
    missing_table_signature_texts = [
        text for text in baseline_table_texts
        if text and text not in filtered_table_texts
    ]
    missing_table_text = sorted(set([
        text for text in missing_body
        if 'body table' in text
    ] + missing_table_signature_texts))
    docx_artifacts_temp_only = (
        _path_inside_root(baseline_path, temp_root) and
        _path_inside_root(filtered_path, temp_root))
    default_unchanged = (
        baseline_text == default_after_metrics.get('all_text', '') and
        baseline_metrics.get('paragraph_count', 0) == default_after_metrics.get('paragraph_count', 0) and
        baseline_metrics.get('table_count', 0) == default_after_metrics.get('table_count', 0))

    safety_warnings = []
    if not baseline_metrics.get('exists') or not baseline_metrics.get('size'):
        safety_warnings.append({'type': 'baseline_docx_missing_or_empty'})
    if not filtered_metrics.get('exists') or not filtered_metrics.get('size'):
        safety_warnings.append({'type': 'filtered_docx_missing_or_empty'})
    if not docx_artifacts_temp_only:
        safety_warnings.append({'type': 'docx_artifact_not_temp_only'})
    if true_residual:
        safety_warnings.append({
            'type': 'true_residual_header_footer_pollution',
            'count': len(true_residual),
        })
    if missing_body:
        safety_warnings.append({
            'type': 'body_text_loss',
            'count': len(missing_body),
        })
    if missing_table_text:
        safety_warnings.append({
            'type': 'table_text_loss',
            'count': len(missing_table_text),
        })
    if not default_unchanged:
        safety_warnings.append({'type': 'default_conversion_changed_after_experiment'})

    for warning in internal_report.get('safety_warnings', []) or []:
        if warning.get('type') in {
                'body_text_signature_changed',
                'body_region_removed_during_internal_filtered_parse',
                'blocked_or_placeholder_removed_during_internal_filtered_parse',
                'mapping_count_mismatch',
                'copied_apply_count_mismatch'}:
            safety_warnings.append({
                'type': f'internal_{warning.get("type")}',
            })

    body_text_block_delta = internal_summary.get('body_text_block_delta')
    table_count_delta = (
        int(filtered_metrics.get('table_count', 0)) -
        int(baseline_metrics.get('table_count', 0)))
    summary = {
        'baseline_docx_path': baseline_path,
        'filtered_docx_path': filtered_path,
        'baseline_docx_exists': bool(baseline_metrics.get('exists')),
        'filtered_docx_exists': bool(filtered_metrics.get('exists')),
        'baseline_docx_size': int(baseline_metrics.get('size', 0)),
        'filtered_docx_size': int(filtered_metrics.get('size', 0)),
        'baseline_docx_paragraph_count': int(baseline_metrics.get('paragraph_count', 0)),
        'filtered_docx_paragraph_count': int(filtered_metrics.get('paragraph_count', 0)),
        'baseline_docx_table_count': int(baseline_metrics.get('table_count', 0)),
        'filtered_docx_table_count': int(filtered_metrics.get('table_count', 0)),
        'docx_table_count_delta': table_count_delta,
        'internal_filtered_parse_applied': bool(internal_report.get('applied_to_parse')),
        'removed_approved_header_footer_page_number_count': (
            int(internal_summary.get('removed_raw_block_count', 0))
            if internal_report.get('applied_to_parse') else 0),
        'removed_header_footer_page_number_residual_count': len(residual_removed),
        'true_residual_header_footer_pollution_count': len(true_residual),
        'body_text_signature_preserved': not missing_body,
        'body_text_loss_warning_count': len(missing_body),
        'table_text_signature_preserved': not missing_table_text,
        'table_text_loss_warning_count': len(missing_table_text),
        'body_text_block_delta': body_text_block_delta,
        'body_text_block_delta_classification': (
            'acceptable_boundary_or_grouping_shift'
            if body_text_block_delta and not missing_body else
            'unsafe_body_text_loss'
            if body_text_block_delta and missing_body else
            'unchanged_or_unavailable'),
        'table_count_delta_classification': (
            'reported_no_table_text_loss'
            if table_count_delta and not missing_table_text else
            'unsafe_table_text_loss'
            if table_count_delta and missing_table_text else
            'unchanged'),
        'generated_docx_artifacts_temp_only': docx_artifacts_temp_only,
        'default_conversion_unchanged_after_experiment': default_unchanged,
    }
    return {
        'summary': summary,
        'baseline_docx_metrics': baseline_metrics,
        'filtered_docx_metrics': filtered_metrics,
        'expected_body_texts': body_fragments,
        'removed_texts': removed_fragments,
        'residual_removed_texts': residual_removed,
        'missing_body_texts': missing_body,
        'safety_warnings': safety_warnings,
        'internal_filtered_parse_report': internal_report,
        'recommendation': {
            'safe_for_internal_filtered_docx': not safety_warnings,
            'reason': (
                'Synthetic filtered DOCX comparison preserved body/table text and removed approved residuals.'
                if not safety_warnings else
                'Synthetic filtered DOCX comparison produced fail-closed warnings.'),
        },
    }


def _local_corpus_docx_smoke_sample_report(
        sample_name,
        baseline_docx_metrics=None,
        filtered_docx_metrics=None,
        internal_filtered_parse_report=None,
        approval_artifacts_available=True,
        bounded_subset_only=False,
        full_document_skipped=False,
        body_text_signature_preserved=True,
        true_residual_header_footer_pollution_count=0,
        body_text_loss_warning_count=0,
        table_text_loss_warning_count=0,
        generated_paths=None,
        skipped_reason=''):
    baseline_docx_metrics = baseline_docx_metrics or {}
    filtered_docx_metrics = filtered_docx_metrics or {}
    internal_report = internal_filtered_parse_report or {}
    internal_summary = internal_report.get('summary') or {}
    generated_paths = list(generated_paths or [])
    diagnostic_warnings = []
    safety_warnings = []

    if skipped_reason:
        safety_warnings.append({
            'type': 'sample_skipped',
            'reason': skipped_reason,
        })
    if not approval_artifacts_available:
        safety_warnings.append({'type': 'missing_approval_artifacts'})
    if not _local_docx_metric_exists(baseline_docx_metrics):
        safety_warnings.append({'type': 'baseline_docx_missing_or_empty'})
    if not _local_docx_metric_exists(filtered_docx_metrics):
        safety_warnings.append({'type': 'filtered_docx_missing_or_empty'})
    if not body_text_signature_preserved:
        safety_warnings.append({'type': 'body_text_signature_not_preserved'})
    if true_residual_header_footer_pollution_count:
        safety_warnings.append({
            'type': 'true_residual_header_footer_pollution',
            'count': true_residual_header_footer_pollution_count,
        })
    if body_text_loss_warning_count:
        safety_warnings.append({
            'type': 'body_text_loss',
            'count': body_text_loss_warning_count,
        })
    if table_text_loss_warning_count:
        safety_warnings.append({
            'type': 'table_text_loss',
            'count': table_text_loss_warning_count,
        })
    if not _local_generated_paths_are_local_only(generated_paths):
        safety_warnings.append({'type': 'generated_docx_path_not_local_only'})

    body_text_block_delta = internal_summary.get('body_text_block_delta')
    if body_text_block_delta:
        diagnostic_warnings.append({
            'type': 'body_text_block_count_changed',
            'delta': body_text_block_delta,
            'classification': (
                'acceptable_boundary_or_grouping_shift'
                if body_text_signature_preserved else
                'unsafe_body_text_loss'),
        })

    table_delta = (
        _local_docx_metric_int(filtered_docx_metrics, 'table_count') -
        _local_docx_metric_int(baseline_docx_metrics, 'table_count'))
    if table_delta:
        diagnostic_warnings.append({
            'type': 'docx_table_count_changed',
            'delta': table_delta,
            'classification': (
                'reported_no_table_text_loss'
                if not table_text_loss_warning_count else
                'unsafe_table_text_loss'),
        })

    status = 'blocked' if safety_warnings else 'passed'
    if bounded_subset_only and not safety_warnings:
        status = 'bounded_subset_passed'
    if skipped_reason and not approval_artifacts_available:
        status = 'skipped'

    return {
        'sample_name': sample_name,
        'status': status,
        'bounded_subset_only': bool(bounded_subset_only),
        'full_document_skipped': bool(full_document_skipped),
        'generated_paths': generated_paths,
        'summary': {
            'baseline_docx_paragraph_count': _local_docx_metric_int(
                baseline_docx_metrics, 'paragraph_count'),
            'filtered_docx_paragraph_count': _local_docx_metric_int(
                filtered_docx_metrics, 'paragraph_count'),
            'baseline_docx_table_count': _local_docx_metric_int(
                baseline_docx_metrics, 'table_count'),
            'filtered_docx_table_count': _local_docx_metric_int(
                filtered_docx_metrics, 'table_count'),
            'removed_approved_header_footer_page_number_count': _local_docx_metric_int(
                internal_summary, 'removed_raw_block_count'),
            'true_residual_header_footer_pollution_count': int(
                true_residual_header_footer_pollution_count or 0),
            'body_text_signature_preserved': bool(body_text_signature_preserved),
            'body_text_loss_warning_count': int(body_text_loss_warning_count or 0),
            'table_text_loss_warning_count': int(table_text_loss_warning_count or 0),
            'body_text_block_delta': body_text_block_delta or 0,
            'generated_docx_artifacts_local_only': _local_generated_paths_are_local_only(
                generated_paths),
            'default_conversion_unchanged': True,
        },
        'diagnostic_warnings': diagnostic_warnings,
        'safety_warnings': safety_warnings,
    }


def _local_corpus_docx_smoke_summary_report(sample_reports):
    sample_reports = list(sample_reports or [])
    return {
        'enabled': True,
        'policy': 'local_corpus_internal_filtered_docx_smoke_only',
        'summary': {
            'sample_count': len(sample_reports),
            'passed_count': sum(
                1 for report in sample_reports
                if report.get('status') == 'passed'),
            'bounded_subset_passed_count': sum(
                1 for report in sample_reports
                if report.get('status') == 'bounded_subset_passed'),
            'blocked_count': sum(
                1 for report in sample_reports
                if report.get('status') == 'blocked'),
            'skipped_count': sum(
                1 for report in sample_reports
                if report.get('status') == 'skipped'),
            'total_removed_approved_count': sum(
                int((report.get('summary') or {}).get(
                    'removed_approved_header_footer_page_number_count', 0))
                for report in sample_reports),
            'total_true_residual_header_footer_pollution_count': sum(
                int((report.get('summary') or {}).get(
                    'true_residual_header_footer_pollution_count', 0))
                for report in sample_reports),
            'total_body_text_loss_warning_count': sum(
                int((report.get('summary') or {}).get(
                    'body_text_loss_warning_count', 0))
                for report in sample_reports),
            'total_table_text_loss_warning_count': sum(
                int((report.get('summary') or {}).get(
                    'table_text_loss_warning_count', 0))
                for report in sample_reports),
        },
        'samples': sample_reports,
        'recommendation': {
            'safe_for_phase_3e': all(
                report.get('status') in {'passed', 'bounded_subset_passed'}
                for report in sample_reports),
            'reason': (
                'Local corpus smoke validation has no fail-closed blockers.'
                if all(
                    report.get('status') in {'passed', 'bounded_subset_passed'}
                    for report in sample_reports) else
                'Resolve blocked or skipped local corpus smoke samples before Phase 3E.'),
        },
    }


def _local_docx_metric(paragraphs=10, tables=1, size=1200):
    return {
        'exists': True,
        'size': size,
        'paragraph_count': paragraphs,
        'table_count': tables,
        'all_text': 'local corpus body text',
    }


def _local_internal_report(removed=0, body_text_block_delta=0):
    return {
        'applied_to_parse': bool(removed),
        'summary': {
            'removed_raw_block_count': removed,
            'body_text_block_delta': body_text_block_delta,
        },
        'safety_warnings': [],
    }


def _local_docx_metric_exists(metrics):
    return bool(
        metrics.get('exists') and
        (
            metrics.get('size', 0) or
            metrics.get('size_bytes', 0) or
            metrics.get('file_size_bytes', 0)))


def _local_docx_metric_int(metrics, key):
    try:
        return int(metrics.get(key, 0) or 0)
    except (TypeError, ValueError):
        return 0


def _local_generated_paths_are_local_only(paths):
    for path in paths or []:
        normalized = str(path).replace('\\', '/')
        if normalized.startswith('/tmp/'):
            continue
        if normalized.startswith('local_reports/') or '/local_reports/' in normalized:
            continue
        return False
    return True


def _body_text_fragments(layout):
    fragments = []
    for page in (layout or {}).get('pages', []) or []:
        for block in page.get('text_blocks', []) or []:
            if block.get('region') != REGION_BODY:
                continue
            text = normalize_text(block.get('text', '')).lower()
            if text:
                fragments.append(text)
    return fragments


def _removed_text_fragments(filtering_report):
    fragments = []
    for page in (filtering_report or {}).get('pages', []) or []:
        for block in page.get('removed_blocks', []) or []:
            text = normalize_text(
                block.get('short_preview') or
                block.get('text') or '')
            text = text.lower()
            if text and text not in fragments:
                fragments.append(text)
    return fragments


def _path_inside_root(path, root):
    try:
        Path(path).resolve().relative_to(Path(root).resolve())
        return True
    except ValueError:
        return False


def _synthetic_body_text_validation(
        layout,
        filtering_report,
        baseline_body_text_block_count=None,
        filtered_body_text_block_count=None):
    original_signature = _body_text_signature(layout.get('pages', []))
    filtered_pages = filtering_report.get('filtered_pages', []) or layout.get('pages', [])
    filtered_signature = _body_text_signature(filtered_pages)
    signature_preserved = original_signature == filtered_signature
    count_delta = None
    classification = 'signature_preserved' if signature_preserved else 'unsafe_body_text_loss'
    if (
            baseline_body_text_block_count is not None and
            filtered_body_text_block_count is not None):
        count_delta = filtered_body_text_block_count - baseline_body_text_block_count
        if count_delta and signature_preserved:
            classification = 'acceptable_boundary_or_grouping_shift'
        elif count_delta:
            classification = 'unsafe_body_text_loss'
    return {
        'original_signature': original_signature,
        'filtered_signature': filtered_signature,
        'body_text_signature_preserved': signature_preserved,
        'body_text_block_delta': count_delta,
        'body_text_block_delta_classification': classification,
    }


def _body_text_signature(pages):
    return normalize_text(' '.join(
        normalize_text(block.get('text', '')).lower()
        for page in pages or []
        for block in page.get('text_blocks', []) or []
        if block.get('region') == REGION_BODY
    ))


def _removed_region_count(filtering_report, region):
    return sum(
        1
        for page in filtering_report.get('pages', []) or []
        for block in page.get('removed_blocks', []) or []
        if block.get('region') == region)


def _removed_text_match_count(filtering_report, text):
    needle = normalize_text(text).lower()
    return sum(
        1
        for page in filtering_report.get('pages', []) or []
        for block in page.get('removed_blocks', []) or []
        if needle in normalize_text(block.get('text_preview', '')).lower())


def _corpus_warning_types(report):
    return {warning['type'] for warning in report['warnings']}


def _filtered_docx_experiment_report():
    return {
        'summary': {
            'baseline_raw_block_count': 790,
            'filtered_raw_block_count': 742,
            'removed_raw_block_count': 48,
            'baseline_parsed_text_block_count': 523,
            'filtered_parsed_text_block_count': 486,
            'baseline_body_text_block_count': 393,
            'filtered_body_text_block_count': 393,
            'baseline_table_count': 139,
            'filtered_table_count': 127,
            'baseline_image_count': 0,
            'filtered_image_count': 0,
            'baseline_section_count': 50,
            'filtered_section_count': 50,
            'body_region_removed_count': 0,
            'rejected_unsure_layout_placeholder_removed_count': 0,
        },
        'header_footer_pollution_reduction': {
            'removed_header_footer_page_number_count': 48,
            'parsed_text_block_delta': 37,
            'body_text_block_delta': 0,
        },
    }


def _passed_table_visual_gate_report():
    return {
        'gate_status': 'passed',
        'summary': {
            'gate_status': 'passed',
            'expected_review_item_count': 8,
            'parsed_review_item_count': 8,
            'approve_count': 8,
            'reject_count': 0,
            'unsure_count': 0,
            'missing_decision_count': 0,
        },
    }


def _docx_metrics(size=1200, paragraphs=10, tables=2):
    return {
        'exists': True,
        'size_bytes': size,
        'paragraph_count': paragraphs,
        'table_count': tables,
    }


def _normal_conversion_check(passed=True):
    return {
        'passed': passed,
        'state_restored_or_reloaded': passed,
        'message': '' if passed else 'conversion failed',
    }


def _write_docx_fixture(path, body_paragraphs=None, table_cells=None, header_texts=None, footer_texts=None):
    body_paragraphs = body_paragraphs or []
    table_cells = table_cells or []
    header_texts = header_texts or []
    footer_texts = footer_texts or []
    document_xml = _docx_document_xml(body_paragraphs, table_cells)
    with zipfile.ZipFile(path, 'w') as archive:
        archive.writestr('word/document.xml', document_xml)
        for index, text in enumerate(header_texts, start=1):
            archive.writestr(f'word/header{index}.xml', _docx_part_xml([text]))
        for index, text in enumerate(footer_texts, start=1):
            archive.writestr(f'word/footer{index}.xml', _docx_part_xml([text]))


def _docx_document_xml(paragraphs, table_cells):
    body = ''.join(_docx_paragraph_xml(text) for text in paragraphs)
    if table_cells:
        cells = ''.join(
            f'<w:tc>{_docx_paragraph_xml(text)}</w:tc>'
            for text in table_cells)
        body += f'<w:tbl><w:tr>{cells}</w:tr></w:tbl>'
    body += '<w:sectPr/>'
    return _docx_part_xml([], body_xml=body)


def _docx_part_xml(paragraphs, body_xml=None):
    content = body_xml if body_xml is not None else ''.join(
        _docx_paragraph_xml(text) for text in paragraphs)
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:body>{content}</w:body>'
        '</w:document>'
    )


def _docx_paragraph_xml(text):
    return f'<w:p><w:r><w:t>{text}</w:t></w:r></w:p>'


def _table_record(
        table_id,
        page_index,
        region,
        bbox,
        rows=2,
        cols=2,
        cells=None,
        cell_texts=None,
        cell_bboxes=None,
        include_cells=True):
    cells = cells if cells is not None else (
        len(cell_texts) if cell_texts is not None else rows * cols)
    record = {
        'table_id': table_id,
        'page_index': page_index,
        'page_number': page_index + 1,
        'table_index': 0,
        'bbox': list(bbox),
        'region': region,
        'row_count': rows,
        'column_count': cols,
        'cell_count': cells,
        'non_empty_cell_count': cells,
        'table_type': 'stream',
        'text_preview': 'short table preview',
    }
    if include_cells:
        texts = list(cell_texts) if cell_texts is not None else [
            f'cell {index + 1}' for index in range(cells)
        ]
        bboxes = list(cell_bboxes) if cell_bboxes is not None else _table_cell_bboxes(
            bbox, rows, cols, cells)
        cell_summaries = []
        for index in range(cells):
            text = texts[index] if index < len(texts) else ''
            cell_bbox = bboxes[index] if index < len(bboxes) else bbox
            cell_summaries.append({
                'row_index': index // max(cols, 1),
                'column_index': index % max(cols, 1),
                'bbox': list(cell_bbox),
                'text_preview': text,
            })
        record['cell_summaries'] = cell_summaries
        record['cell_text_signature'] = [cell.get('text_preview', '') for cell in cell_summaries]
        record['cell_bbox_signature'] = [list(cell.get('bbox', [])) for cell in cell_summaries]
    return record


def _table_cell_bboxes(bbox, rows, cols, cells):
    left, top, right, bottom = [float(value) for value in bbox]
    rows = max(int(rows or 1), 1)
    cols = max(int(cols or 1), 1)
    width = (right - left) / cols
    height = (bottom - top) / rows
    bboxes = []
    for index in range(cells):
        row = index // cols
        col = index % cols
        bboxes.append([
            round(left + col * width, 2),
            round(top + row * height, 2),
            round(left + (col + 1) * width, 2),
            round(top + (row + 1) * height, 2),
        ])
    return bboxes


def _removed_objects_page(objects, page_index=0):
    return {
        'page_index': page_index,
        'page_number': page_index + 1,
        'removed_count': len(objects),
        'objects': objects,
    }


def _removed_raw_object(text, region, bbox, role):
    return {
        'raw_object_id': f'raw-{text}',
        'page_index': 0,
        'block_index': 0,
        'candidate_id': f'candidate-{text}',
        'proposed_role': role,
        'region': region,
        'bbox': list(bbox),
        'text': text,
        'text_preview': text,
    }


def _summary_fingerprint(text, region):
    return make_text_fingerprint(text, y_band=region)['key']


def _raw_object(block_index, text, region):
    bbox_by_region = {
        REGION_TOP: [50.0, 20.0, 300.0, 40.0],
        REGION_BODY: [50.0, 400.0, 520.0, 420.0],
        REGION_BOTTOM: [50.0, 960.0, 300.0, 980.0],
    }
    bbox = list(bbox_by_region[region])
    bbox[1] += block_index * 0.2
    bbox[3] += block_index * 0.2
    return {
        'raw_object_id': f'raw-{block_index}',
        'object_index': block_index,
        'block_index': block_index,
        'object_type': 'Line',
        'fingerprint': _summary_fingerprint(text, region),
        'region': region,
        'text': text,
        'bbox': bbox,
    }


def _mapping_summary_block(block_index, fingerprint, region, text):
    block = _raw_object(block_index, text, region)
    return {
        'block_index': block_index,
        'fingerprint': fingerprint,
        'normalized_text': normalize_page_number(text).lower(),
        'region': region,
        'text': text,
        'bbox': block['bbox'],
    }


class _FakePage:

    def __init__(self, page_id):
        self.id = page_id


class _FakeRawPage:

    def __init__(self, raw_objects):
        self.width = 600
        self.height = 1000
        self.blocks = [_FakeRawBlock(raw_object) for raw_object in raw_objects]


class _FakeRawBlock:

    def __init__(self, raw_object):
        self.text = raw_object['text']
        self.bbox = raw_object['bbox']
        self.spans = []


def _summary_block(block_index, fingerprint, region, text):
    return {
        'block_index': block_index,
        'fingerprint': fingerprint,
        'region': region,
        'text': text,
        'bbox': [50.0, 50.0 + block_index * 20.0, 520.0, 65.0 + block_index * 20.0],
    }


def _dry_run_candidate(candidate_id, fingerprint, role, action, pages, regions):
    return {
        'candidate_id': candidate_id,
        'fingerprint': fingerprint,
        'proposed_role': role,
        'action': action,
        'affected_pages': pages,
        'regions': regions,
    }


def _review_decision(candidate_id, fingerprint, manual_decision):
    return {
        'candidate_id': candidate_id,
        'fingerprint': fingerprint,
        'manual_decision': manual_decision,
    }


def _insertion_points_by_id(analysis):
    return {
        point['candidate_id']: point
        for point in analysis['insertion_points']
    }


def _dry_run_by_fingerprint(report):
    return {
        candidate['fingerprint']: candidate
        for candidate in report['header_footer_exclusion_dry_run']['candidates']
    }


def _review_markdown(candidate_id, fingerprint, role, action, decision):
    markers = {
        'approve_exclude': '[ ]',
        'reject_exclude': '[ ]',
        'unsure': '[ ]',
    }
    markers[decision] = '[x]'
    return '\n'.join([
        f'### {candidate_id} | {role} | {action}',
        f'- fingerprint: `{fingerprint}`',
        '- human_decision: '
        f'approve_exclude: {markers["approve_exclude"]}    '
        f'reject_exclude: {markers["reject_exclude"]}    '
        f'unsure: {markers["unsure"]}',
    ])


def _unsafe_diff_report_for_removed_blocks(pages, removed_keys):
    removed_by_page = []
    for page in pages:
        page_removed = []
        page_index = page['page_index']
        for block in page['text_blocks']:
            if (page_index, block['block_index']) in set(removed_keys):
                page_removed.append({
                    'page_index': page_index,
                    'page_number': page_index + 1,
                    'block_index': block['block_index'],
                    'candidate_id': 'unsafe-body',
                    'fingerprint': block['fingerprint'],
                    'proposed_role': 'header',
                    'manual_decision': 'reject_exclude',
                    'explicit_approval': False,
                    'region': block['region'],
                    'reason': 'unsafe synthetic removal',
                    'short_preview': block['text'],
                })
        removed_by_page.append({
            'page_index': page_index,
            'page_number': page_index + 1,
            'removed_count': len(page_removed),
            'blocks': page_removed,
        })

    return {
        'enabled': True,
        'summary': {
            'original_block_count': sum(len(page['text_blocks']) for page in pages),
        },
        'removed_blocks_by_page': removed_by_page,
        'kept_blocks_by_page': [],
        'removed_blocks_by_candidate': [],
        'retained_candidates': [],
        'safety': {'warnings': []},
    }


if __name__ == '__main__':
    unittest.main()
